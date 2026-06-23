"""
proposals/docx_forms.py

DOCX generation for downloadable proposals.

Design goal:
- DO NOT "rebuild" tables/cells (that destroys Word formatting).
- Instead, overwrite existing template paragraphs/runs and (only when needed)
  clone existing paragraphs to preserve list bullets/numbering, indents, fonts,
  spacing, and alignment.

This file expects the DOCX templates to exist in:
    proposals/template_files/
        form1_program_template.docx
        form1_project_template.docx
        form2_training_design_template.docx
"""
from __future__ import annotations

import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import List, Optional, Sequence, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph

# --- Optional imports (kept defensive) ---
try:
    from accounts.models import Signatory  # type: ignore
except Exception:  # pragma: no cover
    Signatory = None  # type: ignore

# SDG code->title (stable list)
SDG_TITLES = {
    "01": "No Poverty",
    "02": "Zero Hunger",
    "03": "Good Health and Well-being",
    "04": "Quality Education",
    "05": "Gender Equality",
    "06": "Clean Water and Sanitation",
    "07": "Affordable and Clean Energy",
    "08": "Decent Work and Economic Growth",
    "09": "Industry, Innovation and Infrastructure",
    "10": "Reduced Inequalities",
    "11": "Sustainable Cities and Communities",
    "12": "Responsible Consumption and Production",
    "13": "Climate Action",
    "14": "Life Below Water",
    "15": "Life on Land",
    "16": "Peace, Justice and Strong Institutions",
    "17": "Partnerships for the Goals",
}


# =============================
# Low-level Word helpers
# =============================

def _has_num_pr(paragraph: Paragraph) -> bool:
    """True if paragraph is part of a numbered/bulleted list."""
    p_pr = paragraph._p.get_or_add_pPr()
    return p_pr.find(qn("w:numPr")) is not None


def _clear_num_pr(paragraph: Paragraph) -> None:
    """Remove Word list numbering from a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def _set_paragraph_text_keep_style(paragraph: Paragraph, text: str) -> None:
    """
    Set text while preserving paragraph formatting and run styling.

    - Reuses the first run; clears remaining runs.
    - If there are no runs, adds one run (style comes from paragraph).
    """
    text = "" if text is None else str(text)

    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _remove_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph from the document XML (safe for table cells too)."""
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _clone_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """
    Deep-clone a paragraph element and insert it after `paragraph`.
    Preserves numbering/bullets/indents/spacings/fonts from the template.
    """
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _insert_paragraph_clone_after(cursor: Paragraph, source: Paragraph) -> Paragraph:
    """
    Insert a deep-clone of `source` right after `cursor`.
    Use this when you must preserve a specific prototype's formatting
    (e.g., list numbering) rather than cloning the cursor's formatting.
    """
    new_p = deepcopy(source._p)
    cursor._p.addnext(new_p)
    return Paragraph(new_p, cursor._parent)


def _ensure_paragraphs_count(
    paragraphs: List[Paragraph],
    wanted: int,
    *,
    prototype: Optional[Paragraph] = None
) -> List[Paragraph]:
    """
    Ensure list has exactly `wanted` paragraphs.
    - If more: remove extras from end.
    - If fewer: clone from `prototype` or last paragraph.
    Returns the updated list.
    """
    if wanted < 1:
        wanted = 1

    # shrink
    while len(paragraphs) > wanted:
        _remove_paragraph(paragraphs[-1])
        paragraphs.pop()

    # grow
    proto = prototype or (paragraphs[-1] if paragraphs else None)
    while len(paragraphs) < wanted:
        if proto is None:
            break
        newp = _clone_paragraph_after(paragraphs[-1])
        _set_paragraph_text_keep_style(newp, "")
        paragraphs.append(newp)

    return paragraphs


def _find_row_index_by_keyword(table, keyword: str) -> Optional[int]:
    key = (keyword or "").strip().lower()
    if not key:
        return None
    for idx, row in enumerate(table.rows):
        left = (row.cells[0].text or "").strip().lower()
        if key in left:
            return idx
    return None


def _safe_int(v: object, default: int = 0) -> int:
    try:
        return int(v)  # type: ignore[arg-type]
    except Exception:
        return default


# =============================
# Data access helpers (Proposal)
# =============================


def _normalize_text_preserve_newlines(s: str, *, keep_trailing_blank_lines: bool = False) -> str:
    """
    Normalize whitespace without destroying intentional line breaks.

    - Converts CR/Excel markers into \n (handles _x000D_ and variants)
    - Collapses multiple spaces *within each line*
    - Trims trailing blank lines
    - Splits bullet '•' tokens into separate lines for readability
    """
    s = "" if s is None else str(s)

    # Common Excel/OOXML "carriage return" markers (sometimes appear as literal text)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("_x000D_", "\n").replace("\\x000D", "\n")
    s = re.sub(r"(?i)_?x000d_?", "\n", s)

    # Put each bullet on its own line (helps Work Plan / Gender Issues readability)
    if "•" in s:
        # ensure bullets after the first start on a new line
        s = re.sub(r"\s*•\s*", "\n• ", s)
        if s.startswith("\n• "):
            s = s[1:]

    # Normalize common "Others:-" artifact
    s = re.sub(r"(?i)\bothers\s*:\s*-\s*", "Others: ", s)

    lines = [" ".join(line.split()) for line in s.split("\n")]

    # remove trailing empty lines unless explicitly requested (Work Plan alignment)
    if not keep_trailing_blank_lines:
        while lines and lines[-1] == "":
            lines.pop()

    return "\n".join(lines)

def _get_related_list(obj, attr: str) -> List:
    rel = getattr(obj, attr, None)
    if rel is None:
        return []
    try:
        return list(rel.all())
    except Exception:
        try:
            return list(rel)
        except Exception:
            return []


def _get_proponents(proposal) -> List:
    props = _get_related_list(proposal, "proponents")
    try:
        props.sort(key=lambda x: getattr(x, "id", 0))
    except Exception:
        pass
    return props


def _get_program_projects(proposal) -> List:
    prjs = _get_related_list(proposal, "program_projects")
    try:
        prjs.sort(key=lambda x: (_safe_int(getattr(x, "order", 0)), _safe_int(getattr(x, "id", 0))))
    except Exception:
        pass
    return prjs


def _get_specific_objectives_for_project(proposal, prj) -> List[str]:
    qs = getattr(proposal, "specific_objectives", None)
    if qs is None:
        return []
    try:
        items = qs.filter(program_project=prj).values_list("objective", flat=True)
        return [str(x).strip() for x in items if str(x).strip()]
    except Exception:
        try:
            items = qs.filter(program_project=prj)
            return [
                str(getattr(x, "objective", "")).strip()
                for x in items
                if str(getattr(x, "objective", "")).strip()
            ]
        except Exception:
            return []


def _get_methodology_items(proposal) -> List[str]:
    items = _get_related_list(proposal, "methodologies")
    out = []
    for x in items:
        val = getattr(x, "item", "") if hasattr(x, "item") else str(x)
        val = (val or "").strip()
        if val:
            out.append(val)
    return out


def _get_output_items(proposal) -> List[str]:
    items = _get_related_list(proposal, "output_outcomes")
    out = []
    for x in items:
        val = getattr(x, "item", "") if hasattr(x, "item") else str(x)
        val = (val or "").strip()
        if val:
            out.append(val)
    return out


def _get_sdg_lines(proposal) -> List[str]:
    links = _get_related_list(proposal, "sdg_links")
    out = []
    for item in links:
        code = (getattr(item, "sdg_code", "") or "").strip()
        if not code:
            continue
        code = code.zfill(2)
        title = SDG_TITLES.get(code, code)
        expl = (getattr(item, "explanation", "") or "").strip()
        if expl:
            out.append(f"SDG {code} ({title}) – {expl}")
        else:
            out.append(f"SDG {code} ({title})")
    return out


def _get_thrust_lines(proposal) -> List[str]:
    links = _get_related_list(proposal, "thrust_links")
    out = []
    for item in links:
        name = (getattr(item, "thrust_name", "") or "").strip()
        if not name:
            continue
        expl = (getattr(item, "explanation", "") or "").strip()
        out.append(f"{name} – {expl}" if expl else name)
    return out


def _get_gender_issues(proposal) -> Tuple[set, str]:
    links = _get_related_list(proposal, "gender_issue_links")
    keys = set()
    others = ""
    for item in links:
        k = (getattr(item, "issue_key", "") or "").strip()
        if not k:
            continue
        keys.add(k)
        if k == "others":
            others = (getattr(item, "other_text", "") or "").strip()
    return keys, others


def _strip_leading_bullets(text: str) -> str:
    lines = []
    for line in (text or "").splitlines():
        stripped = line.strip()
        while stripped.startswith(("-", "•", "*", "·")):
            stripped = stripped[1:].strip()
        lines.append(stripped)
    return "\n".join(lines).strip()


def _clean_others_text(text: str) -> str:
    lines = []
    for line in (text or "").splitlines():
        stripped = line.strip()
        content = stripped
        while content.startswith(("-", "•", "*", "·")):
            content = content[1:].strip()
        if not content:
            continue
        if "mandates placed in others" in content.lower():
            continue
        lines.append(stripped)
    return "\n".join(lines).strip()


def _file_basename(file_field) -> str:
    """Return only the uploaded file name for a Django FileField-like value."""
    name = (getattr(file_field, "name", "") or "").strip()
    if not name:
        return ""
    return Path(name).name


def _proposal_file_lines(proposal, pairs: Sequence[Tuple[str, str]]) -> List[str]:
    lines = []
    for label, attr in pairs:
        filename = _file_basename(getattr(proposal, attr, None))
        if filename:
            lines.append(f"{label}: {filename}")
    return lines


def _normalize_extracted_lines(lines: Sequence[str], *, max_chars: int = 30000) -> List[str]:
    out = []
    used = 0
    for raw in lines:
        line = " ".join(str(raw or "").replace("\r", "\n").split())
        if not line:
            continue
        remaining = max_chars - used
        if remaining <= 0:
            break
        if len(line) > remaining:
            out.append(line[:remaining].rstrip())
            used = max_chars
            break
        out.append(line)
        used += len(line) + 1
    return out




def _fmt_excel_value(value: object) -> str:
    """
    Format Excel cell values for display.

    - Adds thousand separators for numeric values (e.g., 1000000 -> 1,000,000).
    - Keeps 2 decimals for non-integer floats (e.g., 1234.5 -> 1,234.50).
    - Preserves intended line breaks for multi-line cells (Topic/Persons, Gender Issues, etc.)
      while removing OOXML artifacts like _x000D_.
    """
    if value is None:
        return ""
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"

    try:
        import decimal
        if isinstance(value, decimal.Decimal):
            value = float(value)
    except Exception:
        pass

    if isinstance(value, int):
        return f"{value:,}"
    if isinstance(value, float):
        if value.is_integer():
            return f"{int(value):,}"
        return f"{value:,.2f}"

    return _normalize_text_preserve_newlines(str(value), keep_trailing_blank_lines=True)

def _drop_leading_title_rows(rows: List[List[str]], *, sheet_title: str = "") -> List[List[str]]:
    """
    Remove leading one-cell title rows often present in uploaded XLSX files,
    e.g. 'Line-item Budget', 'Work Plan', 'Gantt Chart', and also a row that
    exactly matches the worksheet title.

    This prevents the generated DOCX from showing those titles at the top of
    the inserted table.
    """
    title_set = {
        "work plan",
        "gantt chart",
        "gantt",
        "line-item budget",
        "line item budget",
        "line-item",
        "budget",
    }
    sheet_key = (sheet_title or "").strip().lower()

    cleaned = list(rows)
    while cleaned:
        first = cleaned[0]
        non_empty = [str(x or "").strip() for x in first if str(x or "").strip()]
        if not non_empty:
            cleaned.pop(0)
            continue
        if len(non_empty) == 1:
            t = non_empty[0].lower()
            if t.startswith("sheet:") or t in title_set or (sheet_key and t == sheet_key):
                cleaned.pop(0)
                continue
        break
    return cleaned

def _extract_xlsx_lines(path: str) -> List[str]:
    """
    Extract a worksheet as plain lines.

    IMPORTANT:
    - Does NOT include file name or sheet name.
    - Formats numeric amounts with thousand separators.
    """
    try:
        from openpyxl import load_workbook
    except Exception:
        return []

    try:
        wb = load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return []

    lines: List[str] = []
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                values = [_fmt_excel_value(v) for v in row]
                while values and not values[-1]:
                    values.pop()
                if any(values):
                    lines.append(" | ".join(values))
    finally:
        try:
            wb.close()
        except Exception:
            pass

    return _normalize_extracted_lines(lines)

def _extract_xlsx_tables(
    path: str,
    *,
    max_rows: int = 80,
    max_cols: int = 12
) -> List[Tuple[str, List[List[str]], List[float]]]:
    """
    Extract worksheets as 2D tables for insertion into the DOCX.

    IMPORTANT:
    - Does NOT add sheet name labels in the output.
    - Formats numeric amounts with thousand separators.
    - Preserves explicit line breaks within cell text.
    - Drops common leading title rows like 'Line-item Budget'.
    - Captures Excel column widths so we can approximate Word column widths
      (helps keep Topic/Activity and Persons Involved aligned by line).
    """
    try:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter
    except Exception:
        return []

    try:
        wb = load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return []

    tables: List[Tuple[str, List[List[str]], List[float]]] = []

    def is_instruction_row(values: Sequence[str]) -> bool:
        joined = " ".join(str(v or "").strip() for v in values if str(v or "").strip()).lower()
        if not joined:
            return False
        instruction_prefixes = (
            "tip:",
            "tips:",
            "note:",
            "notes:",
            "instruction:",
            "instructions:",
            "reminder:",
        )
        return joined.startswith(instruction_prefixes)

    try:
        for ws in wb.worksheets:
            # Capture Excel column widths (character-based units)
            excel_widths: List[float] = []
            for ci in range(1, max_cols + 1):
                letter = get_column_letter(ci)
                try:
                    w = ws.column_dimensions[letter].width
                except Exception:
                    w = None
                excel_widths.append(float(w) if w else 10.0)

            rows: List[List[str]] = []
            for row in ws.iter_rows(values_only=True):
                values: List[str] = []
                for value in row[:max_cols]:
                    values.append(_fmt_excel_value(value))
                while values and not values[-1]:
                    values.pop()
                if any(values) and not is_instruction_row(values):
                    rows.append(values)
                if len(rows) >= max_rows:
                    break

            rows = _drop_leading_title_rows(rows, sheet_title=ws.title)

            if rows:
                width = max(len(r) for r in rows)
                width = max(1, min(width, max_cols))
                normalized = [(r + [""] * width)[:width] for r in rows]
                tables.append((ws.title, normalized, excel_widths[:width]))
    finally:
        try:
            wb.close()
        except Exception:
            pass

    return tables

def _extract_pdf_lines(path: str) -> List[str]:
    try:
        from pypdf import PdfReader
    except Exception:
        return []

    try:
        reader = PdfReader(path)
    except Exception:
        return []

    lines = []
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        page_lines = [line.strip() for line in text.splitlines() if line.strip()]
        if page_lines:
            lines.append(f"Page {page_no}")
            lines.extend(page_lines)
    return _normalize_extracted_lines(lines)


def _extract_uploaded_file_lines(file_field) -> List[str]:
    """
    Extract text lines from an uploaded file field.

    IMPORTANT:
    - Do NOT include the file name at the top.
    - For XLSX, do NOT include sheet names.
    """
    filename = _file_basename(file_field)
    if not filename:
        return []

    path = getattr(file_field, "path", "") or ""
    suffix = Path(filename).suffix.lower()

    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"} and path:
        lines = _extract_xlsx_lines(path)
    elif suffix == ".pdf" and path:
        lines = _extract_pdf_lines(path)
    else:
        lines = []

    if not lines:
        return ["Content could not be extracted."]

    return lines

def _uploaded_content_section(label: str, file_field) -> List[str]:
    lines = _extract_uploaded_file_lines(file_field)
    if not lines:
        return []
    return [f"{label}:", *lines, ""]

def _to_roman(n: int) -> str:
    vals = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    out = []
    for v, s in vals:
        while n >= v:
            out.append(s)
            n -= v
    return "".join(out)


def _compose_role(proposal, proponent) -> str:
    """
    Combine base role + phase leader roles when applicable.
    Output example:
        "Program Leader / Phase III Project Leader"
    """
    base = (getattr(proponent, "role", "") or "").strip()
    roles = [base] if base else []

    user = getattr(proponent, "user", None)
    user_id = getattr(user, "id", None)
    if user_id:
        prjs = _get_program_projects(proposal)
        phases = []
        for prj in prjs:
            leader_id = (
                getattr(prj, "leader_user_id", None)
                or getattr(getattr(prj, "leader_user", None), "id", None)
            )
            if leader_id and leader_id == user_id:
                order_no = _safe_int(getattr(prj, "order", 0), 0)
                roman = _to_roman(order_no) if order_no else ""
                phases.append(f"Phase {roman} Project Leader" if roman else "Project Leader")

        for ph in phases:
            if ph and ph not in roles:
                roles.append(ph)

    # dedupe preserve order
    clean = []
    seen = set()
    for r in roles:
        k = r.lower()
        if k in seen:
            continue
        seen.add(k)
        clean.append(r)

    return " / ".join(clean)


# =============================
# Fill logic: Form 1 (Program)
# =============================

def _fill_cell_single_paragraph(cell, text: str) -> None:
    paras = list(cell.paragraphs)
    if not paras:
        cell.add_paragraph()
        paras = list(cell.paragraphs)
    _set_paragraph_text_keep_style(paras[0], text)
    for p in paras[1:]:
        _set_paragraph_text_keep_style(p, "")



def _fill_cell_single_paragraph_strict(cell, text: str, *, keep_trailing_blank_lines: bool = False) -> None:
    """
    Write text into a cell used in GENERATED content (inserted tables / merged cells)
    without leaving behind extra empty paragraphs (which cause visible gaps).

    IMPORTANT (alignment):
    - We preserve \n line breaks as *manual line breaks inside a single paragraph*
      (not multiple paragraphs). This matches Excel's wrapped/newline behavior much better
      and keeps Topic/Persons line alignment stable even when font is Arial 12.
    - For Work Plan alignment, pass keep_trailing_blank_lines=True so trailing blank
      lines used as padding are preserved.

    Formatting:
    - Paragraph spacing before/after = 0
    - Single line spacing
    - Cell vertical alignment = TOP
    """
    text = _normalize_text_preserve_newlines(text, keep_trailing_blank_lines=keep_trailing_blank_lines)

    # Split into lines; preserve trailing blanks if requested
    if text is None:
        lines = [""]
    else:
        lines = str(text).split("\n")

    if not keep_trailing_blank_lines:
        while lines and lines[-1] == "":
            lines.pop()
    if not lines:
        lines = [""]

    # Clear everything except tcPr, then rebuild a SINGLE paragraph with line breaks
    _clear_cell_content(cell)
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]

    # clear runs first
    for r in p.runs:
        r.text = ""

    # Apply tight paragraph formatting to remove gaps
    try:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    except Exception:
        pass

    run = p.add_run("")
    for i, line in enumerate(lines):
        if i > 0:
            run.add_break()  # manual line break
        if line:
            run.add_text(line)

    # Cell vertical alignment top (prevents extra vertical centering)
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    except Exception:
        pass

def _fill_cell_paragraphs(cell, items: Sequence[str]) -> None:
    items = [str(item) for item in (items or []) if str(item).strip()]
    paras = list(cell.paragraphs)
    if not paras:
        cell.add_paragraph()
        paras = list(cell.paragraphs)

    if not items:
        _fill_cell_single_paragraph(cell, "")
        return

    paras = _ensure_paragraphs_count(paras, len(items), prototype=paras[0])
    for idx, p in enumerate(paras):
        _set_paragraph_text_keep_style(p, items[idx] if idx < len(items) else "")


def _clear_cell_content(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    cell.add_paragraph()


def _apply_proposal_font_to_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)


def _apply_proposal_font_to_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        _apply_proposal_font_to_paragraph(paragraph)
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                _apply_proposal_font_to_cell(nested_cell)


def _add_heading_paragraph(cell, text: str) -> None:
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)


def _add_lines_to_cell(cell, lines: Sequence[str]) -> None:
    for line in lines:
        if not str(line).strip():
            continue
        p = cell.add_paragraph(str(line))
        _apply_proposal_font_to_paragraph(p)


def _merge_full_row(table, row_idx: int) -> None:
    if row_idx < 0 or row_idx >= len(table.rows):
        return
    cells = table.rows[row_idx].cells
    if len(cells) <= 1:
        return

    parts: List[str] = []
    for c in cells:
        t = _normalize_text_preserve_newlines(c.text or "")
        if t:
            parts.append(t)
    keep = "\n".join(parts).strip()

    cells[0].merge(cells[-1])
    _fill_cell_single_paragraph_strict(cells[0], keep)

def _merge_row_from_col(table, row_idx: int, start_col: int) -> None:
    if row_idx < 0 or row_idx >= len(table.rows):
        return
    cells = table.rows[row_idx].cells
    if start_col < 0 or start_col >= len(cells) - 1:
        return

    parts: List[str] = []
    for c in cells[start_col:]:
        t = _normalize_text_preserve_newlines(c.text or "")
        if t:
            parts.append(t)
    keep = "\n".join(parts).strip()

    cells[start_col].merge(cells[-1])
    _fill_cell_single_paragraph_strict(cells[start_col], keep)

def _center_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _format_inserted_workbook_table(table, table_kind: str) -> None:
    kind = (table_kind or "").lower()

    def non_empty_texts(row) -> List[str]:
        out = []
        for c in row.cells:
            t = " ".join((c.text or "").split())
            if t:
                out.append(t)
        return out

    def row_has_only_first_cell(row) -> bool:
        first = " ".join((row.cells[0].text or "").split())
        if not first:
            return False
        others_have = any(" ".join((c.text or "").split()) for c in row.cells[1:])
        return not others_have

    if kind == "work_plan":
        # Merge only true section headers (single-cell rows / Phase headings).
        for row_idx, row in enumerate(table.rows):
            non = non_empty_texts(row)
            if not non:
                _merge_full_row(table, row_idx)
                continue
            if len(non) <= 1:
                _merge_full_row(table, row_idx)
                continue

            first = " ".join((row.cells[0].text or "").split()).lower()
            if ("phase" in first) and row_has_only_first_cell(row):
                _merge_full_row(table, row_idx)

    elif kind == "line_item_budget":
        # In budget templates, "Phase ..." rows are meant to span the full row.
        for row_idx, row in enumerate(table.rows):
            non = non_empty_texts(row)
            if not non:
                _merge_full_row(table, row_idx)
                continue
            if len(non) <= 1:
                _merge_full_row(table, row_idx)
                continue

            first = " ".join((row.cells[0].text or "").split()).lower()
            if first.startswith("phase") and row_has_only_first_cell(row):
                _merge_full_row(table, row_idx)

            # Also merge section headers like "Supplies & Materials" when placed only in col 0
            if row_has_only_first_cell(row) and len(first) > 0:
                # heuristic: if col0 has text and other cols empty, merge
                _merge_full_row(table, row_idx)

    elif kind == "gantt_chart":
        # Merge top title row only if it's a true title row (single cell).
        if table.rows:
            top_non = non_empty_texts(table.rows[0])
            if len(top_non) <= 1:
                _merge_full_row(table, 0)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = " ".join((cell.text or "").split()).lower()
                if "quarter" in text and "implementation" in text:
                    _merge_row_from_col(table, row_idx, col_idx)
                    _center_cell(row.cells[col_idx])
                    break

    # Center checkmarks/slashes
    for row in table.rows:
        for table_cell in row.cells:
            if " ".join((table_cell.text or "").split()) == "/":
                _center_cell(table_cell)

def _apply_table_column_widths(
    table,
    *,
    excel_widths: Optional[Sequence[float]],
    total_inches: float = 6.5,
    ratio: Optional[Sequence[float]] = None,
) -> None:
    """
    Set Word table column widths.

    Priority:
    1) If `ratio` is provided, apply it (e.g., [20, 40, 40]).
    2) Else if `excel_widths` is provided, approximate from Excel widths (character-based units).

    We disable autofit so wrapping is governed primarily by fixed column widths.
    """
    try:
        table.autofit = False
    except Exception:
        pass

    # 1) Explicit ratio (preferred for strict template alignment)
    if ratio:
        widths = [float(w) for w in ratio if w is not None]
        if widths:
            s = sum(widths) or 1.0
            word_widths = [Inches(total_inches * (w / s)) for w in widths]
            for ci, w in enumerate(word_widths):
                if ci >= len(table.columns):
                    break
                try:
                    table.columns[ci].width = w
                except Exception:
                    pass
                try:
                    for c in table.columns[ci].cells:
                        c.width = w
                except Exception:
                    pass
        return

    # 2) Excel-derived widths
    if not excel_widths:
        return

    widths = [float(w) if w else 1.0 for w in excel_widths]
    s = sum(widths) or 1.0
    word_widths = [Inches(total_inches * (w / s)) for w in widths]

    for ci, w in enumerate(word_widths):
        if ci >= len(table.columns):
            break
        try:
            table.columns[ci].width = w
        except Exception:
            pass
        try:
            for c in table.columns[ci].cells:
                c.width = w
        except Exception:
            pass


def _add_table_to_cell(cell, rows: Sequence[Sequence[str]], *, table_kind: str = "", excel_widths: Optional[Sequence[float]] = None) -> None:
    rows = [list(row) for row in rows if any(str(v).strip() for v in row)]
    if not rows:
        return

    # Work Plan alignment: ensure each row's cells have the same number of lines.
    # Excel templates often use blank lines to align Topic/Activity with Persons Involved.
    # Word will only align reliably if we preserve/pad line counts per row.
    if (table_kind or "").lower() == "work_plan":
        padded_rows: List[List[str]] = []
        for r in rows:
            # Normalize with trailing blanks preserved, then pad to max line count
            norm_cells = [
                _normalize_text_preserve_newlines(str(v or ""), keep_trailing_blank_lines=True)
                for v in r
            ]
            line_lists = [c.split("\n") for c in norm_cells]
            max_lines = max((len(ll) for ll in line_lists), default=1)
            new_r = []
            for ll in line_lists:
                if len(ll) < max_lines:
                    ll = ll + [""] * (max_lines - len(ll))
                new_r.append("\n".join(ll))
            padded_rows.append(new_r)
        rows = padded_rows
    cols = max(len(row) for row in rows)
    table = cell.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    kind = (table_kind or "").lower()
    ratio = None
    total_inches = 7.2 if kind == "work_plan" else 6.5

    # WORK PLAN: strict ratio requested
    if kind == "work_plan" and cols == 3:
        ratio = [20, 40, 40]

    # Gantt & Line-item Budget: widen first column for readability
    elif kind == "gantt_chart" and cols >= 2:
        first = 35.0
        rest = (100.0 - first) / float(cols - 1)
        ratio = [first] + [rest] * (cols - 1)

    elif kind == "line_item_budget" and cols >= 2:
        first = 40.0
        rest = (100.0 - first) / float(cols - 1)
        ratio = [first] + [rest] * (cols - 1)

    _apply_table_column_widths(table, excel_widths=excel_widths, ratio=ratio, total_inches=total_inches)
    for row_idx, row in enumerate(rows):
        for col_idx in range(cols):
            val = row[col_idx] if col_idx < len(row) else ""
            tc = table.cell(row_idx, col_idx)
            _fill_cell_single_paragraph_strict(tc, val, keep_trailing_blank_lines=((table_kind or '').lower() == 'work_plan'))
    _format_inserted_workbook_table(table, table_kind)
    for row in table.rows:
        for table_cell in row.cells:
            _apply_proposal_font_to_cell(table_cell)

def _add_xlsx_content_to_cell(cell, label: str, file_field, *, table_kind: str = "") -> None:
    filename = _file_basename(file_field)
    if not filename:
        return

    path = getattr(file_field, "path", "") or ""
    tables = _extract_xlsx_tables(path) if path else []
    if not tables:
        p = cell.add_paragraph("Content could not be extracted.")
        _apply_proposal_font_to_paragraph(p)
        return

    for _, rows, excel_widths in tables:
        _add_table_to_cell(cell, rows, table_kind=table_kind, excel_widths=excel_widths)


def _add_pdf_content_to_cell(cell, label: str, file_field) -> None:
    filename = _file_basename(file_field)
    if not filename:
        return

    _add_heading_paragraph(cell, label)
    # User request: do NOT show filename at the top.
    lines = _extract_uploaded_file_lines(file_field)
    _add_lines_to_cell(cell, lines)

def _fill_list_cell(cell, items: Sequence[str]) -> None:
    """
    Fill a cell that already uses list paragraphs (bullets) in the template.
    We DO NOT add bullet characters; the template provides list formatting.
    """
    items = [s for s in (items or []) if (s or "").strip()]
    paras = list(cell.paragraphs)

    if not paras:
        _fill_cell_single_paragraph(cell, "\n".join(items))
        return

    proto = paras[0]
    paras = _ensure_paragraphs_count(paras, max(1, len(items)), prototype=proto)

    if not items:
        _set_paragraph_text_keep_style(paras[0], "")
        for p in paras[1:]:
            _set_paragraph_text_keep_style(p, "")
        return

    for i, p in enumerate(paras):
        _set_paragraph_text_keep_style(p, items[i] if i < len(items) else "")

    # remove extras beyond item count
    if len(items) >= 1 and len(paras) > len(items):
        for p in paras[len(items):]:
            _remove_paragraph(p)


def _fill_proponents_row(proposal, left_cell, right_cell) -> None:
    """
    Fill proponents using the template's proponent "blocks".

    IMPORTANT FIX:
    - When cloning extra blocks, LEFT must keep label text (so the "Name"
      marker still exists), otherwise the detection loop never increments and
      generation becomes extremely slow / infinite.
    """
    props = _get_proponents(proposal)

    left_paras = list(left_cell.paragraphs)
    right_paras = list(right_cell.paragraphs)
    if not left_paras or not right_paras:
        return

    starts = [i for i, p in enumerate(left_paras) if (p.text or "").strip().lower() == "name"]
    if not starts:
        # safe fallback: plain block in right cell
        lines: List[str] = []
        for pr in props:
            lines.extend([
                (getattr(pr, "full_name", "") or "").strip(),
                (getattr(pr, "designation", "") or "").strip(),
                (getattr(pr, "specialization", "") or "").strip(),
                _compose_role(proposal, pr),
                (getattr(pr, "cp_number", "") or "").strip(),
                (getattr(pr, "email", "") or "").strip(),
                "",
            ])
        _fill_cell_single_paragraph(right_cell, "\n".join([x for x in lines if x is not None]).strip())
        return

    block_size = (starts[1] - starts[0]) if len(starts) >= 2 else 7
    needed_blocks = max(1, len(props))
    existing_blocks = len(starts)

    def clone_block(
        cell_paras: List[Paragraph],
        start_idx: int,
        end_idx: int,
        *,
        clear_text: bool
    ) -> None:
        block = cell_paras[start_idx:end_idx]
        cursor = cell_paras[-1]
        for src in block:
            # clone SOURCE paragraph to preserve its text+formatting (labels)
            newp = _insert_paragraph_clone_after(cursor, src)
            if clear_text:
                _set_paragraph_text_keep_style(newp, "")
            cursor = newp

    # Grow blocks (safe; this will now terminate because LEFT keeps "Name")
    while existing_blocks < needed_blocks:
        last_start = starts[-1]
        last_end = last_start + block_size

        clone_block(left_paras, last_start, last_end, clear_text=False)   # keep labels
        clone_block(right_paras, last_start, last_end, clear_text=True)   # blank values

        left_paras = list(left_cell.paragraphs)
        right_paras = list(right_cell.paragraphs)
        starts = [i for i, p in enumerate(left_paras) if (p.text or "").strip().lower() == "name"]
        existing_blocks = len(starts)

        # hard safety guard (prevents runaway if a template is malformed)
        if existing_blocks > 50:
            break

    # Trim extra blocks
    if existing_blocks > needed_blocks and len(starts) > needed_blocks:
        trim_from = starts[needed_blocks]
        for p in list(left_cell.paragraphs)[trim_from:]:
            _remove_paragraph(p)
        for p in list(right_cell.paragraphs)[trim_from:]:
            _remove_paragraph(p)

        left_paras = list(left_cell.paragraphs)
        right_paras = list(right_cell.paragraphs)
        starts = [i for i, p in enumerate(left_paras) if (p.text or "").strip().lower() == "name"]

    # Fill each block (assumes template aligns left/right block structure)
    right_paras = list(right_cell.paragraphs)
    for bi, start in enumerate(starts[:needed_blocks]):
        pr = props[bi] if bi < len(props) else None
        if pr is None:
            for off in range(0, min(block_size, len(right_paras) - start)):
                _set_paragraph_text_keep_style(right_paras[start + off], "")
            continue

        lines = [
            (getattr(pr, "full_name", "") or "").strip(),
            (getattr(pr, "designation", "") or "").strip(),
            (getattr(pr, "specialization", "") or "").strip(),
            _compose_role(proposal, pr),
            (getattr(pr, "cp_number", "") or "").strip(),
            (getattr(pr, "email", "") or "").strip(),
        ]

        for off, txt in enumerate(lines):
            if start + off < len(right_paras):
                _set_paragraph_text_keep_style(right_paras[start + off], txt)


def _fill_participants_and_gender(proposal, parent_cell) -> None:
    """
    Fill the nested VI (participants) table and VII (gender issues) table.

    FIX:
    - Uses LABEL matching instead of hard-coded row indices to avoid writing into
      wrong merged cells.
    """
    if not getattr(parent_cell, "tables", None):
        return
    if len(parent_cell.tables) < 2:
        return

    sex_male = _safe_int(getattr(proposal, "sex_male", 0), 0)
    sex_female = _safe_int(getattr(proposal, "sex_female", 0), 0)
    sex_total = sex_male + sex_female

    g_lesbian = _safe_int(getattr(proposal, "g_lesbian", 0), 0)
    g_gay = _safe_int(getattr(proposal, "g_gay", 0), 0)
    g_bisexual = _safe_int(getattr(proposal, "g_bisexual", 0), 0)
    g_transgender = _safe_int(getattr(proposal, "g_transgender", 0), 0)
    g_straight = _safe_int(getattr(proposal, "g_straight", 0), 0)
    g_others = _safe_int(getattr(proposal, "g_others", 0), 0)
    g_total = g_lesbian + g_gay + g_bisexual + g_transgender + g_straight + g_others

    # --- VI participants table ---
    t = parent_cell.tables[0]

    def set_cell_text(cell, text: str) -> None:
        _fill_cell_single_paragraph(cell, _normalize_text_preserve_newlines(text))

    # Sex section: labels usually in col 0, values in col 1
    sex_map = {"male": sex_male, "female": sex_female, "total": sex_total}
    for r in t.rows:
        if len(r.cells) < 2:
            continue
        label = (r.cells[0].text or "").strip().lower()
        if label in sex_map:
            set_cell_text(r.cells[1], str(sex_map[label]))

    # Gender section: labels usually in col 2, values in col 3
    gender_map = {
        "lesbian": g_lesbian,
        "gay": g_gay,
        "bisexual": g_bisexual,
        "transgender": g_transgender,
        "straight": g_straight,
        "others": g_others,
        "total": g_total,
    }

    def gender_key_from_label(text: str) -> str:
        label = (text or "").strip().lower()
        if "straight" in label:
            return "straight"
        if "others" in label:
            return "others"
        if "transgender" in label:
            return "transgender"
        if "bisexual" in label:
            return "bisexual"
        if label == "gay":
            return "gay"
        if "lesbian" in label:
            return "lesbian"
        if label == "total":
            return "total"
        return label

    for r in t.rows:
        if len(r.cells) < 4:
            continue
        label = gender_key_from_label(r.cells[2].text)
        if label in gender_map:
            set_cell_text(r.cells[3], str(gender_map[label]))

    # --- VII gender issues table ---
    issues, others_text = _get_gender_issues(proposal)
    gi = parent_cell.tables[1]

    def issue_key_from_text(text: str) -> str:
        ttxt = (text or "").lower()
        if "significant role of women" in ttxt:
            return "women_role_development"
        if "laws affecting family welfare" in ttxt or "family welfare" in ttxt:
            return "family_welfare_laws"
        if "lgbtq" in ttxt:
            return "lgbtq_acceptance"
        if "safe space" in ttxt or "gad" in ttxt:
            return "gad_awareness_safe_spaces"
        if "others" in ttxt:
            return "others"
        return ""

    for r in list(gi.rows):
        if len(r.cells) < 2:
            continue
        key = issue_key_from_text(r.cells[1].text)
        mark = "✓" if key and (key in issues) else ""
        mark = "/" if key and (key in issues) else ""
        set_cell_text(r.cells[0], mark)
        if key == "others":
            if "others" in issues and others_text:
                set_cell_text(r.cells[1], f"Others:\n{_clean_others_text(others_text)}".strip())
            else:
                row_el = r._tr
                parent = row_el.getparent()
                if parent is not None:
                    parent.remove(row_el)


def _fill_objectives_program(proposal, cell) -> None:
    """
    Rebuild the Objectives cell for PROGRAM scope using template clones.

    FIX:
    - Removed duplicated logic.
    - Uses _insert_paragraph_clone_after(proto) to preserve numbering cleanly.
    """
    paras = list(cell.paragraphs)
    if len(paras) < 4:
        gen = (getattr(proposal, "general_objective", "") or "").strip()
        lines = ["General Objective:", gen, "", "Specific Objectives:"]
        for prj in _get_program_projects(proposal):
            title = (getattr(prj, "title", "") or "").strip()
            if title:
                lines.append(title)
            for obj in _get_specific_objectives_for_project(proposal, prj):
                lines.append(obj)
            lines.append("")
        _fill_cell_single_paragraph(cell, "\n".join(lines).strip())
        return

    gen_text = (getattr(proposal, "general_objective", "") or "").strip()
    _set_paragraph_text_keep_style(paras[1], gen_text)

    proto_blank = paras[2]
    proto_phase = None
    proto_item = None

    for p in paras[4:]:
        if proto_phase is None and (not _has_num_pr(p)) and (p.text or "").strip():
            proto_phase = p
            continue
        if proto_item is None and _has_num_pr(p):
            proto_item = p
            break

    proto_phase = proto_phase or paras[4]
    proto_item = proto_item or paras[4]

    # Remove everything after "Specific Objectives:" header (index 3)
    for p in paras[4:]:
        _remove_paragraph(p)

    cursor = cell.paragraphs[-1]  # now ends at header

    prjs = _get_program_projects(proposal)
    for pi, prj in enumerate(prjs):
        title_raw = (getattr(prj, "title", "") or "").strip()
        title = title_raw

        # normalize "Phase I ..." -> "PHASE I. ..."
        if title.lower().startswith("phase "):
            parts = title.split(" ", 2)
            roman = parts[1].strip() if len(parts) >= 2 else ""
            rest = parts[2].strip() if len(parts) == 3 else ""
            dot = "." if roman and not roman.endswith(".") else ""
            title = f"PHASE {roman}{dot} {rest}".strip()

        phase_p = _insert_paragraph_clone_after(cursor, proto_phase)
        _set_paragraph_text_keep_style(phase_p, title)
        cursor = phase_p

        objectives = _get_specific_objectives_for_project(proposal, prj)
        if not objectives:
            item_p = _insert_paragraph_clone_after(cursor, proto_item)
            _clear_num_pr(item_p)
            _set_paragraph_text_keep_style(item_p, "")
            cursor = item_p
        else:
            for obj_idx, obj in enumerate(objectives, start=1):
                item_p = _insert_paragraph_clone_after(cursor, proto_item)
                _clear_num_pr(item_p)
                _set_paragraph_text_keep_style(item_p, f"{obj_idx}. {obj}")
                cursor = item_p

        if pi != len(prjs) - 1:
            blank_p = _insert_paragraph_clone_after(cursor, proto_blank)
            _set_paragraph_text_keep_style(blank_p, "")
            cursor = blank_p


def _fill_details_files_table(doc: Document, proposal) -> None:
    """
    Fill the Details/Funding attachment table with extracted uploaded content.
    """
    if len(doc.tables) < 2:
        return

    t = doc.tables[1]
    if not t.rows:
        return

    if len(t.rows) >= 1:
        details_cell = t.rows[0].cells[0]
        _clear_cell_content(details_cell)
        _add_xlsx_content_to_cell(details_cell, "Work Plan", getattr(proposal, "work_plan_file", None), table_kind="work_plan")
        _add_xlsx_content_to_cell(details_cell, "Gantt Chart", getattr(proposal, "gantt_chart_file", None), table_kind="gantt_chart")

    if len(t.rows) >= 2:
        funding_cell = t.rows[1].cells[0]
        _clear_cell_content(funding_cell)

        # Match the requested header format:
        # XIII. Funding Strategy
        # Line-item Budget: <fund source>
        p0 = funding_cell.paragraphs[0] if funding_cell.paragraphs else funding_cell.add_paragraph()
        _set_paragraph_text_keep_style(p0, "XIII. Funding Strategy")
        _apply_proposal_font_to_paragraph(p0)
        for r in p0.runs:
            r.bold = True

        p1 = funding_cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fund_title = "CBME EXTENSION FUND"
        br = (getattr(proposal, "budgetary_requirement", "") or "").strip()
        # Prefer text-like budgetary requirement as fund label so both sections match.
        if br and any(ch.isalpha() for ch in br):
            fund_title = br

        run = p1.add_run(f"Line-item Budget: {fund_title}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)

        _add_xlsx_content_to_cell(funding_cell, "Line-Item Budget", getattr(proposal, "funding_file", None), table_kind="line_item_budget")

def _signatory_lookup(position_title: str, *, campus: str = "", college: str = "", department: str = ""):
    """
    Find a Signatory by position_title with best-fit scope.

    IMPORTANT:
    - Signatory.position_title is a TextChoices VALUE in DB (e.g., 'CAMPUS_DIRECTOR', 'VPRDE').
    - Do NOT force campus/college/department filters when they are blank; that would miss scoped signatories.
    - Prefer the most specific scope first, then fall back.
    """
    if Signatory is None:
        return None

    pt = (position_title or "").strip()
    campus = (campus or "").strip()
    college = (college or "").strip()
    department = (department or "").strip()

    # Scope ladder: most specific -> least specific
    scopes = []
    if campus and college and department:
        scopes.append((campus, college, department))
    if campus and college:
        scopes.append((campus, college, ""))
    if campus:
        scopes.append((campus, "", ""))
    scopes.append(("", "", ""))

    seen = set()
    for campus_v, college_v, dept_v in scopes:
        key = (campus_v, college_v, dept_v)
        if key in seen:
            continue
        seen.add(key)

        try:
            qs = Signatory.objects.filter(position_title=pt)
            if campus_v:
                qs = qs.filter(campus__iexact=campus_v)
            if college_v:
                qs = qs.filter(college__iexact=college_v)
            if dept_v:
                qs = qs.filter(department__iexact=dept_v)
            sig = qs.first()
        except Exception:
            sig = None

        if sig:
            return sig

    # Final fallback: position only
    try:
        return Signatory.objects.filter(position_title=pt).first()
    except Exception:
        return None


def _display_signatory(sig) -> str:
    if not sig:
        return ""
    full = (getattr(sig, "full_name", "") or "").strip().upper()
    cred = (getattr(sig, "credentials", "") or "").strip()
    return (full + (f", {cred}" if cred else "")).strip()


def _fill_signatories(doc: Document, proposal) -> None:
    if Signatory is None:
        return

    # If the template doesn't have the expected signatory table, still fill bracket placeholders.
    if len(doc.tables) < 3:
        _fill_signatory_placeholders(doc, proposal)
        return

    campus = (getattr(proposal, "campus", "") or "").strip()
    college = (getattr(proposal, "college", "") or "").strip()
    department = (getattr(proposal, "department", "") or "").strip()

    try:
        positions = (
            Signatory.Position.DEPARTMENT_EXTENSION_COORDINATOR,
            Signatory.Position.DEAN,
            Signatory.Position.CAMPUS_EXTENSION_COORDINATOR,
        )
    except Exception:
        positions = (
            "DEPARTMENT_EXTENSION_COORDINATOR",
            "DEAN",
            "CAMPUS_EXTENSION_COORDINATOR",
        )

    names = [
        _display_signatory(_signatory_lookup(positions[0], campus=campus, college=college, department=department)),
        _display_signatory(_signatory_lookup(positions[1], campus=campus, college=college)),
        _display_signatory(_signatory_lookup(positions[2], campus=campus)),
    ]

    t_sign = doc.tables[2]
    for ci in range(min(3, len(t_sign.rows[0].cells))):
        if names[ci]:
            _fill_cell_single_paragraph(t_sign.rows[0].cells[ci], names[ci])
    # Also replace bracket placeholders anywhere in the document.
    _fill_signatory_placeholders(doc, proposal)



def _fill_form1_program(doc: Document, proposal) -> None:
    """Fill Form 1 Program template (form1_program_template.docx)."""
    if not doc.tables:
        return

    t0 = doc.tables[0]

    # Title
    title = (getattr(proposal, "title", "") or getattr(proposal, "research_title", "") or "").strip()
    _fill_cell_single_paragraph(t0.rows[0].cells[1], title)

    # Proponents (Row containing "Proponent")
    proponent_row = _find_row_index_by_keyword(t0, "proponent")
    if proponent_row is not None:
        _fill_proponents_row(proposal, t0.rows[proponent_row].cells[0], t0.rows[proponent_row].cells[1])

    # Implementing agency/unit
    idx = _find_row_index_by_keyword(t0, "implementing")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "implementing_agency", "") or "").strip())

    # Collaborators/Beneficiaries
    idx = _find_row_index_by_keyword(t0, "collaborators")
    if idx is not None:
        who = (getattr(proposal, "beneficiaries_who", "") or "").strip()
        count = getattr(proposal, "beneficiaries_count", None)
        prefix = f"{count} " if count is not None else ""
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (prefix + who).strip())

    # SDGs
    idx = _find_row_index_by_keyword(t0, "sdg")
    if idx is not None:
        _fill_list_cell(t0.rows[idx].cells[1], _get_sdg_lines(proposal))

    # Extension thrust
    idx = _find_row_index_by_keyword(t0, "thrust")
    if idx is not None:
        _fill_list_cell(t0.rows[idx].cells[1], _get_thrust_lines(proposal))

    # Budgetary Requirement
    idx = _find_row_index_by_keyword(t0, "budgetary")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "budgetary_requirement", "") or "").strip())

    # Participants + Gender Issues row (contains nested tables)
    idx = _find_row_index_by_keyword(t0, "participants/proposed")
    if idx is not None:
        _fill_participants_and_gender(proposal, t0.rows[idx].cells[1])

    # Date and Venue / Extension Site
    idx = _find_row_index_by_keyword(t0, "date and venue")
    if idx is None:
        idx = _find_row_index_by_keyword(t0, "date & venue")
    if idx is not None:
        month = (getattr(proposal, "estimated_month", "") or "").strip()
        year = getattr(proposal, "estimated_year", None)
        venue = (getattr(proposal, "extension_venue", "") or "").strip()
        lines = []
        if month or year:
            lines.append(f"{month} {year}".strip())
        if venue:
            lines.append(venue)
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], "\n".join(lines).strip())

    # Rationale / Background
    idx = _find_row_index_by_keyword(t0, "rationale")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "rationale_background", "") or "").strip())

    # Significance
    idx = _find_row_index_by_keyword(t0, "significance")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "significance", "") or "").strip())

    # Objectives (General to Specific)
    idx = _find_row_index_by_keyword(t0, "objectives")
    if idx is not None:
        _fill_objectives_program(proposal, t0.rows[idx].cells[1])

    # Methodology/Mechanics
    idx = _find_row_index_by_keyword(t0, "methodology")
    if idx is not None:
        _fill_list_cell(t0.rows[idx].cells[1], _get_methodology_items(proposal))

    # Output/Outcome
    idx = _find_row_index_by_keyword(t0, "output/outcome")
    if idx is not None:
        outs = _get_output_items(proposal)
        cell = t0.rows[idx].cells[1]
        paras = list(cell.paragraphs)
        if not paras:
            _fill_cell_single_paragraph(cell, "\n".join(outs))
        else:
            paras = _ensure_paragraphs_count(paras, max(1, len(outs)), prototype=paras[0])
            if not outs:
                _set_paragraph_text_keep_style(paras[0], "")
            else:
                for i, p in enumerate(paras):
                    _set_paragraph_text_keep_style(p, outs[i] if i < len(outs) else "")
                if len(paras) > len(outs):
                    for p in paras[len(outs):]:
                        _remove_paragraph(p)

    # --- Signatories tables (bottom) ---
    if Signatory is not None and len(doc.tables) >= 3:
        campus = (getattr(proposal, "campus", "") or "").strip()
        college = (getattr(proposal, "college", "") or "").strip()
        department = (getattr(proposal, "department", "") or "").strip()

        def get_sig(position_title: str, *, campus_v: str = "", college_v: str = "", dept_v: str = ""):
            try:
                return Signatory.objects.filter(
                    position_title__iexact=position_title,
                    campus=campus_v,
                    college=college_v,
                    department=dept_v,
                ).first()
            except Exception:
                return None

        def display(sig) -> str:
            if not sig:
                return ""
            full = (getattr(sig, "full_name", "") or "").strip()
            cred = (getattr(sig, "credentials", "") or "").strip()
            return (full + (f", {cred}" if cred else "")).strip()

        t_sign = doc.tables[2]  # based on template structure
        try:
            dept_pos = Signatory.Position.DEPARTMENT_EXTENSION_COORDINATOR
            dean_pos = Signatory.Position.DEAN
            campus_coord_pos = Signatory.Position.CAMPUS_EXTENSION_COORDINATOR
        except Exception:
            dept_pos = "DEPARTMENT_EXTENSION_COORDINATOR"
            dean_pos = "DEAN"
            campus_coord_pos = "CAMPUS_EXTENSION_COORDINATOR"

        dept_sig = get_sig(dept_pos, campus_v=campus, college_v=college, dept_v=department)
        dean_sig = get_sig(dean_pos, campus_v=campus, college_v=college)
        campus_coord_sig = get_sig(campus_coord_pos, campus_v=campus)

        names = [display(dept_sig), display(dean_sig), display(campus_coord_sig)]
        for ci in range(min(3, len(t_sign.rows[0].cells))):
            if names[ci]:
                _fill_cell_single_paragraph(t_sign.rows[0].cells[ci], names[ci])

    _fill_details_files_table(doc, proposal)
    _fill_signatories(doc, proposal)


def _fill_form1_project(doc: Document, proposal) -> None:
    """Project template fill (non-destructive)."""
    if not doc.tables:
        return
    t0 = doc.tables[0]

    title = (getattr(proposal, "title", "") or getattr(proposal, "research_title", "") or "").strip()
    _fill_cell_single_paragraph(t0.rows[0].cells[1], title)

    proponent_row = _find_row_index_by_keyword(t0, "proponent")
    if proponent_row is not None:
        _fill_proponents_row(proposal, t0.rows[proponent_row].cells[0], t0.rows[proponent_row].cells[1])

    idx = _find_row_index_by_keyword(t0, "implementing")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "implementing_agency", "") or "").strip())

    idx = _find_row_index_by_keyword(t0, "collaborators")
    if idx is not None:
        who = (getattr(proposal, "beneficiaries_who", "") or "").strip()
        count = getattr(proposal, "beneficiaries_count", None)
        prefix = f"{count} " if count is not None else ""
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (prefix + who).strip())

    idx = _find_row_index_by_keyword(t0, "sdg")
    if idx is not None:
        _fill_list_cell(t0.rows[idx].cells[1], _get_sdg_lines(proposal))

    idx = _find_row_index_by_keyword(t0, "thrust")
    if idx is not None:
        _fill_list_cell(t0.rows[idx].cells[1], _get_thrust_lines(proposal))

    idx = _find_row_index_by_keyword(t0, "budgetary")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "budgetary_requirement", "") or "").strip())

    idx = _find_row_index_by_keyword(t0, "participants/proposed")
    if idx is not None:
        _fill_participants_and_gender(proposal, t0.rows[idx].cells[1])

    idx = _find_row_index_by_keyword(t0, "date and venue")
    if idx is None:
        idx = _find_row_index_by_keyword(t0, "date & venue")
    if idx is not None:
        month = (getattr(proposal, "estimated_month", "") or "").strip()
        year = getattr(proposal, "estimated_year", None)
        venue = (getattr(proposal, "extension_venue", "") or "").strip()
        lines = []
        if month or year:
            lines.append(f"{month} {year}".strip())
        if venue:
            lines.append(venue)
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], "\n".join(lines).strip())

    idx = _find_row_index_by_keyword(t0, "rationale")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "rationale_background", "") or "").strip())

    idx = _find_row_index_by_keyword(t0, "significance")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], (getattr(proposal, "significance", "") or "").strip())

    idx = _find_row_index_by_keyword(t0, "objectives")
    if idx is not None:
        gen = (getattr(proposal, "general_objective", "") or "").strip()
        objs = []
        try:
            qs = proposal.specific_objectives.filter(program_project__isnull=True).values_list("objective", flat=True)
            objs = [str(x).strip() for x in qs if str(x).strip()]
        except Exception:
            pass
        block = ["General Objective:", gen, "", "Specific Objectives:"]
        block.extend(objs)
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], "\n".join(block).strip())

    idx = _find_row_index_by_keyword(t0, "methodology")
    if idx is not None:
        _fill_list_cell(t0.rows[idx].cells[1], _get_methodology_items(proposal))

    idx = _find_row_index_by_keyword(t0, "output/outcome")
    if idx is not None:
        _fill_cell_single_paragraph(t0.rows[idx].cells[1], "\n".join(_get_output_items(proposal)))

    _fill_details_files_table(doc, proposal)
    _fill_signatories(doc, proposal)


def _fill_form2_training_design(doc: Document, proposal) -> None:
    """Fill Form 2 Training Design template."""
    if not doc.tables:
        return
    t0 = doc.tables[0]
    title = (getattr(proposal, "title", "") or getattr(proposal, "research_title", "") or "").strip()

    if len(t0.rows) > 0:
        _fill_cell_single_paragraph(t0.rows[0].cells[1], title)
    if len(t0.rows) > 1:
        _fill_proponents_row(proposal, t0.rows[1].cells[0], t0.rows[1].cells[1])
    if len(t0.rows) > 2:
        _fill_cell_single_paragraph(t0.rows[2].cells[1], (getattr(proposal, "implementing_agency", "") or "").strip())
    if len(t0.rows) > 3:
        who = (getattr(proposal, "beneficiaries_who", "") or "").strip()
        count = getattr(proposal, "beneficiaries_count", None)
        prefix = f"{count} " if count is not None else ""
        _fill_cell_single_paragraph(t0.rows[3].cells[1], (prefix + who).strip())
    if len(t0.rows) > 4:
        _fill_list_cell(t0.rows[4].cells[1], _get_sdg_lines(proposal))
    if len(t0.rows) > 5:
        _fill_list_cell(t0.rows[5].cells[1], _get_thrust_lines(proposal))
    if len(t0.rows) > 6:
        month = (getattr(proposal, "estimated_month", "") or "").strip()
        year = getattr(proposal, "estimated_year", None)
        _fill_cell_single_paragraph(t0.rows[6].cells[1], f"{month} {year}".strip())
    if len(t0.rows) > 7:
        venue = (getattr(proposal, "extension_venue", "") or "").strip()
        _fill_cell_single_paragraph(t0.rows[7].cells[1], venue)
    if len(t0.rows) > 8:
        _fill_cell_single_paragraph(t0.rows[8].cells[1], (getattr(proposal, "budgetary_requirement", "") or "").strip())
    if len(t0.rows) > 9:
        _fill_cell_single_paragraph(t0.rows[9].cells[1], "")

    if len(doc.tables) > 1 and doc.tables[1].rows:
        _fill_participants_and_gender(proposal, doc.tables[1].rows[0].cells[0])
    _fill_signatories(doc, proposal)


# =============================
# Public API
# =============================


def _cleanup_document_end(doc: Document) -> None:
    """
    Remove trailing empty/page-break-only paragraphs that commonly cause a blank LAST page.

    IMPORTANT:
    - We avoid removing paragraphs that carry section properties (w:sectPr),
      because those define page layout for the preceding content.
    """
    # Remove "page break before" from body paragraphs
    for p in list(getattr(doc, "paragraphs", [])):
        try:
            p_pr = p._p.get_or_add_pPr()
            pb = p_pr.find(qn("w:pageBreakBefore"))
            if pb is not None:
                p_pr.remove(pb)
        except Exception:
            pass

    # Drop trailing empty paragraphs (but keep the last sectPr paragraph intact)
    removed = 0
    for p in reversed(list(getattr(doc, "paragraphs", []))):
        if removed >= 25:
            break
        if (p.text or "").strip():
            break
        try:
            p_pr = p._p.get_or_add_pPr()
            if p_pr.find(qn("w:sectPr")) is not None:
                break
        except Exception:
            pass
        try:
            _remove_paragraph(p)
            removed += 1
        except Exception:
            break

def build_extension_form_docx(proposal, *args, **kwargs) -> bytes:
    """
    Generate the downloadable Proposal DOCX.

    NOTE:
    - Signature-flexible (accepts *args/**kwargs) because views changed multiple times.
      Extra args/kwargs are ignored.

    Returns:
        bytes of the generated DOCX.
    """
    base_dir = Path(__file__).resolve().parent / "template_files"

    extension_type = (getattr(proposal, "extension_type", "") or "").strip().upper()
    scope_type = (getattr(proposal, "scope_type", "") or "").strip().upper()

    is_research = extension_type in {"RESEARCH_FACULTY", "RESEARCH_STUDENT"}

    if is_research:
        if scope_type == "PROGRAM":
            template_path = base_dir / "form1_program_template.docx"
        else:
            template_path = base_dir / "form1_project_template.docx"

        doc = Document(str(template_path))
        if scope_type == "PROGRAM":
            _fill_form1_program(doc, proposal)
        else:
            _fill_form1_project(doc, proposal)
    else:
        template_path = base_dir / "form2_training_design_template.docx"
        doc = Document(str(template_path))
        _fill_form2_training_design(doc, proposal)
    _cleanup_document_end(doc)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
def _iter_all_paragraphs(doc: Document):
    """Yield all paragraphs in the document, including tables and section headers/footers."""

    def walk_table(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    yield from walk_table(nested)

    # Body paragraphs
    for p in getattr(doc, "paragraphs", []):
        yield p
    for t in getattr(doc, "tables", []):
        yield from walk_table(t)

    # Headers/footers (signatories are often placed here in templates)
    for sec in getattr(doc, "sections", []):
        for hf in (
            getattr(sec, "header", None),
            getattr(sec, "footer", None),
            getattr(sec, "first_page_header", None),
            getattr(sec, "first_page_footer", None),
            getattr(sec, "even_page_header", None),
            getattr(sec, "even_page_footer", None),
        ):
            if hf is None:
                continue
            for p in getattr(hf, "paragraphs", []):
                yield p
            for t in getattr(hf, "tables", []):
                yield from walk_table(t)



def _replace_placeholders_in_doc_xml(doc: Document, compiled: Sequence[Tuple[str, re.Pattern, str]]) -> None:
    """
    Replace placeholders in OOXML, including content inside shapes/textboxes.

    IMPORTANT:
    - Word "text boxes" may store text in either WordprocessingML (w:t) OR DrawingML (a:t).
      The previous implementation only replaced w:t, which misses many modern templates.
    """
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }

    def iter_parts():
        parts = [getattr(doc, "part", None)]
        for sec in getattr(doc, "sections", []):
            for hf in (
                getattr(sec, "header", None),
                getattr(sec, "footer", None),
                getattr(sec, "first_page_header", None),
                getattr(sec, "first_page_footer", None),
                getattr(sec, "even_page_header", None),
                getattr(sec, "even_page_footer", None),
            ):
                if hf is None:
                    continue
                try:
                    parts.append(hf.part)
                except Exception:
                    pass

        out = []
        seen = set()
        for p in parts:
            if p is None:
                continue
            key = id(p)
            if key in seen:
                continue
            seen.add(key)
            out.append(p)
        return out

    def apply_replacements(text: str) -> str:
        new = text
        for _, rx, name in compiled:
            new = rx.sub(name, new)
        # Cleanup: some templates leave a trailing '|' after placeholder lines
        new = re.sub(r"\s*\|\s*$", "", new)
        return new

    for part in iter_parts():
        root = getattr(part, "element", None) or getattr(part, "_element", None)
        if root is None:
            continue

        # --- WordprocessingML paragraphs (w:p / w:t) ---
        try:
            ps = root.xpath(".//w:p", namespaces=ns)
        except Exception:
            ps = []

        for wp in ps:
            try:
                ts = wp.xpath(".//w:t", namespaces=ns)
            except Exception:
                ts = []
            if not ts:
                continue

            full = "".join((t.text or "") for t in ts)
            if not full:
                continue

            new = apply_replacements(full)
            if new != full:
                ts[0].text = new
                for t in ts[1:]:
                    t.text = ""

        # --- DrawingML paragraphs inside shapes/textboxes (a:p / a:t) ---
        try:
            aps = root.xpath(".//a:p", namespaces=ns)
        except Exception:
            aps = []

        for ap in aps:
            try:
                ats = ap.xpath(".//a:t", namespaces=ns)
            except Exception:
                ats = []
            if not ats:
                continue
            full = "".join((t.text or "") for t in ats)
            if not full:
                continue

            new = apply_replacements(full)
            if new != full:
                ats[0].text = new
                for t in ats[1:]:
                    t.text = ""


def _replace_in_paragraph_runs(paragraph: Paragraph, needle: str, repl: str) -> bool:
    """Replace text in runs if possible, preserving formatting. Returns True if replaced."""
    if not needle or repl is None:
        return False
    changed = False
    for run in paragraph.runs:
        if needle in (run.text or ""):
            run.text = (run.text or "").replace(needle, repl)
            changed = True
    if changed:
        return True
    # Fallback: placeholder might be split across runs; if paragraph contains it, rewrite whole paragraph
    if needle in (paragraph.text or ""):
        _set_paragraph_text_keep_style(paragraph, (paragraph.text or "").replace(needle, repl))
        return True
    return False


def _fill_signatory_placeholders(doc: Document, proposal) -> None:
    """
    Replace bracket placeholders like [CAMPUS_DIRECTOR], [VPRDE], etc. using Signatory records.

    Your Signatory model stores TextChoices VALUES in DB, e.g.:
        CAMPUS_DIRECTOR, CAMPUS_EXTENSION_COORDINATOR, DEAN,
        DEPARTMENT_EXTENSION_COORDINATOR, DIRECTOR_EXTENSION, VPRDE, SUC_PRESIDENT_III

    This function replaces placeholders in:
    - normal document paragraphs/tables
    - headers/footers
    - shapes/textboxes (DrawingML a:t) via `_replace_placeholders_in_doc_xml`
    """
    if Signatory is None:
        return

    campus = (getattr(proposal, "campus", "") or "").strip()
    college = (getattr(proposal, "college", "") or "").strip()
    department = (getattr(proposal, "department", "") or "").strip()

    # Prefer the explicit TextChoices values (stable).
    try:
        codes = {c[0] for c in Signatory.Position.choices}
    except Exception:
        try:
            codes = set(Signatory.objects.values_list("position_title", flat=True).distinct())
        except Exception:
            codes = set()

    if not codes:
        return

    # Backward-compatible aliases (in case some templates still use the readable placeholders).
    alias_to_code = {
        "CAMPUS DIRECTOR": "CAMPUS_DIRECTOR",
        "CAMPUS_EXTENSION_COORDINATOR": "CAMPUS_EXTENSION_COORDINATOR",
        "CAMPUS EXTENSION COORDINATOR": "CAMPUS_EXTENSION_COORDINATOR",
        "DEAN": "DEAN",
        "DEPARTMENT EXTENSION COORDINATOR": "DEPARTMENT_EXTENSION_COORDINATOR",
        "DEPARTMENT_EXTENSION_COORDINATOR": "DEPARTMENT_EXTENSION_COORDINATOR",
        "DIRECTOR FOR EXTENSION": "DIRECTOR_EXTENSION",
        "DIRECTOR_EXTENSION": "DIRECTOR_EXTENSION",
        "VICE PRESIDENT FOR RESEARCH DEVELOPMENT AND EXTENSION": "VPRDE",
        "VPRDE": "VPRDE",
        "SUC PRESIDENT": "SUC_PRESIDENT_III",
        "SUC_PRESIDENT_III": "SUC_PRESIDENT_III",
        "SUC PRESIDENT III": "SUC_PRESIDENT_III",
    }

    # Build placeholder->resolved name
    resolved: dict[str, str] = {}

    def resolve_code(code: str) -> str:
        sig = _signatory_lookup(code, campus=campus, college=college, department=department)
        if not sig:
            sig = _signatory_lookup(code, campus="", college="", department="")
        return _display_signatory(sig) if sig else ""

    # Replace for every known code (Option A)
    for code in sorted(codes):
        name = resolve_code(code)
        if name:
            resolved[f"[{code}]"] = name

    # Replace for common readable aliases too
    for alias, code in alias_to_code.items():
        if code not in codes:
            continue
        name = resolve_code(code)
        if name:
            resolved[f"[{alias}]"] = name

    if not resolved:
        return

    def placeholder_regex(ph: str) -> re.Pattern:
        inner = ph.strip()[1:-1]  # remove []
        # allow any whitespace/newlines around the token
        inner_esc = re.escape(inner)
        inner_esc = inner_esc.replace(r"\ ", r"\s+")
        return re.compile(r"\[\s*" + inner_esc + r"\s*\]", flags=re.IGNORECASE)

    compiled: list[tuple[str, re.Pattern, str]] = [(ph, placeholder_regex(ph), name) for ph, name in resolved.items()]

    # Shallow replacement (keeps general formatting in body/table paragraphs)
    for p in _iter_all_paragraphs(doc):
        original = p.text or ""
        new_text = original
        for ph, rx, name in compiled:
            if ph in new_text:
                new_text = new_text.replace(ph, name)
            else:
                new_text = rx.sub(name, new_text)

        # Cleanup: some templates leave a trailing '|' after placeholder lines
        new_text = re.sub(r"\s*\|\s*$", "", new_text)

        if new_text != original:
            _set_paragraph_text_keep_style(p, new_text)

    # Deep replacement for textboxes/shapes (WordprocessingML + DrawingML)
    _replace_placeholders_in_doc_xml(doc, compiled)

