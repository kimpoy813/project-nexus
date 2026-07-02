from collections import Counter, OrderedDict
from copy import copy
from datetime import timedelta
from io import BytesIO
import math
import mimetypes
from pathlib import Path
import re
from urllib.parse import quote
from django.urls import reverse

from django.contrib import messages
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required
from django.core.exceptions import ValidationError
from django.db.models import Prefetch, Q
from django.db import transaction
from django.http import FileResponse, Http404, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from django.views.decorators.http import require_GET, require_POST, require_http_methods

from openpyxl import load_workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from urllib3 import request
from accounts.decorators import faculty_like_required, role_required
from details.models import ExtensionProcess, ProcessStep

from .models import (
    ExtensionThrust,
    ProgramProject,
    Proposal,
    ProposalAttachment,
    ProposalFinalDocument,
    ProposalCollaborator,
    ProposalCommentSummary,
    ProposalEditorPresence,
    ProposalEvaluatorAssignment,
    ProposalGenderIssue,
    ProposalMethodology,
    ProposalOutputOutcome,
    ProposalProponent,
    ProposalReviewRound,
    ProposalSDG,
    ProposalSectionComment,
    ProposalSpecificObjective,
    ProposalThrust,
    SDG,
)

from xhtml2pdf import pisa
from django.template.loader import render_to_string
from docx.shared import Mm, Pt

from .docx_forms import build_extension_form_docx

User = get_user_model()

TOTAL_STEPS = 19

STEP_LABELS = [
    {"no": 1, "title": "Extension Type and Scope", "desc": "Type of extension and proposal scope"},
    {"no": 2, "title": "Title", "desc": "Program, project, or activity title"},
    {"no": 3, "title": "Proponents", "desc": "Proponent details and assigned roles"},
    {"no": 4, "title": "Implementing Agency/Unit", "desc": "Office, agency, or unit responsible"},
    {"no": 5, "title": "Collaborators/Beneficiaries", "desc": "Beneficiary count and target group"},
    {"no": 6, "title": "SDGs / Extension Agenda", "desc": "SDGs covered and extension thrust"},
    {"no": 7, "title": "Budgetary Requirement", "desc": "Funding source and budget"},
    {"no": 8, "title": "Participants / Proposed Clients", "desc": "Participant profiling and counts"},
    {"no": 9, "title": "Gender Issues / Mandates Addressed", "desc": "Applicable GAD mandates"},
    {"no": 10, "title": "Date and Venue / Extension Site", "desc": "Schedule and implementation site"},
    {"no": 11, "title": "Rationale / Background", "desc": "Context and alignment with SDG / thrust / GAD"},
    {"no": 12, "title": "Significance", "desc": "Importance of the proposed extension"},
    {"no": 13, "title": "Objectives", "desc": "General and specific SMART objectives"},
    {"no": 14, "title": "Methodology / Mechanics", "desc": "Implementation approach"},
    {"no": 15, "title": "Output / Outcome", "desc": "Expected outputs and outcomes"},
    {"no": 16, "title": "Details of Activities", "desc": "Work plan, Gantt chart, and related files"},
    {"no": 17, "title": "Funding Strategy", "desc": "Funding strategy template and related supporting files"},
    {"no": 18, "title": "Research Abstract Upload", "desc": "Required for research-based proposals"},
    {"no": 19, "title": "Certificate of Completion Upload", "desc": "Required for research-based proposals"},
]

SDG_LIST = [
    ("01", "No Poverty"),
    ("02", "Zero Hunger"),
    ("03", "Good Health and Well-being"),
    ("04", "Quality Education"),
    ("05", "Gender Equality"),
    ("06", "Clean Water and Sanitation"),
    ("07", "Affordable and Clean Energy"),
    ("08", "Decent Work and Economic Growth"),
    ("09", "Industry, Innovation and Infrastructure"),
    ("10", "Reduced Inequalities"),
    ("11", "Sustainable Cities and Communities"),
    ("12", "Responsible Consumption and Production"),
    ("13", "Climate Action"),
    ("14", "Life Below Water"),
    ("15", "Life on Land"),
    ("16", "Peace, Justice and Strong Institutions"),
    ("17", "Partnerships for the Goals"),
]

THRUST_LIST = [
    "Indigenous Heritage Protection",
    "Environmental Protection",
    "Resource Sharing",
    "Numeracy and Literacy",
    "Governance and Administration",
    "IP-TBM Office Establishment",
    "Trade Fair and Exhibit",
    "Technology Transfer & RD Results Dissemination",
    "Network and Linkage",
    "Adult Education",
    "Calamity & Disaster Rehabilitation",
    "Entrepreneurship & Financial Literacy",
    "Health and Nutrition",
    "Advocacies & Social Justice",
]

GENDER_ISSUE_LIST = [
    (
        "women_role_development",
        "The activity strengthens the advocacy on the significant role of women in development.",
    ),
    (
        "family_welfare_laws",
        "The activity strengthens the understanding of the men and women in barangays on the laws affecting family welfare; decreased incidence of bullying and SH in the barangay.",
    ),
    (
        "lgbtq_acceptance",
        "The activity increases the level of acceptance of the LGBTQ in the society.",
    ),
    (
        "gad_awareness_safe_spaces",
        "The activity enhances the level of awareness on GAD issues and concepts including related laws specifically Safe Space Act.",
    ),
    ("others", "Others"),
]

def _format_numbered_list_paragraph(paragraph):
    """
    Force consistent numbering alignment (hanging indent) for all list items.
    """
    pf = paragraph.paragraph_format
    pf.left_indent = Pt(36)          # 0.5 inch
    pf.first_line_indent = Pt(-18)   # hanging by 0.25 inch
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

def _to_int(value, default=0):
    value = str(value or "").strip()
    return int(value) if value.isdigit() else default


def _to_roman(n: int) -> str:
    vals = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    out = []
    for v, s in vals:
        while n >= v:
            out.append(s)
            n -= v
    return "".join(out)


def _strip_phase_prefix(text: str) -> str:
    return re.sub(r"^Phase\s+[IVXLCDM]+\s+", "", (text or "").strip(), flags=re.I).strip()


def _is_director(user):
    profile = getattr(user, "profile", None)
    return getattr(profile, "role", "") == "DIRECTOR"


def _is_staff(user):
    profile = getattr(user, "profile", None)
    return getattr(profile, "role", "") == "STAFF"


def _is_department_coordinator(user):
    profile = getattr(user, "profile", None)
    return getattr(profile, "role", "") == "DEPARTMENT_COORDINATOR"


def _is_campus_coordinator(user):
    profile = getattr(user, "profile", None)
    return getattr(profile, "role", "") == "CAMPUS_COORDINATOR"


def _has_active_evaluator_assignment(user, proposal=None, review_round=None):
    qs = ProposalEvaluatorAssignment.objects.filter(
        evaluator=user,
        is_active=True,
    )
    if proposal is not None:
        qs = qs.filter(proposal=proposal)
    if review_round is not None:
        qs = qs.filter(review_round=review_round)
    return qs.exists()


def _can_edit(user, proposal):
    if not user.is_authenticated:
        return False

    if proposal.created_by_id == user.id:
        return True

    if ProposalCollaborator.objects.filter(
        proposal=proposal,
        user=user,
        can_edit=True,
    ).exists():
        return True

    if ProposalProponent.objects.filter(
        proposal=proposal,
        user=user,
    ).exists():
        return True

    return False


def _can_review(user, proposal):
    if not user.is_authenticated:
        return False

    profile = getattr(user, "profile", None)
    role = getattr(profile, "role", "")

    if role == "DIRECTOR":
        return True

    if role == "DEPARTMENT_COORDINATOR":
        return (proposal.department or "").strip() == (getattr(profile, "department", "") or "").strip()

    if role == "CAMPUS_COORDINATOR":
        return (proposal.campus or "").strip() == (getattr(profile, "campus", "") or "").strip()

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()

    # Step-level comment counts (for sidebar markers)
    step_comment_counts = {}
    if current_round:
        raw_steps = ProposalSectionComment.objects.filter(
            proposal=proposal,
            review_round=current_round,
        ).values_list("step_no", flat=True)
        temp = Counter()
        for st in raw_steps:
            try:
                st_no = int(st or 1)
            except (TypeError, ValueError):
                st_no = 1
            temp[st_no] += 1
        step_comment_counts = dict(temp)
    return _has_active_evaluator_assignment(user, proposal=proposal, review_round=current_round)


def _can_view_proposal(user, proposal):
    return _can_edit(user, proposal) or _can_review(user, proposal) or _is_staff(user)



def _ensure_open_review_round(proposal, user):
    """Ensure an OPEN review round exists (is_closed=False), creating a new one if needed."""
    rr = ProposalReviewRound.objects.filter(proposal=proposal, is_closed=False).order_by("-round_no", "-id").first()
    if rr:
        return rr

    last_rr = ProposalReviewRound.objects.filter(proposal=proposal).order_by("-round_no", "-id").first()
    next_no = (getattr(last_rr, "round_no", 0) or 0) + 1

    field_names = {f.name for f in ProposalReviewRound._meta.get_fields() if hasattr(f, "name")}
    create_kwargs = {"proposal": proposal}

    if "round_no" in field_names:
        create_kwargs["round_no"] = next_no
    if "is_closed" in field_names:
        create_kwargs["is_closed"] = False
    if "ready_for_staff_summary" in field_names:
        create_kwargs["ready_for_staff_summary"] = False

    for k in ("started_by", "created_by", "opened_by"):
        if k in field_names:
            create_kwargs[k] = user

    return ProposalReviewRound.objects.create(**create_kwargs)

def _get_reviewer_role(user, proposal, review_round):
    if not user.is_authenticated or not review_round:
        return ""

    if _is_director(user):
        return "DIRECTOR"

    if _is_department_coordinator(user):
        profile = getattr(user, "profile", None)
        user_department = (getattr(profile, "department", "") or "").strip()
        proposal_department = (proposal.department or "").strip()
        if user_department == proposal_department:
            return "DEPARTMENT_COORDINATOR"
        return ""

    if _is_campus_coordinator(user):
        profile = getattr(user, "profile", None)
        user_campus = (getattr(profile, "campus", "") or "").strip()
        proposal_campus = (proposal.campus or "").strip()
        if user_campus == proposal_campus:
            return "CAMPUS_COORDINATOR"
        return ""

    if _has_active_evaluator_assignment(user, proposal=proposal, review_round=review_round):
        return "EVALUATOR"

    return ""



def mark_step_completed(proposal, step_no):
    completed = set(proposal.completed_steps or [])
    skipped = set(proposal.skipped_steps or [])
    completed.add(step_no)
    skipped.discard(step_no)
    proposal.completed_steps = sorted(completed)
    proposal.skipped_steps = sorted(skipped)
    messages.success(request, "Step saved successfully.")


def mark_step_skipped(proposal, step_no):
    completed = set(proposal.completed_steps or [])
    skipped = set(proposal.skipped_steps or [])
    completed.discard(step_no)
    skipped.add(step_no)
    proposal.completed_steps = sorted(completed)
    proposal.skipped_steps = sorted(skipped)


def unmark_step_completed(proposal, step_no):
    completed = set(proposal.completed_steps or [])
    completed.discard(step_no)
    proposal.completed_steps = sorted(completed)


def is_step_complete(proposal, step):
    if step == 1:
        if not proposal.extension_type or not proposal.scope_type:
            return False
        if proposal.extension_type in ["RESEARCH_FACULTY", "RESEARCH_STUDENT"] and not (proposal.research_title or "").strip():
            return False
        return True

    if step == 2:
        if not (proposal.title or "").strip():
            return False
        if proposal.scope_type == "PROGRAM":
            return proposal.program_projects.exists()
        return True

    if step == 3:
        return proposal.proponents.exists()

    if step == 4:
        return bool((proposal.implementing_agency or "").strip())

    if step == 5:
        return proposal.beneficiaries_count is not None and bool((proposal.beneficiaries_who or "").strip())

    if step == 6:
        return proposal.sdg_links.exists() or proposal.thrust_links.exists()

    if step == 7:
        return bool((proposal.budgetary_requirement or "").strip())

    if step == 8:
        sex_total = (proposal.sex_male or 0) + (proposal.sex_female or 0)
        gender_total = (
            (proposal.g_lesbian or 0)
            + (proposal.g_gay or 0)
            + (proposal.g_bisexual or 0)
            + (proposal.g_transgender or 0)
            + (proposal.g_straight or 0)
            + (proposal.g_others or 0)
        )
        return sex_total > 0 and sex_total == gender_total

    if step == 9:
        issues = proposal.gender_issue_links.all()
        if not issues.exists():
            return False
        others = issues.filter(issue_key="others").first()
        if others and not (others.other_text or "").strip():
            return False
        return True

    if step == 10:
        return bool((proposal.extension_venue or "").strip())

    if step == 11:
        return bool((proposal.rationale_background or "").strip())

    if step == 12:
        return bool((proposal.significance or "").strip())

    if step == 13:
        if not (proposal.general_objective or "").strip():
            return False

        if proposal.scope_type == "PROGRAM":
            projects = proposal.program_projects.all()
            if not projects.exists():
                return False
            for prj in projects:
                if not proposal.specific_objectives.filter(program_project=prj).exists():
                    return False
            return True

        return proposal.specific_objectives.filter(program_project__isnull=True).exists()

    if step == 14:
        return proposal.methodologies.exists()

    if step == 15:
        return proposal.output_outcomes.exists()

    if step == 16:
        return bool(proposal.work_plan_file) and bool(proposal.gantt_chart_file)

    if step == 17:
        return bool(proposal.funding_file)

    if step == 18:
        if proposal.extension_type in ["RESEARCH_FACULTY", "RESEARCH_STUDENT"]:
            return bool(proposal.research_abstract_file)
        return True

    if step == 19:
        if proposal.extension_type in ["RESEARCH_FACULTY", "RESEARCH_STUDENT"]:
            return bool(proposal.certificate_of_completion_file)
        return True

    return False


def build_wizard_steps(proposal, current_step, comment_counts=None):
    completed = set(proposal.completed_steps or [])
    skipped = set(proposal.skipped_steps or [])
    comment_counts = comment_counts or {}

    steps = []
    for item in STEP_LABELS:
        no = item["no"]
        if no == current_step:
            state = "current"
        elif no in completed:
            state = "completed"
        elif no in skipped:
            state = "skipped"
        else:
            state = "upcoming"

        ccount = int(comment_counts.get(no, 0) or 0)
        steps.append({**item, "state": state, "comment_count": ccount, "has_comment": ccount > 0})
    return steps


def _update_creator_role(proposal):
    creator_pp = proposal.proponents.filter(user=proposal.created_by).first()
    if not creator_pp:
        return

    if proposal.scope_type == "PROGRAM":
        creator_pp.role = "Program Leader"
    elif proposal.scope_type == "PROJECT":
        creator_pp.role = "Project Leader"
    else:
        creator_pp.role = "Proponent"

    creator_pp.save(update_fields=["role"])


def _build_wizard_context(proposal, step, request_user, comment_counts=None):
    completed_count = len(proposal.completed_steps or [])
    progress = int((completed_count / TOTAL_STEPS) * 100) if TOTAL_STEPS else 0

    ctx = {
        "proposal": proposal,
        "step": step,
        "total_steps": TOTAL_STEPS,
        "progress": progress,
        "wizard_steps": build_wizard_steps(proposal, step, comment_counts=comment_counts),
    }

    active_cutoff = timezone.now() - timedelta(seconds=45)
    active_editors = (
        proposal.editor_presences
        .select_related("user", "user__profile")
        .filter(last_seen__gte=active_cutoff)
        .exclude(user=request_user)
    )
    ctx["active_editors"] = active_editors
    return ctx

def _build_status_flow(choices, progress_map):
    return [
        {
            "code": value,
            "label": label,
            "progress": progress_map.get(value, 0),
        }
        for value, label in choices
    ]


def services_home(request):
    sdgs = SDG.objects.all().order_by("code")
    thrusts = ExtensionThrust.objects.all().order_by("name")
    process_records = (
        ExtensionProcess.objects
        .prefetch_related(
            Prefetch(
                "steps",
                queryset=ProcessStep.objects.order_by("order", "id"),
            )
        )
        .order_by("order", "id")
    )

    workflow_phases = [
        {
            "key": "proposal",
            "label": "Proposal",
            "summary": "Drafting, review, revision, printing, signed upload, and approval document release.",
            "weight": "40% of overall progress when MOA is required",
            "statuses": _build_status_flow(
                Proposal.ProposalStatus.choices,
                Proposal.PROPOSAL_PROGRESS_MAP,
            ),
        },
        {
            "key": "moa",
            "label": "MOA",
            "summary": "Optional agreement routing from draft through legal review, certification, agenda, and completion.",
            "weight": "20% of overall progress when required",
            "statuses": _build_status_flow(
                Proposal.MOAStatus.choices,
                Proposal.MOA_PROGRESS_MAP,
            ),
        },
        {
            "key": "implementation",
            "label": "Implementation",
            "summary": "Preparation, monitoring, progress reporting, terminal reporting, review, revision, and completion.",
            "weight": "40% of overall progress when MOA is required",
            "statuses": _build_status_flow(
                Proposal.ImplementationStatus.choices,
                Proposal.IMPLEMENTATION_PROGRESS_MAP,
            ),
        },
    ]

    context = {
        "sdgs": sdgs,
        "thrusts": thrusts,
        "process_records": process_records,
        "workflow_phases": workflow_phases,
        "wizard_steps": STEP_LABELS,
        "total_wizard_steps": TOTAL_STEPS,
    }
    return render(request, "services/services_home.html", context)


@login_required
@faculty_like_required
def proposal_create(request):
    profile = getattr(request.user, "profile", None)

    proposal = Proposal.objects.create(
        created_by=request.user,
        campus=getattr(profile, "campus", "") or "",
        college=getattr(profile, "college", "") or "",
        department=getattr(profile, "department", "") or "",
        current_step=1,
    )

    ProposalProponent.objects.get_or_create(
        proposal=proposal,
        user=request.user,
        defaults={
            "full_name": getattr(profile, "full_name", request.user.username),
            "email": request.user.email or "",
            "role": "Proponent",
        },
    )

    _update_creator_role(proposal)
    return redirect("proposal_wizard", proposal_id=proposal.id, step=1)


@login_required
@faculty_like_required
def proposal_wizard(request, proposal_id, step):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_view_proposal(request.user, proposal):
        messages.error(request, "You don't have access to this proposal.")
        return redirect("dashboard_redirect")

    step = max(1, min(step, TOTAL_STEPS))
    can_edit = _can_edit(request.user, proposal)
    can_review = _can_review(request.user, proposal)

    if can_edit:
        ProposalEditorPresence.objects.update_or_create(
            proposal=proposal,
            user=request.user,
            defaults={
                "last_seen": timezone.now(),
                "step": step,
            },
        )

    if can_edit and step > (proposal.current_step or 1):
        proposal.current_step = step
        proposal.save(update_fields=["current_step"])

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()

    # Sidebar step comment markers (count comments per step for the current review round).
    # We only show badges to proponents when the proposal is FOR_REVISION.
    step_comment_counts = {}
    if current_round and proposal.proposal_status == Proposal.ProposalStatus.FOR_REVISION:
        raw_steps = ProposalSectionComment.objects.filter(
            proposal=proposal,
            review_round=current_round,
        ).values_list("step_no", flat=True)
        temp = Counter()
        for st in raw_steps:
            try:
                st_no = int(st or 1)
            except (TypeError, ValueError):
                st_no = 1
            temp[st_no] += 1
        step_comment_counts = dict(temp)


    is_proponent_view = (
        request.user == proposal.created_by
        or proposal.proponents.filter(user=request.user).exists()
    )

    review_summary = None
    visible_comments = []

    if is_proponent_view and current_round:
        review_summary = ProposalCommentSummary.objects.filter(
            proposal=proposal,
            review_round=current_round,
            sent_to_proponent=True,
        ).first()

        if proposal.proposal_status == Proposal.ProposalStatus.FOR_REVISION:
            visible_comments = (
                ProposalSectionComment.objects.filter(
                    proposal=proposal,
                    review_round=current_round,
                    step_no=step,
                )
                .select_related("reviewer", "reviewer__profile")
                .order_by("step_no", "created_at")
            )
        else:
            visible_comments = []

    # 🔥 IMPORTANT: ctx must exist first
    ctx = _build_wizard_context(proposal, step, request.user, comment_counts=step_comment_counts)

    # add review panel data
    ctx["review_summary"] = review_summary
    ctx["visible_comments"] = visible_comments
    ctx["show_readonly_review_panel"] = bool(review_summary or visible_comments)

    reviewer_role = _get_reviewer_role(request.user, proposal, current_round)
    can_comment = bool(reviewer_role and current_round)

    action = request.POST.get("action", "next") if request.method == "POST" else "next"

    if request.method == "POST":
        if action == "save_comment":
            if not can_comment:
                messages.error(request, "You are not allowed to comment on this proposal.")
                return redirect("proposal_wizard", proposal_id=proposal.id, step=step)
        else:
            if not can_edit:
                messages.error(request, "You can review this proposal, but you cannot edit it.")
                return redirect("proposal_wizard", proposal_id=proposal.id, step=step)

            editable_statuses = {
                Proposal.ProposalStatus.DRAFTING,
                Proposal.ProposalStatus.FOR_REVISION,
            }
            if proposal.is_locked and proposal.proposal_status not in editable_statuses:
                messages.warning(request, "This proposal is currently locked for editing.")
                return redirect("proposal_wizard", proposal_id=proposal.id, step=step)

    ctx["is_edit_mode"] = proposal.proposal_status in [
        Proposal.ProposalStatus.DRAFTING,
        Proposal.ProposalStatus.FOR_REVISION,
    ]
    ctx["is_review_mode"] = can_review and not can_edit
    ctx["current_review_round"] = current_round
    ctx["can_comment"] = can_comment
    ctx["reviewer_role"] = reviewer_role
    ctx["can_edit_proposal"] = can_edit

    if ctx["can_comment"]:
        ctx["existing_step_comment"] = ProposalSectionComment.objects.filter(
            proposal=proposal,
            review_round=current_round,
            reviewer=request.user,
            step_no=step,
        ).first()

        ctx["step_comments"] = ProposalSectionComment.objects.filter(
            proposal=proposal,
            review_round=current_round,
            step_no=step,
        ).select_related("reviewer", "reviewer__profile").order_by("created_at")
    else:
        ctx["existing_step_comment"] = None
        ctx["step_comments"] = []

    template = f"services/wizard/step_{step}.html"

    if request.method == "GET":
        if step == 2 and proposal.scope_type == "PROGRAM":
            ctx["program_projects"] = proposal.program_projects.all().order_by("order", "id")

        if step == 3:
            ctx["proponents"] = proposal.proponents.select_related("user").all().order_by("id")
            if proposal.scope_type == "PROGRAM":
                ctx["program_projects"] = proposal.program_projects.select_related("leader_user").all().order_by("order", "id")

        if step == 6:
            ctx["sdgs"] = SDG_LIST
            ctx["thrusts"] = THRUST_LIST
            sdg_links = proposal.sdg_links.all()
            thrust_links = proposal.thrust_links.all()
            ctx["selected_sdg_codes"] = set(sdg_links.values_list("sdg_code", flat=True))
            ctx["selected_thrust_names"] = set(thrust_links.values_list("thrust_name", flat=True))
            ctx["sdg_explanations"] = {item.sdg_code: item.explanation for item in sdg_links}
            ctx["thrust_explanations"] = {item.thrust_name: item.explanation for item in thrust_links}

        if step == 7:
            ctx["budgetary_requirement"] = proposal.budgetary_requirement or ""

        if step == 8:
            ctx["sex_total"] = (proposal.sex_male or 0) + (proposal.sex_female or 0)
            ctx["gender_total"] = (
                (proposal.g_lesbian or 0)
                + (proposal.g_gay or 0)
                + (proposal.g_bisexual or 0)
                + (proposal.g_transgender or 0)
                + (proposal.g_straight or 0)
                + (proposal.g_others or 0)
            )

        if step == 9:
            ctx["gender_issues"] = GENDER_ISSUE_LIST
            ctx["selected_gender_issue_keys"] = set(
                proposal.gender_issue_links.values_list("issue_key", flat=True)
            )
            others_item = proposal.gender_issue_links.filter(issue_key="others").first()
            ctx["gender_issue_other_text"] = others_item.other_text if others_item else ""

        if step == 10:
            ctx["estimated_month"] = proposal.estimated_month or ""
            ctx["estimated_year"] = proposal.estimated_year or ""
            ctx["extension_venue"] = proposal.extension_venue or ""
            ctx["month_choices"] = [
                "January", "February", "March", "April", "May", "June",
                "July", "August", "September", "October", "November", "December",
            ]

        if step == 11:
            ctx["rationale_background"] = proposal.rationale_background or ""

        if step == 12:
            ctx["significance"] = proposal.significance or ""

        if step == 13:
            ctx["general_objective"] = proposal.general_objective or ""
            if proposal.scope_type == "PROGRAM":
                projects = proposal.program_projects.all().order_by("order", "id")
                ctx["program_projects"] = projects
                ctx["project_objectives_map"] = {
                    prj.id: list(
                        proposal.specific_objectives.filter(program_project=prj)
                        .values_list("objective", flat=True)
                    )
                    for prj in projects
                }
            else:
                ctx["specific_objectives"] = list(
                    proposal.specific_objectives.filter(program_project__isnull=True)
                    .values_list("objective", flat=True)
                )

        if step == 14:
            ctx["methodologies"] = list(proposal.methodologies.values_list("item", flat=True))

        if step == 15:
            ctx["output_outcomes"] = list(proposal.output_outcomes.values_list("item", flat=True))

        if step == 16:
            ctx["existing_attachments"] = proposal.attachments.filter(
                category=ProposalAttachment.Category.DETAILS_OF_ACTIVITIES
            ).order_by("id")
            if proposal.scope_type == "PROGRAM":
                ctx["program_projects"] = proposal.program_projects.all().order_by("order", "id")

        if step == 17:
            ctx["existing_funding_attachments"] = proposal.attachments.filter(
                category=ProposalAttachment.Category.OTHER
            ).order_by("id")

        if step == 18:
            ctx["requires_abstract"] = proposal.extension_type in ["RESEARCH_FACULTY", "RESEARCH_STUDENT"]

        if step == 19:
            ctx["requires_certificate"] = proposal.extension_type in ["RESEARCH_FACULTY", "RESEARCH_STUDENT"]

        return render(request, template, ctx)

    if action == "save_comment":
        comment_text = (request.POST.get("comment") or "").strip()
        if not comment_text:
            messages.error(request, "Comment cannot be empty.")
            return redirect("proposal_wizard", proposal_id=proposal.id, step=step)

        existing_step_comment = ProposalSectionComment.objects.filter(
            proposal=proposal,
            review_round=current_round,
            reviewer=request.user,
            step_no=step,
        ).first()

        if existing_step_comment:
            existing_step_comment.comment = comment_text
            existing_step_comment.reviewer_role = reviewer_role
            existing_step_comment.save(update_fields=["comment", "reviewer_role"])
            messages.success(request, f"Your comment for Step {step} was updated.")
        else:
            ProposalSectionComment.objects.create(
                proposal=proposal,
                review_round=current_round,
                reviewer=request.user,
                reviewer_role=reviewer_role,
                step_no=step,
                comment=comment_text,
            )
            messages.success(request, f"Your comment for Step {step} was saved.")

        if proposal.proposal_status == Proposal.ProposalStatus.SUBMITTED_FOR_REVIEW:
            proposal.mark_in_review(review_level=proposal.review_level or Proposal.ReviewLevel.DEPARTMENT)

        return redirect("proposal_wizard", proposal_id=proposal.id, step=step)

    if step == 1:
        proposal.extension_type = request.POST.get("extension_type", "")
        proposal.scope_type = request.POST.get("scope_type", "")
        research_title = (request.POST.get("research_title") or "").strip()

        if proposal.extension_type in ["RESEARCH_FACULTY", "RESEARCH_STUDENT"]:
            proposal.research_title = research_title
        else:
            proposal.research_title = ""

        if proposal.extension_type in ["RESEARCH_FACULTY", "RESEARCH_STUDENT"] and not research_title and action != "skip":
            messages.error(request, "Research Title is required for research-based extension type.")
            return redirect("proposal_wizard", proposal_id=proposal.id, step=1)

        proposal.save(update_fields=["extension_type", "scope_type", "research_title"])
        _update_creator_role(proposal)

    elif step == 2:
        proposal.title = (request.POST.get("title") or "").strip()
        proposal.save(update_fields=["title"])

        if proposal.scope_type == "PROGRAM":
            raw_ids = request.POST.getlist("project_id[]")
            raw_titles = request.POST.getlist("project_titles[]")

            max_len = max(len(raw_ids), len(raw_titles), 0)
            raw_ids += [""] * (max_len - len(raw_ids))
            raw_titles += [""] * (max_len - len(raw_titles))

            existing = {str(p.id): p for p in proposal.program_projects.all()}
            keep_db_ids = []
            to_update = []
            to_create = []

            for i in range(max_len):
                pid = (raw_ids[i] or "").strip()
                clean_title = _strip_phase_prefix(raw_titles[i])

                if not clean_title:
                    continue

                order_no = len(keep_db_ids) + len(to_create) + 1
                stored_title = f"Phase {_to_roman(order_no)} {clean_title}"

                if pid and pid in existing:
                    prj = existing[pid]
                    prj.title = stored_title
                    prj.order = order_no
                    to_update.append(prj)
                    keep_db_ids.append(prj.id)
                else:
                    to_create.append(
                        ProgramProject(
                            proposal=proposal,
                            title=stored_title,
                            order=order_no,
                        )
                    )

            proposal.program_projects.exclude(id__in=keep_db_ids).delete()

            if to_update:
                ProgramProject.objects.bulk_update(to_update, ["title", "order"])
            if to_create:
                ProgramProject.objects.bulk_create(to_create)

    elif step == 3:
        remove_ids = request.POST.getlist("remove_proponent_ids")
        if remove_ids:
            ProposalProponent.objects.filter(
                proposal=proposal,
                id__in=remove_ids,
            ).exclude(user=proposal.created_by).delete()

        add_user_id = (request.POST.get("add_user_id") or "").strip()
        if add_user_id.isdigit():
            user_obj = User.objects.filter(id=int(add_user_id)).first()
            if user_obj:
                prof = getattr(user_obj, "profile", None)
                ProposalProponent.objects.get_or_create(
                    proposal=proposal,
                    user=user_obj,
                    defaults={
                        "full_name": getattr(prof, "full_name", user_obj.username),
                        "email": user_obj.email or "",
                        "role": "Proponent",
                        "designation": "",
                        "specialization": "",
                        "cp_number": "",
                    },
                )

        for p in proposal.proponents.all():
            prefix = f"p_{p.id}_"
            p.designation = request.POST.get(prefix + "designation", p.designation)
            p.specialization = request.POST.get(prefix + "specialization", p.specialization)
            p.cp_number = request.POST.get(prefix + "cp_number", p.cp_number)
            p.email = request.POST.get(prefix + "email", p.email)
            p.save(update_fields=["designation", "specialization", "cp_number", "email"])

        _update_creator_role(proposal)

        if proposal.scope_type == "PROGRAM":
            proposal.proponents.exclude(user=proposal.created_by).update(role="Proponent")

            for prj in proposal.program_projects.all():
                uid = (request.POST.get(f"project_leader_{prj.id}") or "").strip()
                if uid.isdigit():
                    prj.leader_user_id = int(uid)
                    prj.save(update_fields=["leader_user"])

                    pp = proposal.proponents.filter(user_id=int(uid)).first()
                    if pp and pp.user_id != proposal.created_by_id:
                        pp.role = "Project Leader"
                        pp.save(update_fields=["role"])
                else:
                    prj.leader_user = None
                    prj.save(update_fields=["leader_user"])

        if action in ("add_member", "save_members"):
            if is_step_complete(proposal, step):
                mark_step_completed(proposal, step)
            else:
                unmark_step_completed(proposal, step)

            proposal.save(update_fields=["completed_steps", "skipped_steps"])
            messages.success(request, "Members updated.")
            return redirect("proposal_wizard", proposal_id=proposal.id, step=3)

    elif step == 4:
        proposal.implementing_agency = (request.POST.get("implementing_agency") or "").strip()
        proposal.save(update_fields=["implementing_agency"])

    elif step == 5:
        raw_beneficiaries = request.POST.get("beneficiaries_count")
        proposal.beneficiaries_count = _to_int(raw_beneficiaries, default=None)
        if proposal.beneficiaries_count == 0 and (raw_beneficiaries or "").strip() == "":
            proposal.beneficiaries_count = None
        proposal.beneficiaries_who = (request.POST.get("beneficiaries_who") or "").strip()
        proposal.save(update_fields=["beneficiaries_count", "beneficiaries_who"])

    elif step == 6:
        sdg_codes = request.POST.getlist("sdg_codes")
        thrust_names = request.POST.getlist("thrust_names")

        ProposalSDG.objects.filter(proposal=proposal).delete()
        ProposalThrust.objects.filter(proposal=proposal).delete()

        for code in sdg_codes:
            code = (code or "").strip()
            if code:
                explanation = (request.POST.get(f"sdg_explanation_{code}") or "").strip()
                ProposalSDG.objects.create(
                    proposal=proposal,
                    sdg_code=code,
                    explanation=explanation,
                )

        for name in thrust_names:
            name = (name or "").strip()
            if name:
                explanation = (request.POST.get(f"thrust_explanation_{name}") or "").strip()
                ProposalThrust.objects.create(
                    proposal=proposal,
                    thrust_name=name,
                    explanation=explanation,
                )

    elif step == 7:
        proposal.budgetary_requirement = (request.POST.get("budgetary_requirement") or "").strip()
        proposal.save(update_fields=["budgetary_requirement"])

    elif step == 8:
        sex_male = _to_int(request.POST.get("sex_male"))
        sex_female = _to_int(request.POST.get("sex_female"))
        g_lesbian = _to_int(request.POST.get("g_lesbian"))
        g_gay = _to_int(request.POST.get("g_gay"))
        g_bisexual = _to_int(request.POST.get("g_bisexual"))
        g_transgender = _to_int(request.POST.get("g_transgender"))
        g_straight = _to_int(request.POST.get("g_straight"))
        g_others = _to_int(request.POST.get("g_others"))

        sex_total = sex_male + sex_female
        gender_total = g_lesbian + g_gay + g_bisexual + g_transgender + g_straight + g_others

        proposal.sex_male = sex_male
        proposal.sex_female = sex_female
        proposal.g_lesbian = g_lesbian
        proposal.g_gay = g_gay
        proposal.g_bisexual = g_bisexual
        proposal.g_transgender = g_transgender
        proposal.g_straight = g_straight
        proposal.g_others = g_others
        proposal.save(update_fields=[
            "sex_male", "sex_female", "g_lesbian", "g_gay",
            "g_bisexual", "g_transgender", "g_straight", "g_others",
        ])

        if action == "next" and sex_total != gender_total:
            messages.error(request, "Sex total and Gender total must be the same before you can proceed.")
            return redirect("proposal_wizard", proposal_id=proposal.id, step=8)

    elif step == 9:
        selected_keys = request.POST.getlist("gender_issue_keys")
        other_text = (request.POST.get("gender_issue_other_text") or "").strip()

        ProposalGenderIssue.objects.filter(proposal=proposal).delete()
        label_map = dict(GENDER_ISSUE_LIST)

        for key in selected_keys:
            key = (key or "").strip()
            if not key or key not in label_map:
                continue

            ProposalGenderIssue.objects.create(
                proposal=proposal,
                issue_key=key,
                issue_label=label_map[key],
                other_text=other_text if key == "others" else "",
            )

    elif step == 10:
        estimated_month = (request.POST.get("estimated_month") or "").strip()
        estimated_year_raw = (request.POST.get("estimated_year") or "").strip()
        extension_venue = (request.POST.get("extension_venue") or "").strip()

        proposal.estimated_month = estimated_month or ""
        proposal.estimated_year = int(estimated_year_raw) if estimated_year_raw.isdigit() else None
        proposal.extension_venue = extension_venue
        proposal.save(update_fields=["estimated_month", "estimated_year", "extension_venue"])

    elif step == 11:
        proposal.rationale_background = (request.POST.get("rationale_background") or "").strip()
        proposal.save(update_fields=["rationale_background"])

    elif step == 12:
        proposal.significance = (request.POST.get("significance") or "").strip()
        proposal.save(update_fields=["significance"])

    elif step == 13:
        proposal.general_objective = (request.POST.get("general_objective") or "").strip()
        proposal.save(update_fields=["general_objective"])

        proposal.specific_objectives.all().delete()

        if proposal.scope_type == "PROGRAM":
            for prj in proposal.program_projects.all():
                objectives = request.POST.getlist(f"specific_objectives_{prj.id}[]")
                for obj in objectives:
                    obj = (obj or "").strip()
                    if obj:
                        ProposalSpecificObjective.objects.create(
                            proposal=proposal,
                            program_project=prj,
                            objective=obj,
                        )
        else:
            objectives = request.POST.getlist("specific_objectives[]")
            for obj in objectives:
                obj = (obj or "").strip()
                if obj:
                    ProposalSpecificObjective.objects.create(
                        proposal=proposal,
                        program_project=None,
                        objective=obj,
                    )

    elif step == 14:
        proposal.methodologies.all().delete()
        for item in request.POST.getlist("methodologies[]"):
            item = (item or "").strip()
            if item:
                ProposalMethodology.objects.create(proposal=proposal, item=item)

    elif step == 15:
        proposal.output_outcomes.all().delete()
        for item in request.POST.getlist("output_outcomes[]"):
            item = (item or "").strip()
            if item:
                ProposalOutputOutcome.objects.create(proposal=proposal, item=item)

    elif step == 16:
        remove_attachment_ids = request.POST.getlist("remove_attachment_ids")
        if remove_attachment_ids:
            ProposalAttachment.objects.filter(
                proposal=proposal,
                category=ProposalAttachment.Category.DETAILS_OF_ACTIVITIES,
                id__in=remove_attachment_ids,
            ).delete()

        changed_fields = []
        if request.FILES.get("work_plan_file"):
            proposal.work_plan_file = request.FILES["work_plan_file"]
            changed_fields.append("work_plan_file")

        if request.FILES.get("gantt_chart_file"):
            proposal.gantt_chart_file = request.FILES["gantt_chart_file"]
            changed_fields.append("gantt_chart_file")

        if changed_fields:
            proposal.save(update_fields=changed_fields)

        for f in request.FILES.getlist("attachment_files"):
            if f:
                ProposalAttachment.objects.create(
                    proposal=proposal,
                    file=f,
                    category=ProposalAttachment.Category.DETAILS_OF_ACTIVITIES,
                    label=getattr(f, "name", ""),
                )

    elif step == 17:
        if request.FILES.get("funding_file"):
            proposal.funding_file = request.FILES["funding_file"]
            proposal.save(update_fields=["funding_file"])

    elif step == 18:
        if request.FILES.get("research_abstract_file"):
            proposal.research_abstract_file = request.FILES["research_abstract_file"]
            proposal.save(update_fields=["research_abstract_file"])

    elif step == 19:
        if request.FILES.get("certificate_of_completion_file"):
            proposal.certificate_of_completion_file = request.FILES["certificate_of_completion_file"]
            proposal.save(update_fields=["certificate_of_completion_file"])

    if action == "back":
        return redirect("proposal_wizard", proposal_id=proposal.id, step=max(1, step - 1))

    if action == "skip":
        mark_step_skipped(proposal, step)
        proposal.save(update_fields=["completed_steps", "skipped_steps"])
        messages.info(request, "Skipped.")
        return redirect("proposal_wizard", proposal_id=proposal.id, step=min(TOTAL_STEPS, step + 1))

    if is_step_complete(proposal, step):
        mark_step_completed(proposal, step)
    else:
        unmark_step_completed(proposal, step)

    proposal.save(update_fields=["completed_steps", "skipped_steps"])

    if step >= TOTAL_STEPS:
        messages.success(request, "All steps completed.")
        return redirect("proposal_submit", proposal_id=proposal.id)

    messages.success(request, "Draft saved.")
    return redirect("proposal_wizard", proposal_id=proposal.id, step=step + 1)


@login_required
@faculty_like_required
def proposal_submit(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_edit(request.user, proposal):
        messages.error(request, "No access.")
        return redirect("services_home")

    editable_statuses = {
        Proposal.ProposalStatus.DRAFTING,
        Proposal.ProposalStatus.FOR_REVISION,
    }
    if proposal.is_locked and proposal.proposal_status not in editable_statuses:
        messages.warning(request, "This proposal has already been submitted or is not editable.")
        return redirect("services_home")

    all_required_steps = set(range(1, TOTAL_STEPS + 1))
    completed_steps = set(proposal.completed_steps or [])

    if not all_required_steps.issubset(completed_steps):
        messages.error(request, "Please complete all required steps before submitting.")
        return redirect("proposal_wizard", proposal_id=proposal.id, step=proposal.current_step)

    if request.method == "POST":
        if proposal.proposal_status == Proposal.ProposalStatus.FOR_REVISION:

            current_round = proposal.get_current_review_round()
            if current_round:
                current_round.is_closed = True
                current_round.ready_for_staff_summary = False
                current_round.save()

            proposal.is_locked = False
            proposal.start_review_round()
            _ensure_open_review_round(proposal, request.user)

            proposal.transition_proposal_status(
                Proposal.ProposalStatus.SUBMITTED_FOR_REVIEW
            )

            messages.success(request, "Proposal resubmitted successfully.")
        else:
            proposal.lock_and_submit()
            proposal.start_review_round()
            _ensure_open_review_round(proposal, request.user)
            messages.success(request, "Proposal submitted successfully.")

        return redirect("services_home")

    return render(request, "services/proposal_submit.html", {"proposal": proposal})


@login_required
@faculty_like_required
def proposal_share(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if proposal.created_by_id != request.user.id:
        messages.error(request, "Only the creator can manage sharing.")
        return redirect("proposal_wizard", proposal_id=proposal.id, step=proposal.current_step)

    if request.method == "POST":
        query = (request.POST.get("q") or "").strip()
        if not query:
            messages.error(request, "Enter username or email.")
            return redirect("proposal_share", proposal_id=proposal.id)

        user_obj = User.objects.filter(Q(username__iexact=query) | Q(email__iexact=query)).first()
        if not user_obj:
            messages.error(request, "User not found.")
            return redirect("proposal_share", proposal_id=proposal.id)

        ProposalCollaborator.objects.get_or_create(
            proposal=proposal,
            user=user_obj,
            defaults={"can_edit": True},
        )
        messages.success(request, f"Shared with {user_obj.username}.")
        return redirect("proposal_share", proposal_id=proposal.id)

    collaborators = proposal.collaborators.select_related("user").all()
    return render(
        request,
        "services/proposal_share.html",
        {"proposal": proposal, "collaborators": collaborators},
    )


@login_required
def proposal_editor_ping(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_edit(request.user, proposal):
        return JsonResponse({"ok": False}, status=403)

    step = request.GET.get("step")
    try:
        step = int(step)
    except (TypeError, ValueError):
        step = 1

    ProposalEditorPresence.objects.update_or_create(
        proposal=proposal,
        user=request.user,
        defaults={
            "last_seen": timezone.now(),
            "step": step,
        },
    )

    return JsonResponse({"ok": True})


@login_required
@faculty_like_required
def title_suggest(request):
    query = (request.GET.get("q") or "").strip()
    if len(query) < 3:
        return JsonResponse({"suggestions": []})

    suggestions = (
        Proposal.objects.filter(title__icontains=query)
        .values_list("title", flat=True)
        .distinct()[:8]
    )
    return JsonResponse({"suggestions": list(suggestions)})


@require_GET
@login_required
@faculty_like_required
def proponent_search(request):
    query = (request.GET.get("q") or "").strip()
    if len(query) < 2:
        return JsonResponse({"results": []})

    users = (
        User.objects.select_related("profile")
        .filter(
            Q(username__icontains=query)
            | Q(email__icontains=query)
            | Q(profile__full_name__icontains=query)
        )
        .distinct()[:10]
    )

    results = []
    for user_obj in users:
        prof = getattr(user_obj, "profile", None)
        results.append({
            "id": user_obj.id,
            "username": user_obj.username,
            "name": getattr(prof, "full_name", user_obj.username),
            "email": user_obj.email or "",
            "designation": getattr(prof, "department", "") or "",
            "campus": getattr(prof, "campus", "") or "",
        })

    return JsonResponse({"results": results})


@login_required
@require_POST
def proposal_return_for_revision(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _is_director(request.user):
        messages.error(request, "Only the Director can return this proposal for revision.")
        return redirect("dashboard_redirect")

    proposal.return_for_revision()

    messages.success(request, "Proposal returned for revision.")
    return redirect("dashboard_redirect")


@login_required
@require_POST
def proposal_approve(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _is_director(request.user):
        messages.error(request, "Only the Director can approve proposals.")
        return redirect("dashboard_redirect")

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    if current_round and not current_round.is_closed:
        current_round.is_closed = True
        current_round.save(update_fields=["is_closed"])

        ProposalEvaluatorAssignment.objects.filter(
            proposal=proposal,
            review_round=current_round,
            is_active=True,
        ).update(is_active=False)

    # Director approval here means: cleared for printing (not final approval yet)
    proposal.mark_ready_for_printing()

    messages.success(request, "Proposal cleared for printing.")
    return redirect("dashboard_redirect")


@login_required
def proposal_review_comments(request, proposal_id, step=1):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not (
        _can_view_proposal(request.user, proposal)
        or request.user == proposal.created_by
        or proposal.proponents.filter(user=request.user).exists()
    ):
        messages.error(request, "You don't have access to this proposal.")
        return redirect("dashboard_redirect")

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    if not current_round:
        messages.warning(request, "No review comments are available yet.")
        return redirect("proposal_wizard", proposal_id=proposal.id, step=1)

    step_comments = (
        ProposalSectionComment.objects
        .filter(
            proposal=proposal,
            review_round=current_round,
        )
        .select_related("reviewer", "reviewer__profile", "review_round")
        .order_by("step_no", "created_at")
    )

    grouped_comments = OrderedDict()
    for item in step_comments:
        step_no = item.step_no or 1
        if step_no not in grouped_comments:
            step_meta = next((s for s in STEP_LABELS if s["no"] == step_no), None)
            grouped_comments[step_no] = {
                "step_no": step_no,
                "step_title": step_meta["title"] if step_meta else f"Step {step_no}",
                "step_desc": step_meta["desc"] if step_meta else "",
                "comments": [],
            }
        grouped_comments[step_no]["comments"].append(item)

    summary = ProposalCommentSummary.objects.filter(
        proposal=proposal,
        review_round=current_round,
        sent_to_proponent=True
    ).first()

    context = {
        "proposal": proposal,
        "review_round": current_round,
        "grouped_comments": grouped_comments,
        "current_step": int(step or 1),
        "summary": summary,
    }
    return render(request, "services/review/proposal_review_summary.html", context)


def _serialize_user_for_eval(user):
    profile = getattr(user, "profile", None)
    return {
        "id": user.id,
        "username": getattr(user, "username", ""),
        "name": getattr(profile, "full_name", None) or getattr(user, "username", ""),
        "profile": {
            "full_name": getattr(profile, "full_name", None),
        },
    }


@login_required
@require_POST
def proposal_assign_evaluator(request, proposal_id, evaluator_id=None):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _is_director(request.user):
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "message": "Only the Director can assign evaluators."}, status=403)
        messages.error(request, "Only the Director can assign evaluators.")
        return redirect("dashboard_redirect")

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    if not current_round:
        current_round = proposal.start_review_round()

    selected_evaluator_id = evaluator_id if evaluator_id is not None else (request.POST.get("evaluator_id") or "").strip()
    if not str(selected_evaluator_id).isdigit():
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "message": "Invalid evaluator selected."}, status=400)
        messages.error(request, "Invalid evaluator selected.")
        return redirect("dashboard_redirect")

    evaluator = User.objects.filter(id=int(selected_evaluator_id)).select_related("profile").first()
    if not evaluator:
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "message": "Selected faculty member was not found."}, status=404)
        messages.error(request, "Selected faculty member was not found.")
        return redirect("dashboard_redirect")

    evaluator_profile = getattr(evaluator, "profile", None)
    allowed_roles = {"FACULTY", "DEPARTMENT_COORDINATOR", "CAMPUS_COORDINATOR"}

    if not evaluator_profile or evaluator_profile.role not in allowed_roles:
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "message": "Selected user is not eligible to be assigned as evaluator."}, status=400)
        messages.error(request, "Selected user is not eligible to be assigned as evaluator.")
        return redirect("dashboard_redirect")

    if proposal.proponents.filter(user=evaluator).exists():
        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"ok": False, "message": "A proponent of the proposal cannot be assigned as evaluator."}, status=400)
        messages.error(request, "A proponent of the proposal cannot be assigned as evaluator.")
        return redirect("dashboard_redirect")

    assignment, created = ProposalEvaluatorAssignment.objects.get_or_create(
        proposal=proposal,
        review_round=current_round,
        evaluator=evaluator,
        defaults={
            "assigned_by": request.user,
            "is_active": True,
            "is_completed": False,
        },
    )

    if not created:
        assignment.is_active = True
        assignment.is_completed = False
        assignment.assigned_by = request.user
        assignment.save(update_fields=["is_active", "is_completed", "assigned_by"])

    current_round.evaluator_review_required = True
    current_round.evaluator_review_done = False
    current_round.ready_for_staff_summary = False
    current_round.save(update_fields=["evaluator_review_required", "evaluator_review_done", "ready_for_staff_summary"])

    proposal.mark_in_review(review_level=Proposal.ReviewLevel.DIRECTOR)

    assigned_qs = ProposalEvaluatorAssignment.objects.filter(proposal=proposal, review_round=current_round, is_active=True).select_related("evaluator", "evaluator__profile")
    available_qs = User.objects.filter(profile__role__in=["FACULTY", "DEPARTMENT_COORDINATOR", "CAMPUS_COORDINATOR"]).exclude(id__in=assigned_qs.values_list("evaluator_id", flat=True)).exclude(id__in=proposal.proponents.values_list("user_id", flat=True)).select_related("profile").order_by("profile__full_name", "username")

    response = {
        "ok": True,
        "message": f"{getattr(evaluator_profile, 'full_name', evaluator.username)} has been assigned as evaluator.",
        "assigned_evaluators": [_serialize_user_for_eval(item.evaluator) for item in assigned_qs],
        "available_evaluators": [_serialize_user_for_eval(item) for item in available_qs],
    }

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return JsonResponse(response)

    messages.success(request, response["message"])
    return redirect("dashboard_redirect")


@login_required
@require_POST
def proposal_remove_evaluator(request, proposal_id, evaluator_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _is_director(request.user):
        return JsonResponse({"ok": False, "message": "Only the Director can remove evaluators."}, status=403)

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    if not current_round:
        return JsonResponse({"ok": False, "message": "No active review round found."}, status=404)

    assignment = ProposalEvaluatorAssignment.objects.filter(
        proposal=proposal,
        review_round=current_round,
        evaluator_id=evaluator_id,
        is_active=True,
    ).first()

    if assignment:
        assignment.is_active = False
        assignment.is_completed = False
        assignment.save(update_fields=["is_active", "is_completed"])

    active_assignments = ProposalEvaluatorAssignment.objects.filter(proposal=proposal, review_round=current_round, is_active=True)
    current_round.evaluator_review_required = active_assignments.exists()
    current_round.evaluator_review_done = current_round.evaluator_review_required and not active_assignments.filter(is_completed=False).exists()
    current_round.ready_for_staff_summary = False
    current_round.save(update_fields=["evaluator_review_required", "evaluator_review_done", "ready_for_staff_summary"])

    assigned_qs = ProposalEvaluatorAssignment.objects.filter(proposal=proposal, review_round=current_round, is_active=True).select_related("evaluator", "evaluator__profile")
    available_qs = User.objects.filter(profile__role__in=["FACULTY", "DEPARTMENT_COORDINATOR", "CAMPUS_COORDINATOR"]).exclude(id__in=assigned_qs.values_list("evaluator_id", flat=True)).exclude(id__in=proposal.proponents.values_list("user_id", flat=True)).select_related("profile").order_by("profile__full_name", "username")

    return JsonResponse({
        "ok": True,
        "assigned_evaluators": [_serialize_user_for_eval(item.evaluator) for item in assigned_qs],
        "available_evaluators": [_serialize_user_for_eval(item) for item in available_qs],
    })


@login_required
@require_POST
def proposal_complete_evaluation(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    if not current_round:
        messages.error(request, "No active review round found.")
        return redirect("dashboard_redirect")

    assignment = ProposalEvaluatorAssignment.objects.filter(
        proposal=proposal,
        review_round=current_round,
        evaluator=request.user,
        is_active=True,
    ).first()

    if not assignment:
        messages.error(request, "You are not assigned to evaluate this proposal.")
        return redirect("dashboard_redirect")

    has_comment = ProposalSectionComment.objects.filter(
        proposal=proposal,
        review_round=current_round,
        reviewer=request.user,
        reviewer_role="EVALUATOR",
    ).exists()

    if not has_comment:
        messages.error(request, "Please add at least one evaluator comment before completing the evaluation.")
        return redirect("dashboard_redirect")

    assignment.is_completed = True
    assignment.save(update_fields=["is_completed"])

    current_round.refresh_evaluator_review_done()

    messages.success(request, "Evaluation marked as completed.")
    return redirect("dashboard_redirect")


@login_required
@require_POST
def proposal_mark_ready_for_summary(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _is_director(request.user):
        messages.error(request, "Only the Director can mark a proposal ready for summary.")
        return redirect("dashboard_redirect")

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    # Ensure we operate on an OPEN round
    if not current_round or getattr(current_round, "is_closed", False):
        current_round = _ensure_open_review_round(proposal, request.user)
    if not current_round:
        messages.error(request, "No active review round found.")
        return redirect("dashboard_redirect")

    # Guard: don't re-queue a round that already has a SENT summary
    if ProposalCommentSummary.objects.filter(
        proposal=proposal,
        review_round=current_round,
        sent_to_proponent=True,
    ).exists():
        messages.error(request, "A summary has already been issued for the current review round.")
        return redirect("dashboard_redirect")

    director_has_comment = ProposalSectionComment.objects.filter(
        proposal=proposal,
        review_round=current_round,
        reviewer=request.user,
        reviewer_role="DIRECTOR",
    ).exists()

    if not director_has_comment:
        messages.error(request, "Please add at least one director comment before finishing the review round.")
        return redirect("dashboard_redirect")

    current_round.director_review_done = True
    current_round.save(update_fields=["director_review_done"])

    if current_round.evaluator_review_required:
        current_round.refresh_evaluator_review_done()

    try:
        current_round.mark_ready_for_staff_summary()
    except ValidationError as e:
        messages.error(request, e.messages[0])
        return redirect("dashboard_redirect")

    proposal.mark_in_review(review_level=Proposal.ReviewLevel.DIRECTOR)
    proposal.proposal_status = Proposal.ProposalStatus.IN_REVIEW
    proposal.save(update_fields=["proposal_status"])

    messages.success(request, "Proposal marked ready for staff summary.")
    return redirect("dashboard_redirect")


@login_required
@require_POST
def proposal_mark_department_review_done(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _is_department_coordinator(request.user):
        messages.error(request, "Only the Department Coordinator can mark department review done.")
        return redirect("dashboard_redirect")

    user_department = (getattr(request.user.profile, "department", "") or "").strip()
    proposal_department = (proposal.department or "").strip()

    if user_department != proposal_department:
        messages.error(request, "You can only mark review done for proposals in your department.")
        return redirect("dashboard_redirect")

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    if not current_round:
        messages.error(request, "No active review round found.")
        return redirect("dashboard_redirect")

    current_round.department_review_done = True
    current_round.save(update_fields=["department_review_done"])

    proposal.mark_in_review(review_level=Proposal.ReviewLevel.DEPARTMENT)

    messages.success(request, "Department review marked as done.")
    return redirect("dashboard_redirect")


@login_required
@require_POST
def proposal_mark_campus_review_done(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _is_campus_coordinator(request.user):
        messages.error(request, "Only the Campus Coordinator can mark campus review done.")
        return redirect("dashboard_redirect")

    user_campus = (getattr(request.user.profile, "campus", "") or "").strip()
    proposal_campus = (proposal.campus or "").strip()

    if user_campus != proposal_campus:
        messages.error(request, "You can only mark review done for proposals in your campus.")
        return redirect("dashboard_redirect")

    current_round = proposal.get_active_review_round() or proposal.get_current_review_round()
    if not current_round:
        messages.error(request, "No active review round found.")
        return redirect("dashboard_redirect")

    current_round.campus_review_done = True
    current_round.save(update_fields=["campus_review_done"])

    proposal.mark_in_review(review_level=Proposal.ReviewLevel.CAMPUS)

    messages.success(request, "Campus review marked as done.")
    return redirect("dashboard_redirect")


def set_estimated_row_height(ws, cell_ref, text, chars_per_line=70, base_height=15):
    text = str(text or "")
    row = ws[cell_ref].row

    if not text.strip():
        ws.row_dimensions[row].height = base_height
        ws[cell_ref].alignment = Alignment(
            wrap_text=True,
            vertical="top",
            horizontal="left",
        )
        return

    raw_lines = text.split("\n")
    visual_lines = 0

    for line in raw_lines:
        visual_lines += max(1, math.ceil(len(line or "") / chars_per_line))

    ws.row_dimensions[row].height = max(base_height, visual_lines * base_height + 4)
    ws[cell_ref].alignment = Alignment(
        wrap_text=True,
        vertical="top",
        horizontal="left",
    )


def copy_cell_style(src_cell, dst_cell):
    if src_cell.has_style:
        dst_cell._style = copy(src_cell._style)
    if src_cell.font:
        dst_cell.font = copy(src_cell.font)
    if src_cell.fill:
        dst_cell.fill = copy(src_cell.fill)
    if src_cell.border:
        dst_cell.border = copy(src_cell.border)
    if src_cell.alignment:
        dst_cell.alignment = copy(src_cell.alignment)
    if src_cell.number_format:
        dst_cell.number_format = src_cell.number_format
    if src_cell.protection:
        dst_cell.protection = copy(src_cell.protection)


def copy_row_heights(ws, src_start_row, src_end_row, target_start_row):
    for offset, src_row in enumerate(range(src_start_row, src_end_row + 1)):
        target_row = target_start_row + offset
        src_dim = ws.row_dimensions[src_row]
        if src_dim.height is not None:
            ws.row_dimensions[target_row].height = src_dim.height


def copy_block(ws, src_start_row, src_end_row, src_start_col, src_end_col, target_start_row):
    row_offset = target_start_row - src_start_row

    for src_row in range(src_start_row, src_end_row + 1):
        for src_col in range(src_start_col, src_end_col + 1):
            src_cell = ws.cell(row=src_row, column=src_col)
            dst_cell = ws.cell(row=src_row + row_offset, column=src_col)
            dst_cell.value = src_cell.value
            copy_cell_style(src_cell, dst_cell)

    copy_row_heights(ws, src_start_row, src_end_row, target_start_row)


def copy_merged_ranges_for_block(ws, src_start_row, src_end_row, target_start_row):
    row_offset = target_start_row - src_start_row
    merges_to_add = []

    for merged in list(ws.merged_cells.ranges):
        min_col = merged.min_col
        min_row = merged.min_row
        max_col = merged.max_col
        max_row = merged.max_row

        if src_start_row <= min_row and max_row <= src_end_row:
            new_min_row = min_row + row_offset
            new_max_row = max_row + row_offset
            new_range = f"{get_column_letter(min_col)}{new_min_row}:{get_column_letter(max_col)}{new_max_row}"
            merges_to_add.append(new_range)

    for rng in merges_to_add:
        ws.merge_cells(rng)


def replicate_phase_blocks(ws, phase_titles, block_start_row, block_end_row, block_start_col, block_end_col, title_cell_col=1):
    if not phase_titles:
        return

    block_height = block_end_row - block_start_row + 1

    first_title_cell = ws.cell(row=block_start_row, column=title_cell_col)
    first_title_cell.value = phase_titles[0]
    set_estimated_row_height(ws, first_title_cell.coordinate, phase_titles[0], chars_per_line=70)

    if len(phase_titles) == 1:
        return

    for idx, phase_title in enumerate(phase_titles[1:], start=1):
        new_block_start = block_start_row + (idx * block_height)
        ws.insert_rows(new_block_start, amount=block_height)

        copy_block(ws, block_start_row, block_end_row, block_start_col, block_end_col, new_block_start)
        copy_merged_ranges_for_block(ws, block_start_row, block_end_row, new_block_start)

        title_cell = ws.cell(row=new_block_start, column=title_cell_col)
        title_cell.value = phase_title
        set_estimated_row_height(ws, title_cell.coordinate, phase_title, chars_per_line=70)


def replicate_table_blocks_only(ws, copies_needed, block_start_row, block_end_row, block_start_col, block_end_col):
    if copies_needed <= 1:
        return

    block_height = block_end_row - block_start_row + 1

    for idx in range(1, copies_needed):
        new_block_start = block_start_row + (idx * block_height)
        ws.insert_rows(new_block_start, amount=block_height)

        copy_block(ws, block_start_row, block_end_row, block_start_col, block_end_col, new_block_start)
        copy_merged_ranges_for_block(ws, block_start_row, block_end_row, new_block_start)


def get_best_sheet(workbook, preferred_names):
    for name in preferred_names:
        if name in workbook.sheetnames:
            return workbook[name]
    return workbook[workbook.sheetnames[0]]


def replicate_funding_blocks(ws, phase_titles):
    if not phase_titles:
        return

    block_start_row = 3
    block_end_row = 10
    block_start_col = 1
    block_end_col = 5
    block_height = block_end_row - block_start_row + 1

    ws["A3"] = phase_titles[0]

    if len(phase_titles) == 1:
        return

    for idx, phase_title in enumerate(phase_titles[1:], start=1):
        new_block_start = block_start_row + (idx * block_height)
        ws.insert_rows(new_block_start, amount=block_height)
        copy_block(ws, block_start_row, block_end_row, block_start_col, block_end_col, new_block_start)
        copy_merged_ranges_for_block(ws, block_start_row, block_end_row, new_block_start)
        ws.cell(row=new_block_start, column=1).value = phase_title


@login_required
@faculty_like_required
def download_work_plan_template(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_edit(request.user, proposal):
        messages.error(request, "You don't have access to this draft.")
        return redirect("services_home")

    template_dir = Path(__file__).resolve().parent / "template_files"

    if proposal.scope_type == "PROGRAM":
        template_path = template_dir / "program_work_plan_template.xlsx"
        preferred_sheet_names = ["Program Work Plan", "Work Plan", "Sheet1", "Sheet"]
        filename_prefix = "Program"
    else:
        template_path = template_dir / "project_work_plan_template.xlsx"
        preferred_sheet_names = ["Project Work Plan", "Work Plan", "Sheet1", "Sheet"]
        filename_prefix = "Project"

    if not template_path.exists():
        messages.error(request, "Work Plan template file not found.")
        return redirect("proposal_wizard", proposal_id=proposal.id, step=16)

    wb = load_workbook(template_path)
    ws = get_best_sheet(wb, preferred_sheet_names)

    title_value = proposal.title or ""

    sdg_name_map = {code: name for code, name in SDG_LIST}
    sdg_lines = []
    for item in proposal.sdg_links.all().order_by("sdg_code"):
        sdg_title = sdg_name_map.get(item.sdg_code, item.sdg_code)
        line = f"SDG {item.sdg_code} - {sdg_title}"
        if (item.explanation or "").strip():
            line += f": {item.explanation.strip()}"
        sdg_lines.append(line)
    sdg_value = "\n".join(sdg_lines)

    thrust_lines = []
    for item in proposal.thrust_links.all().order_by("id"):
        line = item.thrust_name or ""
        if (item.explanation or "").strip():
            line += f": {item.explanation.strip()}"
        if line.strip():
            thrust_lines.append(line)
    thrust_value = "\n".join(thrust_lines)

    gender_lines = []
    for item in proposal.gender_issue_links.all().order_by("id"):
        if item.issue_key == "others":
            if (item.other_text or "").strip():
                gender_lines.append(f"• Others:\n      {item.other_text.strip()}")
        else:
            gender_lines.append(f"• {item.issue_label}")
    gender_value = "\n".join(gender_lines)

    ws["A3"] = title_value
    ws["A5"] = sdg_value
    ws["A7"] = thrust_value
    ws["A9"] = gender_value

    set_estimated_row_height(ws, "A3", title_value, chars_per_line=90)
    set_estimated_row_height(ws, "A5", sdg_value, chars_per_line=70)
    set_estimated_row_height(ws, "A7", thrust_value, chars_per_line=70)
    set_estimated_row_height(ws, "A9", gender_value, chars_per_line=65)

    if proposal.scope_type == "PROGRAM":
        phase_titles = [
            prj.title for prj in proposal.program_projects.all().order_by("order", "id")
            if (prj.title or "").strip()
        ]

        replicate_phase_blocks(
            ws=ws,
            phase_titles=phase_titles,
            block_start_row=12,
            block_end_row=24,
            block_start_col=1,
            block_end_col=3,
            title_cell_col=1,
        )

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    safe_title = proposal.title or "Proposal"
    response["Content-Disposition"] = f'attachment; filename="{filename_prefix}_Work_Plan_{safe_title}.xlsx"'
    return response


@login_required
@faculty_like_required
def download_gantt_chart_template(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_edit(request.user, proposal):
        messages.error(request, "You don't have access to this draft.")
        return redirect("services_home")

    template_dir = Path(__file__).resolve().parent / "template_files"

    if proposal.scope_type == "PROGRAM":
        template_path = template_dir / "program_gantt_chart_template.xlsx"
        download_name = f"Program_Gantt_Chart_{proposal.title or 'Proposal'}.xlsx"
    else:
        template_path = template_dir / "project_gantt_chart_template.xlsx"
        download_name = f"Project_Gantt_Chart_{proposal.title or 'Proposal'}.xlsx"

    if not template_path.exists():
        messages.error(request, "Gantt Chart template file not found.")
        return redirect("proposal_wizard", proposal_id=proposal.id, step=16)

    if proposal.scope_type != "PROGRAM":
        return FileResponse(
            open(template_path, "rb"),
            as_attachment=True,
            filename=download_name,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    wb = load_workbook(template_path)
    ws = wb[wb.sheetnames[0]]

    replicate_table_blocks_only(
        ws=ws,
        copies_needed=proposal.program_projects.count(),
        block_start_row=3,
        block_end_row=9,
        block_start_col=1,
        block_end_col=5,
    )

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{download_name}"'
    return response


@login_required
@faculty_like_required
def download_funding_template(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_edit(request.user, proposal):
        messages.error(request, "You don't have access to this draft.")
        return redirect("services_home")

    template_dir = Path(__file__).resolve().parent / "template_files"

    if proposal.scope_type == "PROGRAM":
        template_path = template_dir / "program_funding_template.xlsx"
        download_name = f"Program_Line-Item_Budget_{proposal.title or 'Proposal'}.xlsx"
    else:
        template_path = template_dir / "project_funding_template.xlsx"
        download_name = f"Project_Line-Item_Budget_{proposal.title or 'Proposal'}.xlsx"

    if not template_path.exists():
        messages.error(request, "Funding template file not found.")
        return redirect("proposal_wizard", proposal_id=proposal.id, step=17)

    wb = load_workbook(template_path)
    ws = wb["Work Plan Template"]

    if proposal.scope_type == "PROGRAM":
        phase_titles = [
            prj.title for prj in proposal.program_projects.all().order_by("order", "id")
            if (prj.title or "").strip()
        ]
        replicate_funding_blocks(ws, phase_titles)
    else:
        ws["A3"] = proposal.title or ""

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{download_name}"'
    return response


@login_required
def proposal_file_preview(request, proposal_id, file_type, attachment_id=None):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_view_proposal(request.user, proposal):
        raise Http404("You do not have access to this file.")

    file_field = None

    if file_type == "work_plan":
        file_field = proposal.work_plan_file
    elif file_type == "gantt_chart":
        file_field = proposal.gantt_chart_file
    elif file_type == "funding":
        file_field = proposal.funding_file
    elif file_type == "research_abstract":
        file_field = proposal.research_abstract_file
    elif file_type == "certificate":
        file_field = proposal.certificate_of_completion_file
    elif file_type == "attachment":
        if not attachment_id:
            raise Http404("Attachment not found.")
        attachment = get_object_or_404(ProposalAttachment, id=attachment_id, proposal=proposal)
        file_field = attachment.file
    else:
        raise Http404("Invalid file type.")

    if not file_field:
        raise Http404("File not found.")

    file_path = file_field.path
    filename = Path(file_path).name
    content_type, _ = mimetypes.guess_type(file_path)
    content_type = content_type or "application/octet-stream"

    response = FileResponse(open(file_path, "rb"), content_type=content_type)
    response["Content-Disposition"] = f'inline; filename="{quote(filename)}"'
    return response


# ==============================
# POST-APPROVAL (PRINT / UPLOAD / FINAL DOC RELEASE)
# ==============================

@login_required
@faculty_like_required
def proposal_download_approved_docx(request, proposal_id):
    """
    Proponent downloads the final proposal DOCX for printing.

    Workflow rule:
      - When Director has cleared the proposal (READY_FOR_PRINTING),
        the FIRST download transitions it to FOR_SUBMISSION_AND_UPLOAD.
      - Subsequent downloads do not change status.
    """

    # Prefetch everything the DOCX builder is likely to access to prevent N+1 queries.
    proposal = get_object_or_404(
        Proposal.objects
        .select_related("created_by", "created_by__profile")
        .prefetch_related(
            "proponents", "proponents__user", "proponents__user__profile",
            "collaborators", "collaborators__user", "collaborators__user__profile",
            "program_projects",
            "sdg_links",
            "thrust_links",
            "gender_issue_links",
            "specific_objectives",
            "methodologies",
            "output_outcomes",
            "attachments",
        ),
        id=proposal_id
    )

    # Permission check WITHOUT extra DB hits.
    is_proponent = (
        request.user.id == proposal.created_by_id
        or any((pp.user_id == request.user.id) for pp in proposal.proponents.all())
    )
    if not is_proponent:
        messages.error(request, "You do not have permission to download this proposal.")
        return redirect("dashboard_redirect")

    allowed_statuses = {
        Proposal.ProposalStatus.READY_FOR_PRINTING,
        Proposal.ProposalStatus.FOR_SUBMISSION_AND_UPLOAD,
        Proposal.ProposalStatus.APPROVED,
        Proposal.ProposalStatus.COMPLETED,
    }
    if proposal.proposal_status not in allowed_statuses:
        messages.error(request, "This proposal is not yet cleared for printing.")
        return redirect("dashboard_redirect")

    # Transition on first print/download (FAST PATH: one SQL UPDATE, no heavy save hooks)
    if proposal.proposal_status == Proposal.ProposalStatus.READY_FOR_PRINTING:
        with transaction.atomic():
            updated = Proposal.objects.filter(
                id=proposal.id,
                proposal_status=Proposal.ProposalStatus.READY_FOR_PRINTING,
            ).update(
                proposal_status=Proposal.ProposalStatus.FOR_SUBMISSION_AND_UPLOAD,
                last_saved_at=timezone.now(),
            )
            if updated:
                proposal.proposal_status = Proposal.ProposalStatus.FOR_SUBMISSION_AND_UPLOAD

    sdg_title_map = {code: name for code, name in SDG_LIST}

    # IMPORTANT: call as keyword-only to match your build_extension_form_docx signature
    data = build_extension_form_docx(proposal=proposal, sdg_title_map=sdg_title_map)

    filename_title = (proposal.title or proposal.research_title or "Proposal").strip()
    filename = f"{filename_title}.docx".replace("/", "-").replace("\\", "-")

    resp = HttpResponse(
        data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


@login_required
@faculty_like_required
def proposal_download_approval_document(request, proposal_id, document_type):
    """
    Download one released approval document and advance the proposal phase
    once the three approval documents have all been released and downloaded.
    """
    proposal = get_object_or_404(Proposal, id=proposal_id)

    is_proponent = (
        request.user == proposal.created_by
        or proposal.proponents.filter(user=request.user).exists()
    )
    if not (is_proponent or _is_staff(request.user)):
        messages.error(request, "You do not have permission to download this approval document.")
        return redirect("dashboard_redirect")

    doc_map = {
        "letter_of_award": ProposalFinalDocument.DocumentType.LETTER_OF_AWARD,
        "endorsement_for_approval": ProposalFinalDocument.DocumentType.ENDORSEMENT_FOR_APPROVAL,
        "extension_agreement": ProposalFinalDocument.DocumentType.EXTENSION_AGREEMENT,
    }
    final_doc_type = doc_map.get(document_type)
    if not final_doc_type:
        raise Http404("Invalid approval document type.")

    final_doc = proposal.final_documents.filter(document_type=final_doc_type).first()
    if not final_doc or not getattr(final_doc, "file", None):
        raise Http404("Approval document not found.")

    _mark_approval_documents_completed(proposal)

    file_path = final_doc.file.path
    filename = Path(file_path).name
    content_type, _ = mimetypes.guess_type(file_path)
    content_type = content_type or "application/octet-stream"

    response = FileResponse(open(file_path, "rb"), content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{quote(filename)}"'
    return response


@login_required
@faculty_like_required
@require_http_methods(["GET", "POST"])
def proposal_upload_signed_proposal(request, proposal_id):
    """
    Proponent uploads the signed proposal.

    Updated rule:
      - Upload does NOT auto-approve.
      - Staff verifies the signed proposal, then releases LOA / Endorsement / Extension Agreement.
    """
    proposal = get_object_or_404(Proposal, id=proposal_id)

    is_proponent = (
        request.user == proposal.created_by
        or proposal.proponents.filter(user=request.user).exists()
    )
    if not is_proponent:
        messages.error(request, "You do not have permission to upload files for this proposal.")
        return redirect("dashboard_redirect")

    if proposal.proposal_status != Proposal.ProposalStatus.FOR_SUBMISSION_AND_UPLOAD:
        messages.error(request, "This proposal is not yet ready for signed upload.")
        return redirect("dashboard_redirect")

    if request.method == "POST":
        signed_file = request.FILES.get("signed_proposal_file")
        if not signed_file:
            messages.error(request, "Please choose a signed proposal file to upload.")
            return redirect("proposal_upload_signed_proposal", proposal_id=proposal.id)

        # Keep a single final document record per proposal+type (no duplicates).
        ProposalFinalDocument.objects.update_or_create(
            proposal=proposal,
            document_type=ProposalFinalDocument.DocumentType.SIGNED_PROPOSAL,
            defaults={
                "file": signed_file,
                "uploaded_by": request.user,
                "remarks": "",
                "is_verified": False,  # STAFF will verify
            },
        )

        # Optional: keep a single attachment link for easy download.
        ProposalAttachment.objects.filter(
            proposal=proposal,
            category=ProposalAttachment.Category.OTHER,
            label="Signed Proposal",
        ).delete()
        ProposalAttachment.objects.create(
            proposal=proposal,
            file=signed_file,
            category=ProposalAttachment.Category.OTHER,
            label="Signed Proposal",
        )

        proposal.last_saved_at = timezone.now()
        proposal.save(update_fields=["last_saved_at"])

        messages.success(request, "Signed proposal uploaded. Awaiting staff verification.")
        return redirect("dashboard_redirect")

    signed_doc = ProposalFinalDocument.objects.filter(
        proposal=proposal,
        document_type=ProposalFinalDocument.DocumentType.SIGNED_PROPOSAL,
    ).first()

    return render(
        request,
        "services/post_approval/upload_signed_proposal.html",
        {"proposal": proposal, "signed_doc": signed_doc},
    )

@login_required
@role_required(["STAFF"])
@require_http_methods(["GET", "POST"])
def staff_release_approval_documents(request, proposal_id):
    """
    STAFF verifies signed proposal and uploads/releases the final approval documents:
      - Letter of Award
      - Endorsement for Approval
      - Extension Agreement

    Updated rule:
      - Proposal becomes APPROVED / CLAIMING only when ALL 3 are released.
      - This view does NOT mark the proposal COMPLETED (proponent "Claimed" should do that).
    """
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if proposal.proposal_status not in {
        Proposal.ProposalStatus.FOR_SUBMISSION_AND_UPLOAD,
        Proposal.ProposalStatus.APPROVED,
        Proposal.ProposalStatus.COMPLETED,
    }:
        messages.error(request, "This proposal is not yet ready for document release.")
        return redirect("dashboard_redirect")

    signed_doc = ProposalFinalDocument.objects.filter(
        proposal=proposal,
        document_type=ProposalFinalDocument.DocumentType.SIGNED_PROPOSAL,
    ).first()

    if request.method == "POST":
        now = timezone.now()

        # Optional staff verification controls in your form:
        #   signed_action = "verify" | "reject"
        #   signed_remarks = text
        signed_action = (request.POST.get("signed_action") or "").strip().lower()
        signed_remarks = (request.POST.get("signed_remarks") or "").strip()

        if signed_doc and signed_action in {"verify", "reject"}:
            signed_doc.is_verified = (signed_action == "verify")
            signed_doc.remarks = signed_remarks
            signed_doc.save(update_fields=["is_verified", "remarks"])

        loa_file = request.FILES.get("letter_of_award_file")
        endorsement_file = request.FILES.get("endorsement_file")
        agreement_file = request.FILES.get("extension_agreement_file")

        uploaded_any = bool(loa_file or endorsement_file or agreement_file)

        # If staff is already uploading the final docs, treat that as an implicit verification step
        # (so you don't *have* to add new form controls immediately).
        if signed_doc and uploaded_any and not signed_doc.is_verified and signed_action != "reject":
            signed_doc.is_verified = True
            if signed_remarks:
                signed_doc.remarks = signed_remarks
            signed_doc.save(update_fields=["is_verified", "remarks"])

        # Enforce verification before releasing final docs
        verified_signed = ProposalFinalDocument.objects.filter(
            proposal=proposal,
            document_type=ProposalFinalDocument.DocumentType.SIGNED_PROPOSAL,
            is_verified=True,
        ).first()
        if not verified_signed:
            messages.error(request, "Signed proposal is not yet VERIFIED by staff.")
            return redirect("staff_release_approval_documents", proposal_id=proposal.id)

        if loa_file:
            ProposalFinalDocument.objects.update_or_create(
                proposal=proposal,
                document_type=ProposalFinalDocument.DocumentType.LETTER_OF_AWARD,
                defaults={
                    "file": loa_file,
                    "uploaded_by": request.user,
                    "remarks": "",
                    "is_verified": True,
                },
            )
            ProposalAttachment.objects.filter(
                proposal=proposal,
                category=ProposalAttachment.Category.OTHER,
                label="Letter of Award",
            ).delete()
            ProposalAttachment.objects.create(
                proposal=proposal,
                file=loa_file,
                category=ProposalAttachment.Category.OTHER,
                label="Letter of Award",
            )
            if not proposal.letter_of_award_released_at:
                proposal.letter_of_award_released_at = now

        if endorsement_file:
            ProposalFinalDocument.objects.update_or_create(
                proposal=proposal,
                document_type=ProposalFinalDocument.DocumentType.ENDORSEMENT_FOR_APPROVAL,
                defaults={
                    "file": endorsement_file,
                    "uploaded_by": request.user,
                    "remarks": "",
                    "is_verified": True,
                },
            )
            ProposalAttachment.objects.filter(
                proposal=proposal,
                category=ProposalAttachment.Category.OTHER,
                label="Endorsement for Approval",
            ).delete()
            ProposalAttachment.objects.create(
                proposal=proposal,
                file=endorsement_file,
                category=ProposalAttachment.Category.OTHER,
                label="Endorsement for Approval",
            )
            if not proposal.endorsement_released_at:
                proposal.endorsement_released_at = now

        if agreement_file:
            ProposalFinalDocument.objects.update_or_create(
                proposal=proposal,
                document_type=ProposalFinalDocument.DocumentType.EXTENSION_AGREEMENT,
                defaults={
                    "file": agreement_file,
                    "uploaded_by": request.user,
                    "remarks": "",
                    "is_verified": True,
                },
            )
            ProposalAttachment.objects.filter(
                proposal=proposal,
                category=ProposalAttachment.Category.OTHER,
                label="Extension Agreement",
            ).delete()
            ProposalAttachment.objects.create(
                proposal=proposal,
                file=agreement_file,
                category=ProposalAttachment.Category.OTHER,
                label="Extension Agreement",
            )
            if not proposal.extension_agreement_released_at:
                proposal.extension_agreement_released_at = now

        proposal.last_saved_at = now
        proposal.save(update_fields=[
            "letter_of_award_released_at",
            "endorsement_released_at",
            "extension_agreement_released_at",
            "last_saved_at",
        ])

        # Mark APPROVED / CLAIMING only when all 3 are released
        if (
            proposal.letter_of_award_released_at
            and proposal.endorsement_released_at
            and proposal.extension_agreement_released_at
        ):
            if proposal.proposal_status != Proposal.ProposalStatus.APPROVED:
                proposal.approve_proposal()
            messages.success(request, "All approval documents released. Proposal is now APPROVED / CLAIMING.")
        else:
            messages.success(request, "Documents updated. Release the remaining documents to mark Approved / Claiming.")

        return redirect("dashboard_redirect")

    loa_doc = ProposalFinalDocument.objects.filter(
        proposal=proposal,
        document_type=ProposalFinalDocument.DocumentType.LETTER_OF_AWARD,
    ).first()
    end_doc = ProposalFinalDocument.objects.filter(
        proposal=proposal,
        document_type=ProposalFinalDocument.DocumentType.ENDORSEMENT_FOR_APPROVAL,
    ).first()
    agr_doc = ProposalFinalDocument.objects.filter(
        proposal=proposal,
        document_type=ProposalFinalDocument.DocumentType.EXTENSION_AGREEMENT,
    ).first()

    context = {
        "proposal": proposal,
        "signed_doc": signed_doc,
        "loa_doc": loa_doc,
        "end_doc": end_doc,
        "agr_doc": agr_doc,
        "letter_of_award_released_at": proposal.letter_of_award_released_at,
        "endorsement_released_at": proposal.endorsement_released_at,
        "extension_agreement_released_at": proposal.extension_agreement_released_at,
    }
    return render(request, "services/post_approval/release_documents.html", context)

def _mark_approval_documents_completed(proposal):
    """
    Advance the proposal phase when the staff-released approval documents
    have been received/downloaded and the proposal is already approved.
    """
    if proposal.proposal_status != Proposal.ProposalStatus.APPROVED:
        return False

    if not (
        proposal.letter_of_award_released_at
        and proposal.endorsement_released_at
        and proposal.extension_agreement_released_at
    ):
        return False

    proposal.mark_proposal_completed()

    if getattr(proposal, "requires_moa", False):
        if proposal.moa_status in {Proposal.MOAStatus.NOT_STARTED, Proposal.MOAStatus.NOT_REQUIRED}:
            proposal.moa_status = Proposal.MOAStatus.DRAFT

    if proposal.status == Proposal.OverallStatus.DRAFT:
        proposal.status = Proposal.OverallStatus.ACTIVE

    proposal.save(update_fields=["moa_status", "status", "last_saved_at"])
    return True


@login_required
@faculty_like_required
@require_POST
def proposal_mark_claimed(request, proposal_id):
    """
    Proponent confirms they have received the released approval documents.

    Updated workflow:
      APPROVED / CLAIMING -> Proposal Completed (proposal phase)
      and then MOA phase starts (if requires_moa).
    """
    proposal = get_object_or_404(Proposal, id=proposal_id)

    is_proponent = (
        request.user == proposal.created_by
        or proposal.proponents.filter(user=request.user).exists()
    )
    if not is_proponent:
        messages.error(request, "You do not have permission to claim this proposal.")
        return redirect("dashboard_redirect")

    if proposal.proposal_status != Proposal.ProposalStatus.APPROVED:
        messages.error(request, "This proposal is not yet in Approved / Claiming status.")
        return redirect("dashboard_redirect")

    # Enforce that all 3 approval documents were released by staff.
    if not (
        proposal.letter_of_award_released_at
        and proposal.endorsement_released_at
        and proposal.extension_agreement_released_at
    ):
        messages.error(request, "Approval documents are not yet fully released.")
        return redirect("dashboard_redirect")

    _mark_approval_documents_completed(proposal)

    messages.success(request, "Proposal claimed. MOA phase is now started.")
    return redirect("dashboard_redirect")


# ==============================
# STAFF SUMMARY ACTIONS
# ==============================

def _get_review_round_for_staff_summary(proposal):
    """
    Return the review round that STAFF should work on.

    Priority:
    1) Latest OPEN round that is already marked ready_for_staff_summary
    2) Latest OPEN round
    3) Latest round (fallback)
    """
    qs = proposal.review_rounds.order_by("-round_no")
    ready_open = qs.filter(is_closed=False, ready_for_staff_summary=True).first()
    if ready_open:
        return ready_open
    open_round = qs.filter(is_closed=False).first()
    if open_round:
        return open_round
    return qs.first()


def _get_latest_sent_summary(proposal):
    return (
        proposal.comment_summaries
        .filter(sent_to_proponent=True)
        .select_related("review_round")
        .order_by("-review_round__round_no", "-created_at")
        .first()
    )


def _get_director_name(proposal, review_round):
    director_comment = (
        ProposalSectionComment.objects.filter(
            proposal=proposal,
            review_round=review_round,
            reviewer_role="DIRECTOR",
        )
        .select_related("reviewer__profile")
        .order_by("-created_at")
        .first()
    )

    if not director_comment:
        return "Director"

    profile = getattr(director_comment.reviewer, "profile", None)
    return (getattr(profile, "full_name", "") or director_comment.reviewer.username or "Director")


def summarize_comments(comments_queryset, *, include_step_labels=True):
    """
    Convert raw reviewer comments into concrete "key revision points".

    Goals:
    - Use the *actual* comment text (not generic placeholders).
    - Group by step and extract 1–2 actionable points per step.
    - Keep output short and readable for the Summary letter.
    """
    step_title_map = {s.get("no"): s.get("title") for s in STEP_LABELS if isinstance(s, dict)}

    ACTION_WORDS = (
        "should", "must", "please", "kindly", "revise", "update", "add", "include",
        "clarify", "ensure", "provide", "remove", "align", "correct", "complete",
        "justify", "specify", "strengthen", "reword", "edit", "format",
    )

    def _normalize_ws(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip())

    def _extract_keypoints(text: str, max_points: int = 2):
        """
        Pull the most actionable sentence/bullet(s) from a comment.
        """
        t = (text or "").strip()
        if not t:
            return []

        # Prefer explicit bullets when present.
        lines = [ln.strip(" \t\r") for ln in (text or "").splitlines() if ln.strip()]
        bullet_like = []
        for ln in lines:
            if re.match(r"^(\-|\*|•|\u2022|\d+[\.\)]|[a-zA-Z][\.\)])\s+", ln):
                ln = re.sub(r"^(\-|\*|•|\u2022|\d+[\.\)]|[a-zA-Z][\.\)])\s+", "", ln).strip()
                if ln:
                    bullet_like.append(_normalize_ws(ln))

        if bullet_like:
            out = []
            seen = set()
            for b in bullet_like:
                key = b.lower()
                if key in seen:
                    continue
                seen.add(key)
                out.append(b)
                if len(out) >= max_points:
                    break
            return out

        # Otherwise, score sentences/clauses.
        compact = _normalize_ws(t)
        parts = re.split(r"(?<=[\.\?\!;:])\s+|\s+\-\s+|\s+\u2022\s+", compact)
        parts = [_normalize_ws(p) for p in parts if _normalize_ws(p)]
        if not parts:
            return []

        def score(p: str) -> float:
            lp = p.lower()
            hits = sum(1 for w in ACTION_WORDS if w in lp)
            # bias toward informative-but-not-too-long clauses
            length = len(p)
            length_score = min(length, 220) / 60.0
            return (hits * 3.0) + length_score

        ranked = sorted(parts, key=score, reverse=True)

        out = []
        seen = set()
        for p in ranked:
            if len(p) < 12:
                continue
            key = p.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append(p)
            if len(out) >= max_points:
                break

        if not out and compact:
            out = [compact[:220].rstrip() + ("..." if len(compact) > 220 else "")]

        return out

    # Build (step_no -> [points...]) preserving step order later.
    grouped = {}
    global_seen = set()

    for c in comments_queryset:
        raw = (getattr(c, "comment", "") or "").strip()
        if not raw:
            continue

        step_no = int(getattr(c, "step_no", 0) or 0)

        for kp in _extract_keypoints(raw, max_points=2):
            kp = _normalize_ws(kp)
            if not kp:
                continue
            dedupe_key = (step_no, kp.lower())
            if dedupe_key in global_seen:
                continue
            global_seen.add(dedupe_key)
            grouped.setdefault(step_no, []).append(kp)

    if not grouped:
        return ["No significant comments were provided by reviewers."]

    bullets = []

    # Sort steps numerically; keep step_no==0 ("General") last.
    for step_no in sorted(grouped.keys(), key=lambda n: (n == 0, n)):
        points = grouped.get(step_no) or []
        if not points:
            continue

        title = step_title_map.get(step_no) if step_no else "General"
        joined = "; ".join(points[:2]).strip()

        if len(joined) > 220:
            joined = joined[:217].rstrip() + "..."

        if step_no and include_step_labels:
            bullets.append(f"Step {step_no} – {title}: {joined}")
        else:
            bullets.append(joined)

        if len(bullets) >= 8:
            break

    # Safety net: if bullets are still too few, add a tiny keyword-based hint (non-generic).
    if len(bullets) < 2:
        bullets.append("Please review each step and apply all reviewer notes before resubmission.")

    return bullets



def _generate_default_summary_text(proposal, review_round, director_name):
    """
    Default for STAFF textarea:
    - ONLY the consolidated comment points
    - NO greeting letter
    - NO 'Step X - ...' labels
    This keeps summary_text compatible with Clear Summary Preview/DOCX (points list).
    """
    comments_qs = (
        ProposalSectionComment.objects
        .filter(proposal=proposal, review_round=review_round)
        .order_by("step_no", "created_at")
    )

    points = summarize_comments(comments_qs, include_step_labels=False)
    # Keep as bullet lines; preview/docx will number them cleanly.
    return "\n".join([f"- {p}" for p in points])



def _finalize_summary_and_return_for_revision(proposal, review_round):
    # Deactivate evaluator assignments for this round
    ProposalEvaluatorAssignment.objects.filter(
        proposal=proposal,
        review_round=review_round,
        is_active=True,
    ).update(is_active=False)

    # Clear staff-ready flag (a round with a SENT summary should never remain in the staff queue)
    update_fields = []
    if hasattr(review_round, "ready_for_staff_summary") and review_round.ready_for_staff_summary:
        review_round.ready_for_staff_summary = False
        update_fields.append("ready_for_staff_summary")

    # Close round
    if not getattr(review_round, "is_closed", False):
        review_round.is_closed = True
        update_fields.append("is_closed")

    if update_fields:
        review_round.save(update_fields=update_fields)

    # Return proposal for revision (unlocks)
    proposal.return_for_revision()

@login_required
@role_required(["STAFF"])
def proposal_comment_summary(request, proposal_id):
    """
    STAFF drafts (and optionally sends) the official summary for a review round
    that has been marked ready_for_staff_summary by the Director.
    """
    proposal = get_object_or_404(Proposal, id=proposal_id)
    review_round = _get_review_round_for_staff_summary(proposal)

    if not review_round or not review_round.ready_for_staff_summary:
        messages.error(request, "This review round is not yet ready for staff summary.")
        return redirect("staff_dashboard")

    step_comments = (
        ProposalSectionComment.objects.filter(proposal=proposal, review_round=review_round)
        .select_related("reviewer", "reviewer__profile", "review_round")
        .order_by("step_no", "created_at")
    )

    summary = ProposalCommentSummary.objects.filter(
        proposal=proposal,
        review_round=review_round,
    ).first()

    director_name = _get_director_name(proposal, review_round)

    if request.method == "POST":
        summary_text = (request.POST.get("summary_text") or "").strip()

        # Primary driver: action buttons
        action = (request.POST.get("action") or "").strip() or "save_draft"

        # Optional: support a POST-based DOCX download if you wire the button as a submit
        if action == "download_docx":
            effective_text = summary_text or _generate_default_summary_text(proposal, review_round, director_name)
            points = _extract_points(effective_text) or ["No significant comments were provided by reviewers."]

            data = _build_clear_summary_docx(proposal=proposal, points=points)
            filename = f"Comment_Summary_{proposal.title}.docx"
            resp = HttpResponse(
                data,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            resp["Content-Disposition"] = f'attachment; filename="{filename}"'
            return resp

        if action == "send_to_proponent":
            if not summary_text:
                messages.error(request, "Summary text cannot be empty.")
                return redirect("proposal_comment_summary", proposal_id=proposal.id)

            if not summary:
                summary = ProposalCommentSummary.objects.create(
                    proposal=proposal,
                    review_round=review_round,
                    prepared_by=request.user,
                    summary_text=summary_text,
                    sent_to_proponent=True,
                )
            else:
                summary.summary_text = summary_text
                summary.sent_to_proponent = True
                if not summary.prepared_by_id:
                    summary.prepared_by = request.user
                summary.save(update_fields=["summary_text", "sent_to_proponent", "prepared_by"])

            _finalize_summary_and_return_for_revision(proposal, review_round)
            messages.success(request, f"Summary for '{proposal.display_title}' sent to proponent.")
            return redirect("staff_dashboard")

        # save_draft (default)
        if not summary_text:
            summary_text = _generate_default_summary_text(proposal, review_round, director_name)

        if not summary:
            summary = ProposalCommentSummary.objects.create(
                proposal=proposal,
                review_round=review_round,
                prepared_by=request.user,
                summary_text=summary_text,
                sent_to_proponent=False,
            )
        else:
            summary.summary_text = summary_text
            if not summary.prepared_by_id:
                summary.prepared_by = request.user
            # don't flip a sent summary back to draft silently
            if not summary.sent_to_proponent:
                summary.sent_to_proponent = False
            summary.save(update_fields=["summary_text", "prepared_by", "sent_to_proponent"])

        messages.success(request, "Summary saved.")
        return redirect("proposal_comment_summary", proposal_id=proposal.id)

    default_text = ((summary.summary_text or "").strip() if summary else "") or _generate_default_summary_text(
        proposal, review_round, director_name
    )

    context = {
        "proposal": proposal,
        "review_round": review_round,
        "step_comments": step_comments,
        "summary_text": default_text,
        "is_draft": (summary is None) or (not summary.sent_to_proponent),
        "summary": summary,
        "director_name": director_name,
    }
    return render(request, "services/review/proposal_comment_summary.html", context)



def proposal_send_summary(request, proposal_id):
    """
    Convenience endpoint: send an already-prepared DRAFT summary to the proponent.
    """
    proposal = get_object_or_404(Proposal, id=proposal_id)
    review_round = _get_review_round_for_staff_summary(proposal)

    if not review_round or not review_round.ready_for_staff_summary:
        messages.error(request, "This review round is not yet ready for staff summary.")
        return redirect("staff_dashboard")

    summary = ProposalCommentSummary.objects.filter(
        proposal=proposal,
        review_round=review_round,
        sent_to_proponent=False,
    ).first()

    if not summary:
        messages.error(request, "No draft summary found. Please prepare one first.")
        return redirect("proposal_comment_summary", proposal_id=proposal.id)

    summary.sent_to_proponent = True
    summary.save(update_fields=["sent_to_proponent"])

    _finalize_summary_and_return_for_revision(proposal, review_round)

    messages.success(request, f"Summary for '{proposal.display_title}' sent to proponent.")
    return redirect("staff_dashboard")


def _can_view_summary(user, proposal):
    if not user.is_authenticated:
        return False
    if _is_staff(user):
        return True
    if user == proposal.created_by:
        return True
    if proposal.proponents.filter(user=user).exists():
        return True
    return False


@login_required
def proposal_version_summary(request, proposal_id, round_no):
    """Open a specific proposal version in the existing readonly proposal view."""
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_view_summary(request.user, proposal):
        messages.error(request, "You do not have permission to view this proposal version.")
        return redirect("dashboard_redirect")

    if not proposal.review_rounds.filter(round_no=round_no).exists():
        messages.error(request, "That review version does not exist.")
        return redirect("proposal_storage", proposal_id=proposal.id)

    return redirect(f"{reverse('proposal_wizard', args=[proposal.id, 1])}?readonly=1")


@login_required
def proposal_view_summary(request, proposal_id):
    """Proponent view to read the latest issued summary and visible comments."""
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_view_summary(request.user, proposal):
        messages.error(request, "You do not have permission to view this summary.")
        return redirect("dashboard_redirect")

    summary = _get_latest_sent_summary(proposal)
    if not summary:
        messages.error(request, "No summary has been issued yet.")
        return redirect("dashboard_redirect")

    review_round = summary.review_round

    comments_qs = ProposalSectionComment.objects.filter(
        proposal=proposal,
        review_round=review_round,
    ).select_related("reviewer__profile").order_by("step_no", "created_at")

    if not _is_staff(request.user):
        comments_qs = comments_qs.filter(is_visible_to_proponent=True)

    context = {
        "proposal": proposal,
        "summary": summary,
        "comments": comments_qs,
        "review_round": review_round,
        "director_name": _get_director_name(proposal, review_round),
    }
    return render(request, "proposals/view_summary.html", context)


@login_required
def proposal_print_summary_page(request, proposal_id):
    """HTML print-friendly summary view (staff/proponent)."""
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_view_summary(request.user, proposal):
        messages.error(request, "You do not have permission to view this summary.")
        return redirect("dashboard_redirect")

    # Prefer latest sent summary; STAFF can still print the current draft.
    summary = _get_latest_sent_summary(proposal)
    if not summary and _is_staff(request.user):
        review_round = _get_review_round_for_staff_summary(proposal)
        summary = ProposalCommentSummary.objects.filter(
            proposal=proposal, review_round=review_round
        ).order_by("-created_at").first()
    else:
        review_round = summary.review_round if summary else None

    if not summary or not review_round:
        messages.error(request, "No summary available for printing.")
        return redirect("dashboard_redirect")

    comments_qs = ProposalSectionComment.objects.filter(
        proposal=proposal,
        review_round=review_round,
    ).select_related("reviewer__profile").order_by("step_no", "created_at")

    if not _is_staff(request.user):
        comments_qs = comments_qs.filter(is_visible_to_proponent=True)

    context = {
        "proposal": proposal,
        "summary": summary,
        "comments": comments_qs,
        "director_name": _get_director_name(proposal, review_round),
        "review_round": review_round,
    }
    return render(request, "proposals/print_summary.html", context)


@login_required
def proposal_print_summary(request, proposal_id):
    """Generate PDF summary using xhtml2pdf."""
    proposal = get_object_or_404(Proposal, id=proposal_id)

    if not _can_view_summary(request.user, proposal):
        messages.error(request, "You do not have permission to view this summary.")
        return redirect("dashboard_redirect")

    # Prefer latest sent summary; STAFF can still export the current draft.
    summary = _get_latest_sent_summary(proposal)
    if not summary and _is_staff(request.user):
        review_round = _get_review_round_for_staff_summary(proposal)
        summary = ProposalCommentSummary.objects.filter(
            proposal=proposal, review_round=review_round
        ).order_by("-created_at").first()
    else:
        review_round = summary.review_round if summary else None

    if not summary or not review_round:
        messages.error(request, "No summary available.")
        return redirect("dashboard_redirect")

    comments_qs = ProposalSectionComment.objects.filter(
        proposal=proposal,
        review_round=review_round,
    ).select_related("reviewer__profile").order_by("step_no", "created_at")

    if not _is_staff(request.user):
        comments_qs = comments_qs.filter(is_visible_to_proponent=True)

    html_string = render_to_string(
        "proposals/summary_pdf.html",
        {
            "proposal": proposal,
            "summary": summary,
            "comments": comments_qs,
            "director_name": _get_director_name(proposal, review_round),
            "review_round": review_round,
        },
    )

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f'inline; filename="summary_{proposal.id}.pdf"'

    pisa_status = pisa.CreatePDF(html_string, dest=response)
    if pisa_status.err:
        return HttpResponse("Error generating PDF", status=500)

    return response

# ==========================================
# STAFF: VIEW + PRINT CLEAR SUMMARY (DOCX)
# ==========================================
# Template docx: proposals/template_files/clear_summary_template.docx
# Based on your "Clear Summary.docx" structure.
# Output matches the layout shown in your "Sample Summary.docx".

import re
from io import BytesIO
from pathlib import Path

from django.http import HttpResponse
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, render

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from accounts.decorators import role_required
from accounts.models import Signatory

from .models import Proposal, ProposalCommentSummary, ProposalSectionComment


def _extract_last_name(full_name: str) -> str:
    s = (full_name or "").strip()
    if not s:
        return ""

    # Remove common prefixes and punctuation
    s = re.sub(r"^(DR\.|MR\.|MS\.|MRS\.|ENGR\.|ATTY\.)\s+", "", s, flags=re.IGNORECASE).strip()
    s = re.sub(r"[,\.\(\)]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()

    parts = s.split(" ")
    if not parts:
        return ""

    # Last "word" is usually the last name
    return parts[-1].title()


def _get_signatory(position_title: str, *, campus: str = "", college: str = "", department: str = ""):
    campus = (campus or "").strip()
    college = (college or "").strip()
    department = (department or "").strip()

    return Signatory.objects.filter(
        position_title=position_title,
        campus=campus,
        college=college,
        department=department,
    ).first()


def _proponent_line(proposal: Proposal) -> str:
    props = list(proposal.proponents.all().order_by("id"))
    names = [(p.full_name or "").strip() for p in props]
    names = [n for n in names if n]

    if names:
        if len(names) == 1:
            return names[0]
        return f"{names[0]}, et.al."

    prof = getattr(proposal.created_by, "profile", None)
    return getattr(prof, "full_name", "") or proposal.created_by.get_username()


def _extract_points(summary_text: str):
    """
    Converts summary_text into clean numbered points.

    - Accepts lines like:
        Step 1 – Extension Type and Scope: Sample
        - Step 2 - Title: Sample
        1. Step 3 – Proponents: Sample; Comment
        • Sample
    - Removes bullet/number prefixes
    - Removes "Step X – <title>:" prefix entirely, leaving only the comment text
    """
    points = []
    for line in (summary_text or "").splitlines():
        s = (line or "").strip()
        if not s:
            continue

        # remove bullet / numbering prefixes
        s = re.sub(r"^[-•\*\u2022]\s*", "", s).strip()
        s = re.sub(r"^\d+[\.\)]\s*", "", s).strip()

        # remove "Step X – Title:" prefix (keep only the actual comment)
        s = re.sub(r"^Step\s+\d+\s*[-–—]\s*[^:]{0,200}:\s*", "", s, flags=re.I).strip()
        # fallback: "Step X: ..." or "Step X - ..."
        s = re.sub(r"^Step\s+\d+\s*[:\-–—]\s*", "", s, flags=re.I).strip()

        if s:
            points.append(s)

    # if staff wrote a single paragraph without line breaks, keep it as one point
    if not points and (summary_text or "").strip():
        points = [(summary_text or "").strip()]

    return points


def _apply_numbering(paragraph, *, num_id: int = 1, ilvl: int = 0):
    # Apply Word numbering using an existing numId in the template doc.
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)

    ilvl_el = numPr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))

    numId_el = numPr.find(qn("w:numId"))
    if numId_el is None:
        numId_el = OxmlElement("w:numId")
        numPr.append(numId_el)
    numId_el.set(qn("w:val"), str(num_id))


def _insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def _build_clear_summary_docx(*, proposal: Proposal, points):
    template_path = Path(__file__).resolve().parent / "template_files" / "clear_summary_template.docx"
    doc = Document(str(template_path))
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    title = proposal.title or proposal.research_title or "Untitled Proposal"
    campus = (proposal.campus or "").strip()

    campus_director = _get_signatory(Signatory.Position.CAMPUS_DIRECTOR, campus=campus) if campus else None
    campus_coordinator = _get_signatory(Signatory.Position.CAMPUS_EXTENSION_COORDINATOR, campus=campus) if campus else None
    director_ext = _get_signatory(Signatory.Position.DIRECTOR_EXTENSION)

    # Full names (for last name extraction)
    cd_full = (campus_director.full_name if campus_director else "").strip() or "[NAME OF THE CAMPUS DIRECTOR]"
    cc_full = (campus_coordinator.full_name if campus_coordinator else "").strip() or "[NAME OF CAMPUS COORDINATOR]"

    # Display names (include credentials)
    cd_cred = (campus_director.credentials if campus_director else "").strip()
    cc_cred = (campus_coordinator.credentials if campus_coordinator else "").strip()
    cd_display = cd_full + (f", {cd_cred}" if cd_cred else "")
    cc_display = cc_full + (f", {cc_cred}" if cc_cred else "")

    director_name = (director_ext.full_name if director_ext else "").strip() or "[NAME OF THE DIRECTOR]"
    director_cred = (director_ext.credentials if director_ext else "").strip()
    proponent_line = _proponent_line(proposal)
    date_str = timezone.localdate().strftime("%B %d, %Y")

    # Placeholder → value
    repl = {
        "[TITLE]": title,
        "[Date]": date_str,
        "[Campus]": campus or "[Campus]",
        "[NAME OF THE CAMPUS DIRECTOR]": cd_display,
        "[NAME OF CAMPUS COORDINATOR]": cc_display,
        "[NAME OF PROPONENT]": proponent_line,
        "[LAST NAME OF THE CAMPUS DIRECTOR]": _extract_last_name(cd_full) or "[LAST NAME OF THE CAMPUS DIRECTOR]",
        "[NAME OF THE DIRECTOR]": director_name,
        "[Credentials]": director_cred,
    }

    def replace_in_paragraph_runs(paragraph):
        for run in paragraph.runs:
            if not run.text:
                continue
            for k, v in repl.items():
                if k in run.text:
                    run.text = run.text.replace(k, v)

        # cleanup if credentials missing -> remove dangling ", "
        if not director_cred and "[NAME OF THE DIRECTOR]" not in paragraph.text:
            if paragraph.text.strip().endswith(","):
                # remove trailing comma by trimming last run that contains it
                for run in reversed(paragraph.runs):
                    if run.text and run.text.strip().endswith(","):
                        run.text = run.text.rstrip().rstrip(",")
                        break

    # Replace in body paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph_runs(p)

    # Replace in tables too (safe)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph_runs(p)

    # Insert comments where [Comments] is
    comments_para = None
    for p in doc.paragraphs:
        if "[Comments]" in p.text:
            comments_para = p
            break

    if comments_para is not None:
        if not points:
            points = ["No significant comments were provided by reviewers."]

        def _apply_manual_list_format(paragraph):
            pf = paragraph.paragraph_format
            pf.left_indent = Pt(68)   # indent the number to the right
            pf.first_line_indent = Pt(0)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            try:
                pf.tab_stops.add_tab_stop(Pt(80))  # align comment text
            except Exception:
                pass

        # Replace [Comments] run text WITHOUT resetting paragraph.runs
        first_text = f"1.	{points[0]}"
        replaced = False
        for run in comments_para.runs:
            if "[Comments]" in (run.text or ""):
                run.text = run.text.replace("[Comments]", first_text)
                replaced = True

        # If the placeholder was split oddly, fallback to setting the whole paragraph text
        if not replaced:
            comments_para.text = first_text

        _apply_manual_list_format(comments_para)

        # Force Arial 12 in list item runs (numbers + text)
        for run in comments_para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)

        last = comments_para
        for idx, pt in enumerate(points[1:], start=2):
            newp = _insert_paragraph_after(last, text=f"{idx}.	{pt}", style=comments_para.style)
            _apply_manual_list_format(newp)

            for run in newp.runs:
                run.font.name = "Arial"
                run.font.size = Pt(12)

            last = newp


    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


@login_required
@role_required(["STAFF"])
def staff_comment_summary_preview(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)

    # ✅ Use the same round as the staff summary editor
    review_round = _get_review_round_for_staff_summary(proposal)

    summary = (
        ProposalCommentSummary.objects.filter(proposal=proposal, review_round=review_round).first()
        if review_round else None
    )

    comments_qs = (
        ProposalSectionComment.objects.filter(proposal=proposal, review_round=review_round)
        .order_by("step_no", "created_at")
        if review_round else ProposalSectionComment.objects.none()
    )

    # ✅ Prefer saved summary; otherwise generate from comments (no Step labels)
    points = _extract_points(summary.summary_text) if (summary and summary.summary_text) else []
    if not points:
        try:
            points = summarize_comments(comments_qs, include_step_labels=False)
        except Exception:
            points = [c.comment.strip() for c in comments_qs[:12] if (c.comment or "").strip()]

    campus = (proposal.campus or "").strip()
    campus_director = _get_signatory(Signatory.Position.CAMPUS_DIRECTOR, campus=campus) if campus else None
    campus_coordinator = _get_signatory(Signatory.Position.CAMPUS_EXTENSION_COORDINATOR, campus=campus) if campus else None
    director_ext = _get_signatory(Signatory.Position.DIRECTOR_EXTENSION)

    # Full names (for last name extraction)
    campus_director_name = (campus_director.full_name if campus_director else "").strip()
    campus_coordinator_name = (campus_coordinator.full_name if campus_coordinator else "").strip()

    # ✅ Display names with credentials
    campus_director_cred = (campus_director.credentials if campus_director else "").strip()
    campus_coordinator_cred = (campus_coordinator.credentials if campus_coordinator else "").strip()

    campus_director_display = (campus_director_name + (f", {campus_director_cred}" if campus_director_cred else "")).strip() \
        or "[NAME OF THE CAMPUS DIRECTOR]"
    campus_coordinator_display = (campus_coordinator_name + (f", {campus_coordinator_cred}" if campus_coordinator_cred else "")).strip() \
        or "[NAME OF CAMPUS COORDINATOR]"

    director_name = (director_ext.full_name if director_ext else "").strip()
    director_cred = (director_ext.credentials if director_ext else "").strip()
    director_signature = (director_name + (f", {director_cred}" if director_cred else "")).strip() \
        or "[NAME OF THE DIRECTOR], [Credentials]"

    ctx = {
        "proposal": proposal,
        "review_round": review_round,
        "title": proposal.title or proposal.research_title or "Untitled Proposal",
        "date_str": timezone.localdate().strftime("%B %d, %Y"),
        "campus": campus or "[Campus]",
        "campus_director_name": campus_director_name,              # keep for last name extraction
        "campus_director_display": campus_director_display,        # ✅ show with credentials
        "campus_coordinator_display": campus_coordinator_display,  # ✅ show with credentials
        "proponent_line": _proponent_line(proposal),
        "dear_last_name": _extract_last_name(campus_director_name) or "[LAST NAME OF THE CAMPUS DIRECTOR]",
        "points": points,
        "director_signature": director_signature,                  # ✅ already includes credentials
    }
    return render(request, "services/review/proposal_comment_summary_preview.html", ctx)


@login_required
@role_required(["STAFF"])
def staff_comment_summary_docx(request, proposal_id):
    proposal = get_object_or_404(Proposal, id=proposal_id)
    review_round = _get_review_round_for_staff_summary(proposal)

    summary = (
        ProposalCommentSummary.objects.filter(proposal=proposal, review_round=review_round).first()
        if review_round else None
    )
    comments_qs = (
        ProposalSectionComment.objects.filter(
            proposal=proposal,
            review_round=review_round,
        ).order_by("step_no", "created_at")
        if review_round else ProposalSectionComment.objects.none()
    )

    # Prefer staff-written summary; otherwise generate from comments (NO "Step X ...")
    points = _extract_points(summary.summary_text) if (summary and summary.summary_text) else []
    if not points:
        try:
            points = summarize_comments(comments_qs, include_step_labels=False)
        except Exception:
            points = [c.comment.strip() for c in comments_qs[:8] if (c.comment or "").strip()]

    data = _build_clear_summary_docx(proposal=proposal, points=points)

    filename = f"Comment_Summary_{proposal.title}.docx"
    resp = HttpResponse(
        data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


# Paste this replacement into proposals/views.py
# It fixes the VariableDoesNotExist error by ensuring every file row has
# filename/file_url/label keys before the template tries to render them.

@login_required
@login_required
def proposal_moa_draft(request, proposal_id):
    """
    Guided MOA draft page for proponents and staff.
    It generates a structured draft from the proposal details and allows
    the final signed MOA to be uploaded into the system.
    """
    proposal = get_object_or_404(Proposal, id=proposal_id)

    is_proponent = (
        request.user == proposal.created_by
        or proposal.proponents.filter(user=request.user).exists()
    )
    if not (is_proponent or _is_staff(request.user)):
        messages.error(request, "You do not have permission to manage MOA for this proposal.")
        return redirect("dashboard_redirect")

    if not getattr(proposal, "requires_moa", False):
        messages.info(request, "This proposal does not require a MOA.")
        return redirect("proposal_storage", proposal_id=proposal.id)

    moa_doc = proposal.final_documents.filter(document_type=ProposalFinalDocument.DocumentType.MOA).first()

    default_fields = {
        "parties": (proposal.created_by.profile.full_name if getattr(proposal.created_by, "profile", None) else getattr(proposal.created_by, "username", "")) or "Proponent Institution",
        "purpose": f"This MOA establishes the collaboration and implementation arrangement for the extension project titled '{proposal.title or proposal.research_title or 'Proposal'}'.",
        "scope": "The parties shall implement the approved extension activities, deliverables, and monitoring arrangements as reflected in the approved proposal.",
        "responsibilities": "The proponent shall coordinate implementation and reporting; the partner institution shall provide support, coordination, and oversight as needed.",
        "timeline": "The term of this MOA shall cover the approved implementation period stated in the proposal and any approved extension thereof.",
        "signatories": "Prepared by the proponent and approved by the authorized signatory of the partner institution.",
    }

    if request.method == "POST":
        action = (request.POST.get("action") or "draft").strip().lower()

        if action == "generate":
            draft_text = "\n\n".join([
                f"MEMORANDUM OF AGREEMENT (MOA)\n",
                f"Project: {proposal.title or proposal.research_title or 'Proposal'}\n",
                f"Prepared for: {default_fields['parties']}\n\n",
                "1. Parties\n" + (request.POST.get("parties") or default_fields["parties"]) + "\n\n",
                "2. Purpose\n" + (request.POST.get("purpose") or default_fields["purpose"]) + "\n\n",
                "3. Scope of Work\n" + (request.POST.get("scope") or default_fields["scope"]) + "\n\n",
                "4. Responsibilities\n" + (request.POST.get("responsibilities") or default_fields["responsibilities"]) + "\n\n",
                "5. Timeline\n" + (request.POST.get("timeline") or default_fields["timeline"]) + "\n\n",
                "6. Signatories\n" + (request.POST.get("signatories") or default_fields["signatories"]) + "\n",
            ])
            response = HttpResponse(draft_text, content_type="text/plain; charset=utf-8")
            filename = f"{(proposal.title or proposal.research_title or 'Proposal').replace('/', '-').replace('\\', '-')}_MOA_Draft.txt"
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            return response

        moa_file = request.FILES.get("moa_file")
        if not moa_file:
            messages.error(request, "Please upload the signed MOA file to save it in the system.")
            return redirect("proposal_moa_draft", proposal_id=proposal.id)

        ProposalFinalDocument.objects.update_or_create(
            proposal=proposal,
            document_type=ProposalFinalDocument.DocumentType.MOA,
            defaults={
                "file": moa_file,
                "uploaded_by": request.user,
                "remarks": (request.POST.get("remarks") or "").strip(),
                "is_verified": False,
            },
        )

        proposal.moa_status = Proposal.MOAStatus.DRAFT
        proposal.last_saved_at = timezone.now()
        proposal.save(update_fields=["moa_status", "last_saved_at"])

        messages.success(request, "MOA uploaded successfully. You can now review and track it in the storage area.")
        return redirect("proposal_storage", proposal_id=proposal.id)

    context = {
        "proposal": proposal,
        "moa_doc": moa_doc,
        "default_fields": default_fields,
    }
    return render(request, "services/moa/moa_draft.html", context)


@login_required
def proposal_storage(request, proposal_id):
    proposal = get_object_or_404(
        Proposal.objects
        .select_related("created_by", "created_by__profile")
        .prefetch_related(
            "proponents",
            "proponents__user",
            "proponents__user__profile",
            "attachments",
            "final_documents",
        ),
        id=proposal_id,
    )

    profile = getattr(request.user, "profile", None)
    role = (getattr(profile, "role", "") or "").upper()
    is_staff = bool(getattr(request.user, "is_staff", False) or role in {"STAFF", "DIRECTOR", "ADMIN"})

    is_proponent = (
        request.user.id == proposal.created_by_id
        or proposal.proponents.filter(user=request.user).exists()
    )

    if not (is_staff or is_proponent):
        messages.error(request, "You do not have permission to view this proposal storage.")
        return redirect("dashboard_redirect")

    def field_url(field_name: str):
        f = getattr(proposal, field_name, None)
        if not f:
            return None
        try:
            return f.url
        except Exception:
            return None

    def file_name_from_path(url_or_name: str | None) -> str:
        if not url_or_name:
            return ""
        return Path(str(url_or_name)).name

    signed_doc = proposal.final_documents.filter(
        document_type=ProposalFinalDocument.DocumentType.SIGNED_PROPOSAL
    ).first()

    loa_doc = proposal.final_documents.filter(
        document_type=ProposalFinalDocument.DocumentType.LETTER_OF_AWARD
    ).first()

    end_doc = proposal.final_documents.filter(
        document_type=ProposalFinalDocument.DocumentType.ENDORSEMENT_FOR_APPROVAL
    ).first()

    agr_doc = proposal.final_documents.filter(
        document_type=ProposalFinalDocument.DocumentType.EXTENSION_AGREEMENT
    ).first()

    attachment_rows = [
        {
            "label": "Work Plan",
            "filename": file_name_from_path(field_url("work_plan_file")),
            "file_url": field_url("work_plan_file"),
        },
        {
            "label": "Gantt Chart",
            "filename": file_name_from_path(field_url("gantt_chart_file")),
            "file_url": field_url("gantt_chart_file"),
        },
        {
            "label": "Line-item Budget",
            "filename": file_name_from_path(field_url("funding_file")),
            "file_url": field_url("funding_file"),
        },
        {
            "label": "Abstract (optional)",
            "filename": file_name_from_path(field_url("research_abstract_file")),
            "file_url": field_url("research_abstract_file"),
        },
        {
            "label": "Certificate of Completion (optional)",
            "filename": file_name_from_path(field_url("certificate_of_completion_file")),
            "file_url": field_url("certificate_of_completion_file"),
        },
    ]

    approval_rows = [
        {
            "label": "Letter of Award",
            "filename": file_name_from_path(getattr(getattr(loa_doc, "file", None), "name", None)),
            "file_url": getattr(getattr(loa_doc, "file", None), "url", None) if loa_doc else None,
            "download_url": reverse("proposal_download_approval_document", args=[proposal.id, "letter_of_award"]),
        },
        {
            "label": "Endorsement for Approval",
            "filename": file_name_from_path(getattr(getattr(end_doc, "file", None), "name", None)),
            "file_url": getattr(getattr(end_doc, "file", None), "url", None) if end_doc else None,
            "download_url": reverse("proposal_download_approval_document", args=[proposal.id, "endorsement_for_approval"]),
        },
        {
            "label": "Extension Agreement",
            "filename": file_name_from_path(getattr(getattr(agr_doc, "file", None), "name", None)),
            "file_url": getattr(getattr(agr_doc, "file", None), "url", None) if agr_doc else None,
            "download_url": reverse("proposal_download_approval_document", args=[proposal.id, "extension_agreement"]),
        },
    ]

    other_files = []
    for a in proposal.attachments.all().order_by("id"):
        file_url = None
        filename = ""
        try:
            file_url = a.file.url
            filename = Path(a.file.name).name
        except Exception:
            pass
        if file_url:
            other_files.append({
                "label": a.label or a.get_category_display() or "Attachment",
                "filename": filename,
                "file_url": file_url,
            })

    review_history = []
    for round_obj in proposal.review_rounds.order_by("-round_no"):
        summary = ProposalCommentSummary.objects.filter(
            proposal=proposal,
            review_round=round_obj,
        ).order_by("-created_at").first()

        review_history.append({
            "round_no": round_obj.round_no,
            "created_at": round_obj.created_at,
            "is_closed": round_obj.is_closed,
            "ready_for_staff_summary": round_obj.ready_for_staff_summary,
            "department_review_done": round_obj.department_review_done,
            "campus_review_done": round_obj.campus_review_done,
            "director_review_done": round_obj.director_review_done,
            "evaluator_review_done": round_obj.evaluator_review_done,
            "view_url": reverse("proposal_version_summary", args=[proposal.id, round_obj.round_no]),
        })

    uploaded_files = [
        {
            "label": "Extension Proposal (Approved DOCX)",
            "filename": f"{(proposal.title or proposal.research_title or 'Proposal').strip()}.docx",
            "file_url": reverse("proposal_download_approved_docx", args=[proposal.id]),
        }
    ]

    if signed_doc and getattr(signed_doc, "file", None):
        uploaded_files.append({
            "label": "Signed Proposal",
            "filename": Path(signed_doc.file.name).name,
            "file_url": signed_doc.file.url,
        })

    uploaded_files.extend([
        *attachment_rows,
        *approval_rows,
        *other_files,
    ])

    can_claim = (
        proposal.proposal_status == Proposal.ProposalStatus.APPROVED
        and bool(loa_doc and getattr(loa_doc, "file", None))
        and bool(end_doc and getattr(end_doc, "file", None))
        and bool(agr_doc and getattr(agr_doc, "file", None))
        and is_proponent
    )

    context = {
        "proposal": proposal,
        "signed_doc": signed_doc,
        "moa_doc": proposal.final_documents.filter(document_type=ProposalFinalDocument.DocumentType.MOA).first(),
        "approval_rows": approval_rows,
        "attachment_rows": attachment_rows,
        "other_files": other_files,
        "review_history": review_history,
        "uploaded_files": uploaded_files,
        "can_claim": can_claim,
    }
    return render(request, "services/post_approval/proposal_storage.html", context)


