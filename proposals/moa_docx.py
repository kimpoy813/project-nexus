"""
proposals/moa_docx.py

Generates a Memorandum of Agreement (MOA) .docx from the structured data
submitted through the guided "MOA Drafting" form.

The section order, headings, and signature/acknowledgment blocks follow the
sample MOA format used by ISPSC (Ilocos Sur Polytechnic State College) and
its partner institutions:

    MEMORANDUM OF AGREEMENT
    <intro "entered into by and between">
    <Party A block> -and- <Party B block>
    -WITNESSETH, that:
    WHEREAS clauses
    NOW, THEREFORE ... agree as follows:
    OBJECTIVES
    OBLIGATIONS OF THE PARTIES (A. ISPSC / B. Partner)
    INTELLECTUAL PROPERTY
    TERM AND EFFECTIVITY
    FUNDING
    DATA PRIVACY AND CONFIDENTIALITY
    AMENDMENTS
    TERMINATION
    MISCELLANEOUS
    Signature blocks (ISPSC | Partner) + "Signed in the presence of:" witnesses
    ACKNOWLEDGEMENT (notarial jurat page)

This is built from scratch with python-docx (no template file is required),
since the MOA is a legal instrument generated fresh for every proposal
rather than filled into a fixed institutional form.
"""
from __future__ import annotations

from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

try:
    from accounts.models import Signatory  # type: ignore
except Exception:  # pragma: no cover
    Signatory = None  # type: ignore

try:
    # Reuse the same best-fit Signatory lookup already used for other
    # generated documents, so the ISPSC signature block is consistent
    # app-wide.
    from .docx_forms import _signatory_lookup, _display_signatory
except Exception:  # pragma: no cover
    _signatory_lookup = None  # type: ignore
    _display_signatory = None  # type: ignore


ISPSC_FULL_NAME = "ILOCOS SUR POLYTECHNIC STATE COLLEGE (ISPSC)"
ISPSC_DESCRIPTION = (
    "a state college duly organized and existing under Philippine laws, "
    "with Main-campus and principal address at San Nicolas, Candon City, "
    "Ilocos Sur"
)
ISPSC_SHORT_NAME = "ISPSC"

DEFAULT_DATA_PRIVACY_TEXT = (
    "Both Parties commit to full compliance with R.A. 10173 (Data Privacy "
    "Act of 2012) and related laws to safeguard all personal and sensitive "
    "information collected and processed under this Agreement. All data "
    "shall be used solely for the purposes stated in this Agreement."
)
DEFAULT_AMENDMENTS_TEXT = (
    "Any modification to this MOA shall be made in writing and signed by "
    "duly authorized representatives of both Parties."
)
DEFAULT_MISC_TEXT = (
    "Nothing in this Agreement shall be construed as creating a joint "
    "venture, partnership, or employer-employee relationship between the "
    "Parties."
)


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # Ensure east-asian font attr doesn't override on some Word versions
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def _add_heading(doc: Document, text: str, *, size: int = 12, space_before: int = 18, space_after: int = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def _add_body(doc: Document, text: str, *, bold_lead: Optional[str] = None, indent: bool = True, justify: bool = True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_lead:
        lead_run = p.add_run(bold_lead)
        lead_run.bold = True
        lead_run.font.name = "Times New Roman"
        lead_run.font.size = Pt(12)

    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def _add_section_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def _add_numbered_paragraphs(doc: Document, items: list[str], *, start: int = 1):
    for i, item in enumerate(items, start=start):
        item = (item or "").strip()
        if not item:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{i}. {item}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _lines(text: str) -> list[str]:
    if not text:
        return []
    return [ln.strip() for ln in text.replace("\r\n", "\n").split("\n") if ln.strip()]


def _resolve_ispsc_signatory(proposal, position_code: str) -> str:
    if not (_signatory_lookup and _display_signatory and Signatory is not None):
        return ""
    campus = (getattr(proposal, "campus", "") or "").strip()
    college = (getattr(proposal, "college", "") or "").strip()
    department = (getattr(proposal, "department", "") or "").strip()
    sig = _signatory_lookup(position_code, campus=campus, college=college, department=department)
    return _display_signatory(sig) if sig else ""


def _signature_block_table(doc: Document, left_title: str, left_name: str, left_role: str,
                            right_title: str, right_name: str, right_role: str):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
    tbl_pr.append(tbl_w)

    usable_width = Inches(6.5)
    col_width = Inches(3.25)
    table.columns[0].width = col_width
    table.columns[1].width = col_width
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    left_cell, right_cell = table.rows[0].cells

    def fill(cell, title, name, role):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(4)
        r0 = p0.add_run(title)
        r0.bold = True
        r0.font.name = "Times New Roman"
        r0.font.size = Pt(11)

        p_by = cell.add_paragraph()
        p_by.paragraph_format.space_before = Pt(24)
        p_by.paragraph_format.space_after = Pt(2)
        r_by = p_by.add_run("By:")
        r_by.font.name = "Times New Roman"
        r_by.font.size = Pt(12)

        p_name = cell.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(18)
        p_name.paragraph_format.space_after = Pt(0)
        r_name = p_name.add_run((name.upper() if name else "_______________________"))
        r_name.bold = True
        r_name.underline = bool(name)
        r_name.font.name = "Times New Roman"
        r_name.font.size = Pt(12)

        p_role = cell.add_paragraph()
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_role.paragraph_format.space_after = Pt(0)
        r_role = p_role.add_run(role or "")
        r_role.font.name = "Times New Roman"
        r_role.font.size = Pt(11)

    fill(left_cell, left_title, left_name, left_role)
    fill(right_cell, right_title, right_name, right_role)
    return table


def _witness_table(doc: Document, witnesses: list[tuple[str, str]]):
    """witnesses: list of (name, role) tuples, rendered two per row."""
    witnesses = [w for w in witnesses if w and (w[0] or w[1])]
    if not witnesses:
        return

    rows = (len(witnesses) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    for idx, (name, role) in enumerate(witnesses):
        cell = table.rows[idx // 2].cells[idx % 2]
        p_name = cell.paragraphs[0]
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_name = p_name.add_run((name or "_______________________").upper() if name else "_______________________")
        r_name.bold = True
        r_name.font.name = "Times New Roman"
        r_name.font.size = Pt(11)

        p_role = cell.add_paragraph()
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_role = p_role.add_run(role or "")
        r_role.font.name = "Times New Roman"
        r_role.font.size = Pt(10)


def build_moa_document(proposal, data: dict) -> BytesIO:
    """
    Build a MOA .docx for the given proposal using the guided-form `data`
    dict (see proposal_moa_draft view for the expected keys) and return it
    as an in-memory BytesIO ready to be saved/attached/downloaded.
    """
    doc = Document()
    _set_default_style(doc)

    project_title = (proposal.title or proposal.research_title or "the approved extension proposal").strip()

    partner_name = (data.get("partner_name") or proposal.implementing_agency or "PARTNER INSTITUTION").strip().upper()
    partner_description = (data.get("partner_description") or "").strip()
    partner_address = (data.get("partner_address") or "").strip()
    partner_short_name = (data.get("partner_short_name") or partner_name).strip().upper()
    partner_rep_name = (data.get("partner_rep_name") or "").strip()
    partner_rep_title = (data.get("partner_rep_title") or "").strip()

    # ---------- Title ----------
    _add_heading(doc, "MEMORANDUM OF AGREEMENT", size=14, space_before=0)

    _add_body(
        doc,
        "This Memorandum of Agreement (MOA) is entered into by and between:",
        indent=False,
    )

    # ---------- Party A: ISPSC ----------
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(ISPSC_FULL_NAME + ", ")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = p.add_run(
        f"{ISPSC_DESCRIPTION}, represented herein by its President, "
    )
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    ispsc_president = _resolve_ispsc_signatory(proposal, "SUC_PRESIDENT_III") or "___________________________"
    r3 = p.add_run(f"{ispsc_president}, ")
    r3.bold = True
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)
    r4 = p.add_run(f'hereinafter referred to as the "{ISPSC_SHORT_NAME}";')
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(12)

    p_and = doc.add_paragraph()
    p_and.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_and.paragraph_format.space_after = Pt(10)
    r_and = p_and.add_run("-and-")
    r_and.italic = True
    r_and.font.name = "Times New Roman"
    r_and.font.size = Pt(12)

    # ---------- Party B: Partner ----------
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.4)
    p2.paragraph_format.space_after = Pt(10)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r5 = p2.add_run(partner_name + ", ")
    r5.bold = True
    r5.font.name = "Times New Roman"
    r5.font.size = Pt(12)
    desc_bits = []
    if partner_description:
        desc_bits.append(partner_description)
    if partner_address:
        desc_bits.append(f"with office/business address at {partner_address}")
    desc_text = ", ".join(desc_bits)
    if desc_text:
        r6 = p2.add_run(desc_text + ", ")
        r6.font.name = "Times New Roman"
        r6.font.size = Pt(12)
    if partner_rep_name:
        r7 = p2.add_run("represented herein by ")
        r7.font.name = "Times New Roman"
        r7.font.size = Pt(12)
        r8 = p2.add_run(partner_rep_name + (f", {partner_rep_title}" if partner_rep_title else "") + ", ")
        r8.bold = True
        r8.font.name = "Times New Roman"
        r8.font.size = Pt(12)
    r9 = p2.add_run(f'hereinafter referred to as the "{partner_short_name}."')
    r9.font.name = "Times New Roman"
    r9.font.size = Pt(12)

    _add_body(doc, f'Collectively, {ISPSC_SHORT_NAME} and {partner_short_name} are referred to as the "Parties."', indent=True)

    _add_heading(doc, "-WITNESSETH, that:", size=12, space_before=6, space_after=10)

    # ---------- WHEREAS clauses ----------
    whereas_items = _lines(data.get("whereas_clauses", "")) or [
        f"{ISPSC_SHORT_NAME}, through its Office of the Vice President for Research, Development, "
        f"and Extension (OVPRDE), approved and implemented the extension project entitled "
        f'"{project_title}," as detailed in the approved proposal made an integral part of this Agreement;',
        f"the Parties recognize the importance of a formal partnership for the implementation, "
        f"monitoring, and sustainability of the said project;",
    ]
    for clause in whereas_items:
        _add_body(doc, clause, bold_lead="WHEREAS, ")

    _add_body(
        doc,
        "for and in consideration of the foregoing premises, the Parties hereby agree as follows:",
        bold_lead="NOW, THEREFORE, ",
        indent=True,
    )

    # ---------- OBJECTIVES ----------
    _add_section_title(doc, "OBJECTIVES")
    objectives = _lines(data.get("objectives", "")) or [
        f"To establish a formal partnership for the implementation of the project;",
        "To define the roles, responsibilities, and obligations of both Parties;",
        "To ensure compliance with applicable laws, standards, and reporting requirements.",
    ]
    _add_numbered_paragraphs(doc, objectives)

    # ---------- OBLIGATIONS OF THE PARTIES ----------
    _add_section_title(doc, "OBLIGATIONS OF THE PARTIES")

    p_a = doc.add_paragraph()
    p_a.paragraph_format.space_before = Pt(6)
    p_a.paragraph_format.space_after = Pt(6)
    r_a = p_a.add_run(f"A. Obligations of {ISPSC_SHORT_NAME}")
    r_a.bold = True
    r_a.font.name = "Times New Roman"
    r_a.font.size = Pt(12)

    obligations_ispsc = _lines(data.get("obligations_ispsc", "")) or [
        f"Implement the project in coordination with {partner_short_name};",
        "Provide technical orientation and capacity-building support to designated personnel;",
        "Ensure compliance with the Data Privacy Act of 2012 (R.A. 10173) and other applicable standards.",
    ]
    _add_numbered_paragraphs(doc, obligations_ispsc)

    p_b = doc.add_paragraph()
    p_b.paragraph_format.space_before = Pt(6)
    p_b.paragraph_format.space_after = Pt(6)
    r_b = p_b.add_run(f"B. Obligations of {partner_short_name}")
    r_b.bold = True
    r_b.font.name = "Times New Roman"
    r_b.font.size = Pt(12)

    obligations_partner = _lines(data.get("obligations_partner", "")) or [
        f"Deploy and support the project as an official partner of {ISPSC_SHORT_NAME};",
        "Designate a lead implementing unit and provide the necessary support for operation and maintenance;",
        "Ensure that data privacy and security measures are in place consistent with R.A. 10173 and other relevant laws.",
    ]
    _add_numbered_paragraphs(doc, obligations_partner)

    # ---------- INTELLECTUAL PROPERTY ----------
    _add_section_title(doc, "INTELLECTUAL PROPERTY")
    ip_ownership_text = (data.get("ip_ownership_text") or "").strip() or (
        f"All outputs, software, materials, and related intellectual property developed under this project, "
        f"including all documentation, are owned by {ISPSC_SHORT_NAME} under Philippine copyright law and the "
        f"Intellectual Property Code of the Philippines (R.A. No. 8293)."
    )
    _add_numbered_paragraphs(doc, [ip_ownership_text], start=1)

    license_years = data.get("ip_license_years") or "3"
    license_terms = []
    if data.get("ip_license_royalty_free", True):
        license_terms.append("royalty-free")
    if data.get("ip_license_exclusive", True):
        license_terms.append("exclusive")
    if data.get("ip_license_irrevocable", True):
        license_terms.append("irrevocable")
    terms_text = ", ".join(license_terms) if license_terms else "non-exclusive"

    ip_license_text = (
        f"{ISPSC_SHORT_NAME} grants {partner_short_name} a {license_years}-year, {terms_text} license "
        f"to use and operate the project outputs for its intended purpose, subject to renewal or further "
        f"agreement by the Parties."
    )
    _add_numbered_paragraphs(doc, [ip_license_text], start=2)

    # ---------- TERM AND EFFECTIVITY ----------
    _add_section_title(doc, "TERM AND EFFECTIVITY")
    term_years = data.get("term_years") or "3"
    _add_body(
        doc,
        f"This MOA shall take effect upon signing and shall remain in force for {term_years} year(s), "
        f"renewable upon mutual agreement.",
        indent=True,
    )

    # ---------- FUNDING ----------
    _add_section_title(doc, "FUNDING")
    funding_ispsc = (data.get("funding_ispsc") or "").strip() or (
        f"{ISPSC_SHORT_NAME} shall cover costs related to project development and initial implementation "
        f"as approved in its project budget."
    )
    funding_partner = (data.get("funding_partner") or "").strip() or (
        f"{partner_short_name} shall shoulder expenses for maintenance, upgrades requested beyond the "
        f"approved scope, and operational costs after turnover."
    )
    _add_body(doc, funding_ispsc, indent=True)
    _add_body(doc, funding_partner, indent=True)

    # ---------- DATA PRIVACY AND CONFIDENTIALITY ----------
    _add_section_title(doc, "DATA PRIVACY AND CONFIDENTIALITY")
    _add_body(doc, (data.get("data_privacy_text") or "").strip() or DEFAULT_DATA_PRIVACY_TEXT, indent=True)

    # ---------- AMENDMENTS ----------
    _add_section_title(doc, "AMENDMENTS")
    _add_body(doc, (data.get("amendments_text") or "").strip() or DEFAULT_AMENDMENTS_TEXT, indent=True)

    # ---------- TERMINATION ----------
    _add_section_title(doc, "TERMINATION")
    notice_days = data.get("termination_notice_days") or "30"
    _add_body(
        doc,
        f"Either Party may terminate this Agreement upon thirty ({notice_days}) days' written notice, "
        f"subject to the settlement of obligations incurred prior to termination.",
        indent=True,
    )

    # ---------- MISCELLANEOUS ----------
    _add_section_title(doc, "MISCELLANEOUS")
    _add_body(doc, (data.get("misc_text") or "").strip() or DEFAULT_MISC_TEXT, indent=True)

    _add_body(
        doc,
        "the Parties have hereunto affixed their signatures on the date and place indicated below.",
        bold_lead="IN WITNESS WHEREOF, ",
        indent=True,
    )

    # ---------- Signature blocks ----------
    doc.add_paragraph()
    _signature_block_table(
        doc,
        left_title=ISPSC_FULL_NAME,
        left_name=ispsc_president,
        left_role="SUC President III",
        right_title=partner_name,
        right_name=partner_rep_name,
        right_role=partner_rep_title,
    )

    doc.add_paragraph()
    p_wit = doc.add_paragraph()
    r_wit = p_wit.add_run("Signed in the presence of:")
    r_wit.font.name = "Times New Roman"
    r_wit.font.size = Pt(12)

    ispsc_witnesses = [
        (_resolve_ispsc_signatory(proposal, "DIRECTOR_EXTENSION"), "Director for Extension"),
        (_resolve_ispsc_signatory(proposal, "VPRDE"), "Vice President for Research, Development, and Extension"),
    ]
    partner_witnesses = []
    w1_name = (data.get("partner_witness1_name") or "").strip()
    w1_role = (data.get("partner_witness1_title") or "").strip()
    w2_name = (data.get("partner_witness2_name") or "").strip()
    w2_role = (data.get("partner_witness2_title") or "").strip()
    if w1_name or w1_role:
        partner_witnesses.append((w1_name, w1_role))
    if w2_name or w2_role:
        partner_witnesses.append((w2_name, w2_role))

    all_witnesses = [w for w in ispsc_witnesses if w[0]] + partner_witnesses
    _witness_table(doc, all_witnesses)

    # ---------- ACKNOWLEDGEMENT (notarial page) ----------
    doc.add_page_break()
    _add_heading(doc, "ACKNOWLEDGEMENT", size=13, space_before=0)

    for line in ["REPUBLIC OF THE PHILIPPINES", "PROVINCE/CITY OF _______________"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(line + " )")
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    doc.add_paragraph()

    _add_body(
        doc,
        f'BEFORE ME, this day, personally appeared {ispsc_president or "_______________________"} and '
        f'{partner_rep_name or "_______________________"}, who are personally known to me, and to me '
        f"known to be the same persons who executed the foregoing Memorandum of Agreement, and they "
        f"acknowledged to me that they affixed their signatures for the purpose stated in the instrument "
        f"as their free and voluntary act and deed.",
        bold_lead="",
        indent=True,
    )

    _add_body(
        doc,
        "WITNESS MY HAND AND SEAL on the date and place above written.",
        indent=True,
    )

    doc.add_paragraph()
    doc.add_paragraph()

    for line in ["Doc. No. _____", "Page No. _____", "Book No. _____", "Series of _____"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer